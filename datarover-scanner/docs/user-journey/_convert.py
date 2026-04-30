#!/usr/bin/env python3
"""Convert DataRover-User-Journey.md to DOCX and PDF.

Azerbaijani-friendly: uses Calibri/DejaVuSans for full Unicode coverage.
Minimal markdown subset (H1-H4, bold-in-cell, bullet lists, pipe tables, hr).
"""

import os
import re
from pathlib import Path

SRC = Path(__file__).parent / "DataRover-User-Journey.md"
OUT_DOCX = Path(__file__).parent / "DataRover-User-Journey.docx"
OUT_PDF = Path(__file__).parent / "DataRover-User-Journey.pdf"


def parse_blocks(md_text):
    """Parse markdown into a list of ('type', payload) blocks."""
    lines = md_text.split("\n")
    blocks = []
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code_buf)))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip() == "---":
            blocks.append(("hr", None))
            i += 1
            continue

        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if img_match:
            alt, path = img_match.group(1), img_match.group(2)
            blocks.append(("image", (alt, path)))
            i += 1
            continue

        bq_match = re.match(r"^>\s*(.*)$", line)
        if bq_match:
            buf = [bq_match.group(1)]
            i += 1
            while i < len(lines):
                m2 = re.match(r"^>\s*(.*)$", lines[i])
                if m2:
                    buf.append(m2.group(1))
                    i += 1
                else:
                    break
            blocks.append(("quote", " ".join(buf).strip()))
            continue
        if line.startswith("#### "):
            blocks.append(("h4", line[5:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-\|]+\|\s*$", lines[i + 1]):
            rows = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header)
            j = i + 2
            while j < len(lines) and lines[j].startswith("|"):
                cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(cells)
                j += 1
            blocks.append(("table", rows))
            i = j
            continue

        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            blocks.append(("ul", items))
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append(("ol", items))
            continue

        if line.strip() == "":
            blocks.append(("blank", None))
            i += 1
            continue

        para = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != "" and not lines[i].startswith(
            ("#", "|", "- ", "* ", "```", "---")
        ) and not re.match(r"^\s*\d+\.\s+", lines[i]):
            para.append(lines[i])
            i += 1
        blocks.append(("p", " ".join(para)))

    return blocks


INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE = re.compile(r"`([^`]+)`")
INLINE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def md_inline_to_runs(text):
    """Split inline text into (text, is_bold, is_code) runs. Links become 'text'."""
    text = INLINE_LINK.sub(r"\1", text)
    runs = []
    pos = 0
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        chunk = m.group(0)
        if chunk.startswith("**"):
            runs.append((chunk[2:-2], True, False))
        else:
            runs.append((chunk[1:-1], False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return runs


# =====================================================================
# DOCX generation
# =====================================================================
def build_docx(blocks, out_path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for h, sz, color in [
        ("Heading 1", 20, RGBColor(0x1A, 0x56, 0xDB)),
        ("Heading 2", 16, RGBColor(0x1E, 0x40, 0xAF)),
        ("Heading 3", 13, RGBColor(0x23, 0x73, 0xCC)),
        ("Heading 4", 11, RGBColor(0x47, 0x55, 0x69)),
    ]:
        style = doc.styles[h]
        style.font.name = "Calibri"
        style.font.size = Pt(sz)
        style.font.bold = True
        style.font.color.rgb = color

    def add_runs(paragraph, text):
        for chunk, bold, code in md_inline_to_runs(text):
            run = paragraph.add_run(chunk)
            run.bold = bold
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)

    def shade_cell(cell, color_hex):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)

    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for kind, payload in blocks:
        if kind == "h1":
            p = doc.add_paragraph(payload, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(payload, style="Heading 2")
        elif kind == "h3":
            doc.add_paragraph(payload, style="Heading 3")
        elif kind == "h4":
            doc.add_paragraph(payload, style="Heading 4")
        elif kind == "image":
            alt, path = payload
            img_path = Path(__file__).parent / path
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img_path), width=Inches(6.3))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(alt)
                cap_run.italic = True
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            else:
                p = doc.add_paragraph(f"[Şəkil: {alt} — {path}]")
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_runs(p, payload)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        elif kind == "p":
            p = doc.add_paragraph()
            add_runs(p, payload)
        elif kind == "ul":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, item)
        elif kind == "ol":
            for item in payload:
                p = doc.add_paragraph(style="List Number")
                add_runs(p, item)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "hr":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif kind == "table":
            rows = payload
            n_cols = max(len(r) for r in rows)
            for r in rows:
                while len(r) < n_cols:
                    r.append("")
            t = doc.add_table(rows=len(rows), cols=n_cols)
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = t.rows[ri].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_runs(p, cell_text)
                    if ri == 0:
                        shade_cell(cell, "1E40AF")
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            doc.add_paragraph()
        elif kind == "blank":
            pass

    doc.save(out_path)
    print(f"DOCX yazıldı: {out_path}")


# =====================================================================
# PDF generation
# =====================================================================
def build_pdf(blocks, out_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
        HRFlowable,
        ListFlowable,
        ListItem,
        Image,
        KeepTogether,
    )
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    base_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    mono_font = "Courier"
    for candidate in (
        r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
    ):
        if os.path.exists(candidate):
            try:
                pdfmetrics.registerFont(TTFont("AZBody", candidate))
                base_font = "AZBody"
                bold_path = candidate.replace(".ttf", "b.ttf")
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont("AZBodyBold", bold_path))
                    bold_font = "AZBodyBold"
                else:
                    bold_font = base_font
                break
            except Exception:
                continue
    if os.path.exists(r"C:\Windows\Fonts\consola.ttf"):
        try:
            pdfmetrics.registerFont(TTFont("AZMono", r"C:\Windows\Fonts\consola.ttf"))
            mono_font = "AZMono"
        except Exception:
            pass

    def esc(s):
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    def md_to_html(text):
        text = esc(text)
        text = re.sub(r"\*\*(.+?)\*\*", r'<b>\1</b>', text)
        text = re.sub(r"`([^`]+)`", lambda m: f'<font name="{mono_font}" size="8">{m.group(1)}</font>', text)
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'\1', text)
        return text

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName=base_font,
        fontSize=9.5,
        leading=13,
        spaceAfter=4,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName=bold_font,
        fontSize=20,
        textColor=colors.HexColor("#1A56DB"),
        spaceAfter=12,
        spaceBefore=10,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=15,
        textColor=colors.HexColor("#1E40AF"),
        spaceAfter=8,
        spaceBefore=14,
    )
    h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontName=bold_font,
        fontSize=12.5,
        textColor=colors.HexColor("#2373CC"),
        spaceAfter=6,
        spaceBefore=10,
    )
    h4 = ParagraphStyle(
        "H4",
        parent=styles["Heading4"],
        fontName=bold_font,
        fontSize=11,
        textColor=colors.HexColor("#475569"),
        spaceAfter=4,
        spaceBefore=8,
    )
    code_style = ParagraphStyle(
        "Code",
        parent=body,
        fontName=mono_font,
        fontSize=8,
        leading=10,
        backColor=colors.HexColor("#F3F4F6"),
        borderPadding=6,
        leftIndent=6,
        rightIndent=6,
        spaceAfter=6,
    )
    cell_style = ParagraphStyle(
        "Cell",
        parent=body,
        fontSize=8.5,
        leading=11,
        spaceAfter=0,
    )
    cell_head = ParagraphStyle(
        "CellHead",
        parent=cell_style,
        fontName=bold_font,
        textColor=colors.white,
    )

    caption_style = ParagraphStyle(
        "Caption",
        parent=body,
        alignment=1,
        fontSize=8.5,
        textColor=colors.HexColor("#64748B"),
        spaceBefore=2,
        spaceAfter=10,
        fontName=base_font,
    )
    quote_style = ParagraphStyle(
        "Quote",
        parent=body,
        leftIndent=14,
        rightIndent=8,
        textColor=colors.HexColor("#475569"),
        borderPadding=6,
        spaceBefore=4,
        spaceAfter=6,
    )

    page_width = A4[0] - 2 * 2 * cm

    story = []
    for kind, payload in blocks:
        if kind == "h1":
            story.append(Paragraph(md_to_html(payload), h1))
        elif kind == "h2":
            story.append(Paragraph(md_to_html(payload), h2))
        elif kind == "h3":
            story.append(Paragraph(md_to_html(payload), h3))
        elif kind == "h4":
            story.append(Paragraph(md_to_html(payload), h4))
        elif kind == "image":
            alt, path = payload
            img_path = Path(__file__).parent / path
            if img_path.exists():
                reader = ImageReader(str(img_path))
                iw, ih = reader.getSize()
                max_w = page_width
                max_h = A4[1] * 0.55
                scale = min(max_w / iw, max_h / ih, 1.0)
                w = iw * scale
                h = ih * scale
                img = Image(str(img_path), width=w, height=h)
                img.hAlign = "CENTER"
                story.append(KeepTogether([
                    img,
                    Paragraph(f"<i>{esc(alt)}</i>", caption_style),
                ]))
            else:
                story.append(Paragraph(f"[Şəkil yoxdur: {esc(alt)} — {esc(path)}]", body))
        elif kind == "quote":
            story.append(Paragraph(f"▎ {md_to_html(payload)}", quote_style))
        elif kind == "p":
            story.append(Paragraph(md_to_html(payload), body))
        elif kind == "ul":
            items = [ListItem(Paragraph(md_to_html(x), body), leftIndent=12) for x in payload]
            story.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=14))
            story.append(Spacer(1, 4))
        elif kind == "ol":
            items = [ListItem(Paragraph(md_to_html(x), body), leftIndent=12) for x in payload]
            story.append(ListFlowable(items, bulletType="1", leftIndent=14))
            story.append(Spacer(1, 4))
        elif kind == "code":
            for line in payload.split("\n"):
                story.append(Paragraph(esc(line).replace(" ", "&nbsp;") or "&nbsp;", code_style))
        elif kind == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC"), spaceBefore=4, spaceAfter=6))
        elif kind == "table":
            rows = payload
            n_cols = max(len(r) for r in rows)
            for r in rows:
                while len(r) < n_cols:
                    r.append("")
            data = []
            for ri, row in enumerate(rows):
                style = cell_head if ri == 0 else cell_style
                data.append([Paragraph(md_to_html(c), style) for c in row])

            col_widths = [page_width / n_cols] * n_cols
            t = Table(data, colWidths=col_widths, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 6))
        elif kind == "blank":
            story.append(Spacer(1, 3))

    pdf_doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        title="MIM — İstifadəçi Səyahəti",
        author="MIM",
    )
    pdf_doc.build(story)
    print(f"PDF yazıldı: {out_path}")


if __name__ == "__main__":
    md = SRC.read_text(encoding="utf-8")
    blocks = parse_blocks(md)
    print(f"Blok sayı: {len(blocks)}")
    build_docx(blocks, OUT_DOCX)
    build_pdf(blocks, OUT_PDF)
