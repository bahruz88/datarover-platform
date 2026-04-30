#!/usr/bin/env python3
"""Detailed technical spec for Glossary, Data Catalog, Reporting modules.
Each API endpoint includes the exact SELECT/INSERT/UPDATE/DELETE SQL from backend.php."""

import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)
doc.styles['Heading 1'].font.size = Pt(22); doc.styles['Heading 1'].font.color.rgb = RGBColor(0x1a,0x56,0xdb); doc.styles['Heading 1'].font.bold = True
doc.styles['Heading 2'].font.size = Pt(15); doc.styles['Heading 2'].font.color.rgb = RGBColor(0x1a,0x56,0xdb)
doc.styles['Heading 3'].font.size = Pt(12); doc.styles['Heading 3'].font.color.rgb = RGBColor(0x33,0x33,0x33)
doc.styles['Heading 4'].font.size = Pt(11); doc.styles['Heading 4'].font.color.rgb = RGBColor(0x55,0x55,0x55)

def p(text, bold=False, size=10):
    para = doc.add_paragraph()
    r = para.add_run(text); r.font.size = Pt(size); r.bold = bold

def code(text):
    para = doc.add_paragraph()
    r = para.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    pf = para.paragraph_format; pf.left_indent = Cm(0.4)
    pf.space_before = Pt(2); pf.space_after = Pt(4)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for pa in c.paragraphs:
            for r in pa.runs: r.bold = True; r.font.size = Pt(9)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for pa in c.paragraphs:
                for r in pa.runs: r.font.size = Pt(8.5)
    doc.add_paragraph()

def pb(): doc.add_page_break()

def endpoint_block(num, method, path, purpose, auth, req_desc, sql_blocks, params_py=None,
                   response=None, curl=None, frontend=None, errors=None):
    """Render one endpoint with full SQL list."""
    doc.add_heading(f'{num}  {method} {path}', level=3)
    tbl(['Xüsusiyyət','Dəyər'],[
        ['Action','?action=' + path.split("?action=")[1] if "?action=" in path else path],
        ['HTTP Method', method],
        ['Təyinat', purpose],
        ['Auth', auth],
    ])
    if req_desc:
        p('Sorğu body / parametr:', bold=True)
        code(req_desc)
    p('SQL sorğuları (backend.php-dən):', bold=True)
    for title, sql in sql_blocks:
        p(f'• {title}', bold=False)
        code(sql)
    if params_py:
        p('PHP parametr bindinqi:', bold=True)
        code(params_py)
    if response:
        p('Cavab nümunəsi:', bold=True)
        code(response)
    if errors:
        p('Xəta halları:', bold=True)
        tbl(['HTTP','Səbəb'], errors)
    if curl:
        p('cURL nümunəsi:', bold=True)
        code(curl)
    if frontend:
        p(f'Frontend binding: {frontend}')

# ============================================================
# COVER
# ============================================================
for _ in range(6): doc.add_paragraph()
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = para.add_run('Texniki Spesifikasiya — Detallı'); r.font.size = Pt(32); r.bold=True; r.font.color.rgb = RGBColor(0x1a,0x56,0xdb)
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = para.add_run('Business Glossary · Data Catalog · Reporting Catalog'); r.font.size = Pt(17); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = para.add_run('Hər API üçün tam SELECT · INSERT · UPDATE · DELETE sorğuları'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x77,0x77,0x77)
doc.add_paragraph()
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = para.add_run(f'v56  |  {datetime.date.today().strftime("%B %Y")}'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x99,0x99,0x99)
pb()

# ============================================================
# INTRO
# ============================================================
doc.add_heading('Giriş', level=1)
p('Bu sənəd üç əsas modulun — Business Glossary, Data Catalog və Reporting Catalog — tam backend implementasiyasını təsvir edir. Hər API endpoint-i üçün backend.php faylından çıxarılmış real SQL sorğuları (SELECT, INSERT, UPDATE, DELETE) göstərilir. Sənəd yeni qoşulan backend developer üçün nəzərdə tutulub — yalnız bunu oxumaqla modulları patch etməyə və ya genişləndirməyə başlaya biləcək.')

doc.add_heading('Ümumi Texniki Struktur', level=2)
tbl(['Aspekt','Detal'],[
    ['Runtime','PHP 8.x (XAMPP)'],
    ['DB','MySQL 8.x, InnoDB, utf8mb4_unicode_ci'],
    ['Tək giriş nöqtəsi','backend.php'],
    ['Routing','?action=<endpoint_name>&id=<optional>'],
    ['HTTP method','$_SERVER[\'REQUEST_METHOD\'] ilə oxunur'],
    ['DB bağlantısı','db() PDO singleton (sətir 28–46)'],
    ['Input parse','getInput() → json_decode($_POST body)'],
    ['Cavab','respond($data) → JSON success envelope'],
    ['Xəta','error($msg, $code) → JSON error envelope'],
    ['Prepared statement','Bütün sorğular PDO positional ? placeholder istifadə edir'],
    ['Transaction','Explicit START/COMMIT yoxdur — autocommit rejimi'],
    ['Timezone','Asia/Baku (UTC+4)'],
])

doc.add_heading('Əsas Helper Funksiyalar', level=2)
code('''// sətir 28–46
function db() {
  static $pdo = null;
  if ($pdo === null) {
    $pdo = new PDO("mysql:host=localhost;dbname=datarover;charset=utf8mb4",
                   "root", "", [
      PDO::ATTR_ERRMODE => PDO::ERRMODE_EXCEPTION,
      PDO::ATTR_DEFAULT_FETCH_MODE => PDO::FETCH_ASSOC,
    ]);
  }
  return $pdo;
}

function getInput()    { return json_decode(file_get_contents('php://input'), true) ?: []; }
function respond($d)   { echo json_encode(['success'=>true, 'data'=>$d]); exit; }
function error($m,$c=400) { http_response_code($c); echo json_encode(['success'=>false,'error'=>$m]); exit; }''')

doc.add_heading('Cavab Formatı', level=2)
code('''// uğur
{"success": true, "data": { ... }}

// xəta
{"success": false, "error": "mesaj"}''')

pb()

# ============================================================
# MODULE 1: BUSINESS GLOSSARY
# ============================================================
doc.add_heading('1. Modul: Business Glossary', level=1)

doc.add_heading('1.1 Overview', level=2)
p('Business Glossary biznes terminlərini (KPI, metrik, entity), onların tərifləri, sahibləri, fiziki atributları və keyfiyyət qaydaları ilə birlikdə saxlayır. Hər term workflow status-da olur: draft → under_review → approved (və ya rejected). Dəyişiklik tarixçəsi glossary_term_history cədvəlində saxlanılır. Silinmə soft-delete prinsipi ilə həyata keçirilir (deleted_at, deleted_by sahələri).')

doc.add_heading('1.2 Cədvəllər (Data Model)', level=2)

doc.add_heading('1.2.1 glossary_terms', level=3)
code('''CREATE TABLE glossary_terms (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255) NOT NULL UNIQUE,
  abbreviation VARCHAR(100),
  definition TEXT,
  domain VARCHAR(255),
  data_type VARCHAR(100),
  example TEXT,
  formula TEXT,
  business_logic TEXT,
  technical_description TEXT,
  owner VARCHAR(255),
  stewards JSON,
  security_classification VARCHAR(100),
  physical_attributes JSON,
  quality_rules JSON,
  synonyms JSON,
  related_terms JSON,
  source_system VARCHAR(255),
  notes TEXT,
  history JSON,
  status VARCHAR(50) DEFAULT 'draft',
  reject_reason TEXT,
  process_comments TEXT,
  deleted_at DATETIME NULL,
  deleted_by VARCHAR(255) NULL,
  restored_at DATETIME NULL,
  restored_by VARCHAR(255) NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  INDEX idx_name (name),
  INDEX idx_status (status),
  INDEX idx_domain (domain),
  INDEX idx_deleted (deleted_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')
p('Sahə izahları:')
tbl(['Sahə','Tip','Qeyd'],[
    ['name','VARCHAR(255)','Unikal, məcburi'],
    ['stewards','JSON','["steward1@x.com","steward2@x.com"]'],
    ['physical_attributes','JSON','[{"layer":"dwh","table":"x","column":"y"}]'],
    ['quality_rules','JSON','["RULE-001","RULE-042"] — data_quality_rules.rule_id ref'],
    ['synonyms','JSON','Sinonim termlər array'],
    ['related_terms','JSON','Bağlı term ID-ləri'],
    ['history','JSON','UI cache üçün (əsas tarixçə ayrı cədvəldədir)'],
    ['status','VARCHAR(50)','draft, under_review, approved, rejected, deleted'],
    ['deleted_at','DATETIME','Soft delete markeri'],
])

doc.add_heading('1.2.2 glossary_term_history', level=3)
code('''CREATE TABLE glossary_term_history (
  id INT PRIMARY KEY AUTO_INCREMENT,
  term_id INT NOT NULL,
  action VARCHAR(100),
  comment TEXT,
  user VARCHAR(255),
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE CASCADE,
  INDEX idx_term_id (term_id),
  INDEX idx_created_at (created_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('1.2.3 term_physical_attributes', level=3)
code('''CREATE TABLE term_physical_attributes (
  id INT PRIMARY KEY AUTO_INCREMENT,
  term_id INT NOT NULL,
  layer VARCHAR(255),
  table_name VARCHAR(255),
  column_name VARCHAR(255),
  attribute_value TEXT,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE CASCADE,
  INDEX idx_term (term_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('1.3 ER Diagram', level=2)
code('''erDiagram
    glossary_terms ||--o{ glossary_term_history : changes
    glossary_terms ||--o{ term_physical_attributes : maps
    glossary_terms }o--|| domains : belongs_to
    glossary_terms ||--o{ report_terms : used_in
    catalog_columns }o--o{ glossary_terms : semantic_link''')

doc.add_heading('1.4 API Endpoint-ləri', level=2)

# ----- 1.4.1 GET glossary_terms (list) -----
endpoint_block(
    '1.4.1',
    'GET',
    'backend.php?action=glossary_terms',
    'Term-ləri filter ilə sadalayır (q, status, domain). Hər term-ə tarixçə sayını əlavə edir.',
    'Session (optional filter)',
    '''Query parametrlər:
  q       (string)  — name və definition üzrə axtarış
  status  (string)  — draft | under_review | approved | rejected | deleted
  domain  (string)  — domain filter''',
    [
        ('Əsas SELECT (history_count subquery ilə)',
'''SELECT t.*,
  (SELECT COUNT(*) FROM glossary_term_history WHERE term_id = t.id) AS history_count
FROM glossary_terms t
WHERE 1=1'''),
        ('+ search şərti (əgər q verilib)',
'''AND (t.name LIKE ? OR t.definition LIKE ?)
-- params: "%$q%", "%$q%"'''),
        ('+ status şərti',
'''AND t.status = ?
-- param: $status'''),
        ('+ domain şərti',
'''AND t.domain = ?
-- param: $domain'''),
        ('+ sıralama',
'ORDER BY t.created_at DESC;'),
    ],
    params_py='''// PHP dinamik sorğu tikir
$sql = "SELECT t.*, (SELECT COUNT(*) FROM glossary_term_history WHERE term_id = t.id) as history_count
        FROM glossary_terms t WHERE 1=1";
$params = [];
if ($search) { $sql .= " AND (t.name LIKE ? OR t.definition LIKE ?)";
               $params[] = "%$search%"; $params[] = "%$search%"; }
if ($status) { $sql .= " AND t.status = ?"; $params[] = $status; }
if ($domain) { $sql .= " AND t.domain = ?"; $params[] = $domain; }
$sql .= " ORDER BY t.created_at DESC";
$stmt = db()->prepare($sql); $stmt->execute($params);''',
    response='''{
  "success": true,
  "data": [
    {
      "id": 12,
      "name": "Customer Lifetime Value",
      "abbreviation": "CLV",
      "definition": "Müştərinin...",
      "domain": "Sales",
      "status": "approved",
      "owner": "ceo@corp.com",
      "history_count": 7,
      "created_at": "2024-11-10 09:22:00"
    }
  ]
}''',
    curl='''curl "http://localhost/datarover/backend.php?action=glossary_terms&status=approved&domain=Sales" \\
  -H "Authorization: Bearer $TOKEN"''',
    frontend='app.js → loadGlossary() (3567–3705)'
)

# ----- 1.4.2 GET single -----
endpoint_block(
    '1.4.2',
    'GET',
    'backend.php?action=glossary_terms&id={id}',
    'Bir termi tarixçəsi ilə birlikdə gətirir.',
    'Session',
    'URL parametri: id (integer, məcburi)',
    [
        ('Term detalı',
'''SELECT * FROM glossary_terms WHERE id = ?;
-- param: $id'''),
        ('Termin tarixçəsi',
'''SELECT * FROM glossary_term_history
WHERE term_id = ?
ORDER BY created_at DESC;
-- param: $id'''),
    ],
    params_py='''$stmt = db()->prepare("SELECT * FROM glossary_terms WHERE id = ?");
$stmt->execute([$id]);
$term = $stmt->fetch();
if (!$term) error('Term not found', 404);

$histStmt = db()->prepare("SELECT * FROM glossary_term_history WHERE term_id = ? ORDER BY created_at DESC");
$histStmt->execute([$id]);
$term['history'] = $histStmt->fetchAll();''',
    errors=[['404','Term tapılmadı']],
    curl='curl "http://localhost/datarover/backend.php?action=glossary_terms&id=12" -H "Authorization: Bearer $TOKEN"'
)

# ----- 1.4.3 POST create -----
endpoint_block(
    '1.4.3',
    'POST',
    'backend.php?action=glossary_terms',
    'Yeni term yaradır + history cədvəlinə "Yaradıldı" qeydi INSERT edir.',
    'Session (glossary.create)',
    '''JSON body (bütün sahələr):
{
  "name": "string (required, unique)",
  "abbreviation": "CLV",
  "definition": "Tam təsvir...",
  "domain": "Sales",
  "dataType": "DECIMAL(10,2)",
  "example": "€1500",
  "formula": "SUM(...)",
  "businessLogic": "...",
  "technicalDescription": "...",
  "owner": "owner@x.com",
  "stewards": ["s1@x.com", "s2@x.com"],
  "securityClassification": "internal",
  "physicalAttributes": [{"layer":"dwh","table":"x","column":"y"}],
  "qualityRules": ["RULE-001"],
  "synonyms": ["LTV"],
  "relatedTerms": [15, 27],
  "sourceSystem": "ERP",
  "notes": "...",
  "status": "draft",
  "rejectReason": null,
  "processComments": null,
  "comment": "İlk yaradılma qeydi"
}''',
    [
        ('Əsas INSERT (22 sahə + NOW() × 2)',
'''INSERT INTO glossary_terms (
  name, abbreviation, definition, domain, data_type,
  example, formula, business_logic, technical_description, owner,
  stewards, security_classification, physical_attributes, quality_rules,
  synonyms, related_terms, source_system, notes, history,
  status, reject_reason, process_comments,
  created_at, updated_at
) VALUES (?, ?, ?, ?, ?,
          ?, ?, ?, ?, ?,
          ?, ?, ?, ?,
          ?, ?, ?, ?, ?,
          ?, ?, ?,
          NOW(), NOW());'''),
        ('History qeydi (yaradıldı)',
'''INSERT INTO glossary_term_history
  (term_id, action, comment, user, created_at)
VALUES (?, ?, ?, ?, NOW());
-- params: $newId, 'Yaradıldı', $data['comment'] ?? null, $data['owner'] ?? 'System' '''),
    ],
    params_py='''$newId = db()->lastInsertId();

// JSON-encoded dəyərlər
$stewards           = is_array($data['stewards'] ?? null)           ? json_encode($data['stewards'])           : ($data['stewards'] ?? null);
$physicalAttributes = is_array($data['physicalAttributes'] ?? null) ? json_encode($data['physicalAttributes']) : null;
$qualityRules       = is_array($data['qualityRules'] ?? null)       ? json_encode($data['qualityRules'])       : null;
$synonyms           = is_array($data['synonyms'] ?? null)           ? json_encode($data['synonyms'])           : null;
$relatedTerms       = is_array($data['relatedTerms'] ?? null)       ? json_encode($data['relatedTerms'])       : null;
$history            = is_array($data['history'] ?? null)            ? json_encode($data['history'])            : null;

$stmt->execute([
    $data['name'],
    $data['abbreviation'] ?? '', $data['definition'] ?? '',
    $data['domain'] ?? '',        $data['dataType'] ?? '',
    $data['example'] ?? '',       $data['formula'] ?? '',
    $data['businessLogic'] ?? '', $data['technicalDescription'] ?? '',
    $data['owner'] ?? '',
    $stewards, $data['securityClassification'] ?? null,
    $physicalAttributes, $qualityRules,
    $synonyms, $relatedTerms,
    $data['sourceSystem'] ?? '', $data['notes'] ?? '', $history,
    $data['status'] ?? 'draft',
    $data['rejectReason'] ?? null, $data['processComments'] ?? null
]);''',
    response='{"success": true, "data": {"id": 13, "message": "Term created"}}',
    errors=[['400','name boş'], ['409','name artıq mövcud (unique violation)']],
    curl='''curl -X POST "http://localhost/datarover/backend.php?action=glossary_terms" \\
  -H "Authorization: Bearer $TOKEN" \\
  -H "Content-Type: application/json" \\
  -d '{"name":"CLV","definition":"Customer LTV","domain":"Sales","status":"draft"}' ''',
    frontend='app.js → saveTermToDB(term, isNew=true) (4206–4297)'
)

# ----- 1.4.4 PUT update -----
endpoint_block(
    '1.4.4',
    'PUT',
    'backend.php?action=glossary_terms&id={id}',
    'Mövcud termi yeniləyir. 22 sahəni tam (overwrite) yeniləyir.',
    'Session (glossary.edit)',
    'POST body ilə eyni strukturda JSON.',
    [
        ('Əsas UPDATE (22 sahə + NOW())',
'''UPDATE glossary_terms SET
  name = ?, abbreviation = ?, definition = ?, domain = ?, data_type = ?,
  example = ?, formula = ?, business_logic = ?, technical_description = ?, owner = ?,
  stewards = ?, security_classification = ?, physical_attributes = ?, quality_rules = ?,
  synonyms = ?, related_terms = ?, source_system = ?, notes = ?, history = ?,
  status = ?, reject_reason = ?, process_comments = ?,
  updated_at = NOW()
WHERE id = ?;'''),
    ],
    params_py='''$stmt->execute([
    $data['name'],
    $data['abbreviation'] ?? '', $data['definition'] ?? '',
    $data['domain'] ?? '',        $data['dataType'] ?? '',
    $data['example'] ?? '',       $data['formula'] ?? '',
    $data['businessLogic'] ?? '', $data['technicalDescription'] ?? '',
    $data['owner'] ?? '',
    $stewards, $data['securityClassification'] ?? null,
    $physicalAttributes, $qualityRules,
    $synonyms, $relatedTerms,
    $data['sourceSystem'] ?? '', $data['notes'] ?? '', $history,
    $data['status'] ?? 'draft',
    $data['rejectReason'] ?? null, $data['processComments'] ?? null,
    $id
]);''',
    response='{"success": true, "data": {"message": "Term updated"}}',
    curl='''curl -X PUT "http://localhost/datarover/backend.php?action=glossary_terms&id=13" \\
  -H "Authorization: Bearer $TOKEN" \\
  -H "Content-Type: application/json" \\
  -d '{"name":"CLV","definition":"Yenilənmiş","status":"under_review"}' ''',
    frontend='app.js → updateTermInDB(index) (4298–4310)'
)

# ----- 1.4.5 DELETE soft -----
endpoint_block(
    '1.4.5',
    'DELETE',
    'backend.php?action=glossary_terms&id={id}',
    'Soft delete — status="deleted", deleted_at=NOW() qoyur. Term fiziki silinmir.',
    'Session (glossary.delete)',
    'Optional JSON body: {"deletedBy":"admin@x.com"}',
    [
        ('Soft delete UPDATE',
'''UPDATE glossary_terms SET
  status = 'deleted',
  deleted_at = NOW(),
  deleted_by = ?
WHERE id = ?;
-- params: $data['deletedBy'] ?? 'System', $id'''),
    ],
    params_py='''$stmt = db()->prepare("UPDATE glossary_terms SET status = 'deleted',
    deleted_at = NOW(), deleted_by = ? WHERE id = ?");
$stmt->execute([$data['deletedBy'] ?? 'System', $id]);''',
    response='{"success": true, "data": {"message": "Term soft-deleted"}}',
    curl='''curl -X DELETE "http://localhost/datarover/backend.php?action=glossary_terms&id=13" \\
  -H "Authorization: Bearer $TOKEN" \\
  -d '{"deletedBy":"admin@corp.com"}' '''
)

# ----- 1.4.6 Hard delete -----
doc.add_heading('1.4.6  DELETE (hard) — internal helper', level=3)
p('Hard delete üçün kaskadlı 3 DELETE icra olunur (tranzaksiyasız):')
code('''DELETE FROM glossary_term_history WHERE term_id = ?;
DELETE FROM term_physical_attributes WHERE term_id = ?;
DELETE FROM glossary_terms WHERE id = ?;
-- hər biri üçün param: $id''')
p('PHP:')
code('''db()->prepare("DELETE FROM glossary_term_history WHERE term_id = ?")->execute([$id]);
db()->prepare("DELETE FROM term_physical_attributes WHERE term_id = ?")->execute([$id]);
db()->prepare("DELETE FROM glossary_terms WHERE id = ?")->execute([$id]);''')
p('Xəbərdarlıq: Explicit transaction yoxdur — orta DELETE-də exception atılsa inconsistent state yarana bilər. Production-da START TRANSACTION / COMMIT əlavə edilməlidir.')

# ----- 1.4.7 GET glossary_history -----
endpoint_block(
    '1.4.7',
    'GET',
    'backend.php?action=glossary_history&id={id}',
    'Termin tam tarixçəsi (yaradılma, status dəyişikliyi, redaktə).',
    'Session',
    'URL param: id (term ID)',
    [
        ('History select',
'''SELECT * FROM glossary_term_history
WHERE term_id = ?
ORDER BY created_at DESC;
-- param: $id'''),
    ],
    response='''{
  "success": true,
  "data": [
    {"id":34,"term_id":13,"action":"status_change","comment":"Approved by steward","user":"s1@x.com","created_at":"..."}
  ]
}'''
)

# ----- 1.4.8 POST glossary_history -----
endpoint_block(
    '1.4.8',
    'POST',
    'backend.php?action=glossary_history',
    'Termə manual tarixçə qeydi əlavə edir (məsələn, kommentariya).',
    'Session',
    '''JSON body:
{
  "term_id": 13,
  "action": "comment",
  "comment": "Review tamamlandı",
  "user": "steward@x.com"
}''',
    [
        ('History insert',
'''INSERT INTO glossary_term_history
  (term_id, action, comment, user, created_at)
VALUES (?, ?, ?, ?, NOW());'''),
    ],
    params_py='''$stmt->execute([
    $data['term_id'],
    $data['action'] ?? '',
    $data['comment'] ?? null,
    $data['user'] ?? 'System'
]);'''
)

# ----- 1.4.9 POST restore_term -----
endpoint_block(
    '1.4.9',
    'POST',
    'backend.php?action=restore_term&id={id}',
    'Soft silinmiş termi bərpa edir. status="draft" (və ya body-dən gələn) qoyulur, deleted_at NULL.',
    'Session (glossary.restore)',
    '''JSON body:
{
  "status": "draft",
  "restoredBy": "admin@x.com"
}''',
    [
        ('Restore UPDATE',
'''UPDATE glossary_terms SET
  status = ?,
  deleted_at = NULL,
  deleted_by = NULL,
  restored_at = NOW(),
  restored_by = ?
WHERE id = ?;'''),
    ],
    params_py='''$stmt->execute([
    $data['status'] ?? 'draft',
    $data['restoredBy'] ?? 'System',
    $id
]);'''
)

doc.add_heading('1.5 Status Workflow', level=2)
tbl(['From','To','Kim','Şərt'],[
    ['draft','under_review','Steward','name + definition boş olmamalı'],
    ['under_review','approved','Gov Manager','Ən azı 1 approver'],
    ['under_review','rejected','Gov Manager','reject_reason məcburi'],
    ['rejected','draft','Steward','Resubmit — dəyişiklik sonra'],
    ['approved','under_review','Steward','Yenidən review üçün'],
    ['*','deleted','Admin / Gov','Soft delete (UPDATE)'],
    ['deleted','draft (və ya digər)','Admin','restore_term endpoint-i'],
])

doc.add_heading('1.6 Permission Matrix', level=2)
tbl(['Action','SUPER_ADMIN','GOV_MGR','STEWARD','ENGINEER','ANALYST','VIEWER'],[
    ['glossary.view','✓','✓','✓','✓','✓','✓'],
    ['glossary.create','✓','✓','✓','—','—','—'],
    ['glossary.edit','✓','✓','✓','—','—','—'],
    ['glossary.approve','✓','✓','—','—','—','—'],
    ['glossary.delete','✓','✓','—','—','—','—'],
    ['glossary.restore','✓','✓','—','—','—','—'],
])

doc.add_heading('1.7 Fayl Strukturu', level=2)
tbl(['Fayl','Sətir aralığı','Məzmun'],[
    ['backend.php','2522–2706','glossary_terms CRUD handler'],
    ['backend.php','3864–3907','restore_term, glossary_history'],
    ['js/app.js','3567–3705','loadGlossary()'],
    ['js/app.js','4206–4310','saveTermToDB, updateTermInDB'],
    ['js/app.js','4903–5168','Workflow: advance, reject, resubmit'],
])

doc.add_heading('1.8 İnteqrasiya', level=2)
tbl(['Modul','Əlaqə'],[
    ['Data Catalog','physical_attributes JSON → catalog_columns ref'],
    ['Data Quality','quality_rules JSON → data_quality_rules.rule_id'],
    ['Reports','report_terms + report_kpis.term_id FK'],
    ['Domain','domain sahəsi domains.name ilə match'],
    ['Governance','Status keçidləri governance_workflow_steps ilə idarə oluna bilər'],
])

pb()

# ============================================================
# MODULE 2: DATA CATALOG
# ============================================================
doc.add_heading('2. Modul: Data Catalog', level=1)

doc.add_heading('2.1 Overview', level=2)
p('Data Catalog platformanın metadata repozitoriyasıdır. Üç səviyyəli iyerarxiya: Layer → Table → Column. Xarici mənbələrdən (external_sources) cədvəllər idxal olunur və live profilləşdirilir. Cascade delete məntiqı tətbiq olunur: layer silinsə, tabelər və sütunlar da silinir.')

doc.add_heading('2.2 Cədvəllər (Data Model)', level=2)

doc.add_heading('2.2.1 catalog_layers', level=3)
code('''CREATE TABLE catalog_layers (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255) UNIQUE NOT NULL,
  icon VARCHAR(50) DEFAULT '📁',
  color VARCHAR(10) DEFAULT '#6366f1',
  description TEXT,
  `order` INT DEFAULT 0,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_order (`order`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('2.2.2 catalog_tables', level=3)
code('''CREATE TABLE catalog_tables (
  id INT PRIMARY KEY AUTO_INCREMENT,
  layer_id INT NOT NULL,
  name VARCHAR(500) NOT NULL,
  description TEXT,
  row_count BIGINT DEFAULT 0,
  owner VARCHAR(255),
  source_id INT NULL,
  schema_name VARCHAR(255),
  object_type VARCHAR(50) DEFAULT 'TABLE',
  is_active BOOLEAN DEFAULT TRUE,
  last_profiled_at DATETIME NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  UNIQUE KEY unique_layer_table (layer_id, name),
  FOREIGN KEY (layer_id) REFERENCES catalog_layers(id) ON DELETE CASCADE,
  FOREIGN KEY (source_id) REFERENCES external_sources(id) ON DELETE SET NULL,
  INDEX idx_name (name)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('2.2.3 catalog_columns', level=3)
code('''CREATE TABLE catalog_columns (
  id INT PRIMARY KEY AUTO_INCREMENT,
  table_id INT NOT NULL,
  name VARCHAR(255) NOT NULL,
  data_type VARCHAR(100) DEFAULT 'VARCHAR',
  description TEXT,
  is_pk BOOLEAN DEFAULT FALSE,
  is_fk BOOLEAN DEFAULT FALSE,
  is_nullable BOOLEAN DEFAULT TRUE,
  icon VARCHAR(50) DEFAULT '📊',
  `order` INT DEFAULT 0,
  profile_stats JSON,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  UNIQUE KEY unique_table_column (table_id, name),
  FOREIGN KEY (table_id) REFERENCES catalog_tables(id) ON DELETE CASCADE,
  INDEX idx_order (`order`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('2.3 ER Diagram', level=2)
code('''erDiagram
    catalog_layers ||--o{ catalog_tables : contains
    catalog_tables ||--o{ catalog_columns : has
    external_sources ||--o{ catalog_tables : sourced_from
    catalog_tables ||--o{ column_mappings : ref
    catalog_tables ||--o{ report_data_sources : ref''')

doc.add_heading('2.4 API Endpoint-ləri', level=2)

# ----- Layers -----
endpoint_block(
    '2.4.1',
    'GET',
    'backend.php?action=catalog_layers',
    'Bütün layer-lərin siyahısı, `order` ASC ilə.',
    'Session (catalog.view)',
    'Parametr yoxdur',
    [
        ('Layer siyahısı',
'''SELECT * FROM catalog_layers ORDER BY `order` ASC;'''),
    ],
    params_py='$stmt = db()->query("SELECT * FROM catalog_layers ORDER BY `order` ASC");',
    response='{"success":true,"data":[{"id":1,"name":"Raw","icon":"📥","color":"#ef4444","order":0}]}',
    curl='curl "http://localhost/datarover/backend.php?action=catalog_layers" -H "Authorization: Bearer $TOKEN"'
)

endpoint_block(
    '2.4.2',
    'POST',
    'backend.php?action=catalog_layers',
    'Yeni layer yaradır.',
    'Session (catalog.create)',
    '''JSON body:
{
  "name": "Mart",
  "icon": "🎯",
  "color": "#10b981",
  "description": "Business-facing marts",
  "order": 3
}''',
    [
        ('INSERT',
'''INSERT INTO catalog_layers
  (name, icon, color, description, `order`)
VALUES (?, ?, ?, ?, ?);'''),
    ],
    params_py='''$stmt->execute([
    $data['name'],
    $data['icon']        ?? '📁',
    $data['color']       ?? '#6366f1',
    $data['description'] ?? '',
    $data['order']       ?? 0
]);''',
    errors=[['400','name məcburi'],['409','name duplicate']]
)

endpoint_block(
    '2.4.3',
    'PUT',
    'backend.php?action=catalog_layers&id={id}',
    'Layer-i yeniləyir.',
    'Session (catalog.edit)',
    'POST ilə eyni body',
    [
        ('UPDATE',
'''UPDATE catalog_layers SET
  name = ?, icon = ?, color = ?, description = ?, `order` = ?
WHERE id = ?;'''),
    ],
    params_py='''$stmt->execute([
    $data['name'],
    $data['icon']        ?? '📁',
    $data['color']       ?? '#6366f1',
    $data['description'] ?? '',
    $data['order']       ?? 0,
    $id
]);'''
)

endpoint_block(
    '2.4.4',
    'DELETE',
    'backend.php?action=catalog_layers&id={id}',
    'Layer-i və bütün tabe cədvəl/sütunları kaskadla silir.',
    'Session (catalog.delete)',
    'URL param: id',
    [
        ('Addım 1 — Layer-də hansı cədvəllər var',
'''SELECT id FROM catalog_tables WHERE layer_id = ?;
-- param: $id'''),
        ('Addım 2 — O cədvəllərin sütunlarını sil',
'''DELETE FROM catalog_columns
WHERE table_id IN (?, ?, ?, ...);
-- params: tableIds array (dinamik placeholder)'''),
        ('Addım 3 — Cədvəlləri sil',
'''DELETE FROM catalog_tables WHERE layer_id = ?;
-- param: $id'''),
        ('Addım 4 — Layer-i özünü sil',
'''DELETE FROM catalog_layers WHERE id = ?;
-- param: $id'''),
    ],
    params_py='''$tablesStmt = db()->prepare("SELECT id FROM catalog_tables WHERE layer_id = ?");
$tablesStmt->execute([$id]);
$tableIds = $tablesStmt->fetchAll(PDO::FETCH_COLUMN);

if (!empty($tableIds)) {
    $ph = implode(',', array_fill(0, count($tableIds), '?'));
    db()->prepare("DELETE FROM catalog_columns WHERE table_id IN ($ph)")->execute($tableIds);
    db()->prepare("DELETE FROM catalog_tables WHERE layer_id = ?")->execute([$id]);
}
db()->prepare("DELETE FROM catalog_layers WHERE id = ?")->execute([$id]);''',
    errors=[['500','Exception orta addımda — inconsistent data (transaction yoxdur)']]
)

# ----- Tables -----
endpoint_block(
    '2.4.5',
    'GET',
    'backend.php?action=catalog_tables',
    'Cədvəllərin siyahısı (opsional layer filteri ilə), hər cədvələ columns array əlavə olunur.',
    'Session (catalog.view)',
    'Query: layer (string, opsional — layer adı)',
    [
        ('Layer filter varsa',
'''SELECT t.*, l.name AS layer_name
FROM catalog_tables t
JOIN catalog_layers l ON t.layer_id = l.id
WHERE l.name = ?
ORDER BY t.name;
-- param: $layer'''),
        ('Filter yoxdursa',
'''SELECT t.*, l.name AS layer_name
FROM catalog_tables t
JOIN catalog_layers l ON t.layer_id = l.id
ORDER BY l.`order`, t.name;'''),
        ('Hər cədvəl üçün sütunlar (loop)',
'''SELECT * FROM catalog_columns
WHERE table_id = ?
ORDER BY `order`;
-- param: $table['id']'''),
    ],
    params_py='''if ($layer) {
    $stmt = db()->prepare("SELECT t.*, l.name as layer_name FROM catalog_tables t
        JOIN catalog_layers l ON t.layer_id = l.id
        WHERE l.name = ? ORDER BY t.name");
    $stmt->execute([$layer]);
} else {
    $stmt = db()->query("SELECT t.*, l.name as layer_name FROM catalog_tables t
        JOIN catalog_layers l ON t.layer_id = l.id ORDER BY l.`order`, t.name");
}
$tables = $stmt->fetchAll();

foreach ($tables as &$table) {
    $colStmt = db()->prepare("SELECT * FROM catalog_columns WHERE table_id = ? ORDER BY `order`");
    $colStmt->execute([$table['id']]);
    $table['columns'] = $colStmt->fetchAll();
}''',
    curl='curl "http://localhost/datarover/backend.php?action=catalog_tables&layer=Mart" -H "Authorization: Bearer $TOKEN"'
)

endpoint_block(
    '2.4.6',
    'POST',
    'backend.php?action=catalog_tables',
    'Yeni cədvəl yaradır.',
    'Session (catalog.create)',
    '''JSON body:
{
  "layer_id": 3,
  "name": "fact_sales",
  "description": "Daily sales fact table",
  "row_count": 1500000,
  "owner": "dwh_team"
}''',
    [
        ('INSERT',
'''INSERT INTO catalog_tables
  (layer_id, name, description, row_count, owner, created_at)
VALUES (?, ?, ?, ?, ?, NOW());'''),
    ],
    params_py='''$stmt->execute([
    $data['layer_id'],
    $data['name'],
    $data['description'] ?? '',
    $data['row_count']   ?? 0,
    $data['owner']       ?? ''
]);''',
    errors=[['409','(layer_id, name) duplicate']]
)

endpoint_block(
    '2.4.7',
    'PUT',
    'backend.php?action=catalog_tables&id={id}',
    'Cədvəl metadatasını yeniləyir.',
    'Session (catalog.edit)',
    'Body: name, description, row_count, owner',
    [
        ('UPDATE',
'''UPDATE catalog_tables SET
  name = ?, description = ?, row_count = ?, owner = ?
WHERE id = ?;'''),
    ],
    params_py='''$stmt->execute([
    $data['name'],
    $data['description'] ?? '',
    $data['row_count']   ?? 0,
    $data['owner']       ?? '',
    $id
]);'''
)

endpoint_block(
    '2.4.8',
    'DELETE',
    'backend.php?action=catalog_tables&id={id}',
    'Cədvəli və onun sütunlarını silir.',
    'Session (catalog.delete)',
    'URL param: id',
    [
        ('Əvvəl sütunları sil',
'''DELETE FROM catalog_columns WHERE table_id = ?;
-- param: $id'''),
        ('Sonra cədvəli sil',
'''DELETE FROM catalog_tables WHERE id = ?;
-- param: $id'''),
    ],
    params_py='''db()->prepare("DELETE FROM catalog_columns WHERE table_id = ?")->execute([$id]);
db()->prepare("DELETE FROM catalog_tables WHERE id = ?")->execute([$id]);'''
)

# ----- Columns -----
endpoint_block(
    '2.4.9',
    'GET',
    'backend.php?action=catalog_columns',
    'Sütunları sadalayır (opsional table_id filter).',
    'Session (catalog.view)',
    'Query: table_id (int, opsional)',
    [
        ('Table filter varsa',
'''SELECT * FROM catalog_columns
WHERE table_id = ?
ORDER BY `order`;
-- param: $tableId'''),
        ('Filter yoxdursa — join ilə',
'''SELECT c.*, t.name AS table_name, l.name AS layer_name
FROM catalog_columns c
JOIN catalog_tables t ON c.table_id = t.id
JOIN catalog_layers l ON t.layer_id = l.id
ORDER BY l.`order`, t.name, c.`order`;'''),
    ]
)

endpoint_block(
    '2.4.10',
    'POST',
    'backend.php?action=catalog_columns',
    'Sütun yaradır.',
    'Session (catalog.create)',
    '''JSON body:
{
  "table_id": 45,
  "name": "customer_id",
  "data_type": "INT",
  "description": "FK → dim_customer.id",
  "is_pk": false,
  "is_fk": true,
  "is_nullable": false,
  "icon": "🔑",
  "order": 3
}''',
    [
        ('INSERT',
'''INSERT INTO catalog_columns
  (table_id, name, data_type, description, is_pk, is_fk, is_nullable, icon, `order`)
VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?);'''),
    ],
    params_py='''$stmt->execute([
    $data['table_id'],
    $data['name'],
    $data['data_type']   ?? 'VARCHAR',
    $data['description'] ?? '',
    $data['is_pk']       ?? 0,
    $data['is_fk']       ?? 0,
    $data['is_nullable'] ?? 1,
    $data['icon']        ?? '📊',
    $data['order']       ?? 0
]);'''
)

endpoint_block(
    '2.4.11',
    'DELETE',
    'backend.php?action=catalog_columns&id={id}',
    'Sütunu silir.',
    'Session (catalog.delete)',
    'URL param: id',
    [
        ('DELETE',
'''DELETE FROM catalog_columns WHERE id = ?;'''),
    ],
    params_py='db()->prepare("DELETE FROM catalog_columns WHERE id = ?")->execute([$id]);'
)

# ----- Profile -----
endpoint_block(
    '2.4.12',
    'GET',
    'backend.php?action=profile_table',
    'Xarici mənbədən live profiling — real row count, sütun statistikası.',
    'Session (catalog.profile)',
    'Query: table (string, adı), layer (string, layer adı)',
    [
        ('Addım 1 — Cədvəl məlumatı',
'''SELECT t.*, l.name AS layer_name
FROM catalog_tables t
JOIN catalog_layers l ON t.layer_id = l.id
WHERE t.name = ? AND l.name = ?;
-- params: $tableName, $layer'''),
        ('Addım 2 — Sütunlar',
'''SELECT * FROM catalog_columns
WHERE table_id = ?
ORDER BY `order`;
-- param: $table['id']'''),
        ('Addım 3 — Xarici mənbə credentials',
'''SELECT * FROM external_sources WHERE id = ?;
-- param: $table['source_id']'''),
        ('Addım 4 — Xarici DB-də row count (dinamik, identifier quoted)',
'''SELECT COUNT(*) AS cnt FROM `schema`.`table`;
-- xarici PDO connection üzərindən'''),
        ('Addım 5 — Yenilənmiş row_count-u catalog-a yaz',
'''UPDATE catalog_tables SET row_count = ? WHERE id = ?;
-- params: $realRowCount, $tableId'''),
    ],
    params_py='''// connectToExternalSource($source) → external PDO
$externalPdo = connectToExternalSource($sourceConnection);
$countStmt = $externalPdo->query("SELECT COUNT(*) as cnt FROM $queryTableName");
$realRowCount = $countStmt->fetch()['cnt'];

// hər sütun üçün: getRealColumnStatsFromExternal() funksiyası
// (backend.php sətir 105–214) — null_count, distinct, min, max qaytarır

$updateStmt = db()->prepare("UPDATE catalog_tables SET row_count = ? WHERE id = ?");
$updateStmt->execute([$realRowCount, $tableId]);''',
    response='''{
  "success": true,
  "data": {
    "table": {"name":"fact_sales","row_count":1500342},
    "columns": [
      {"name":"customer_id","data_type":"INT","null_count":0,"distinct_count":9855}
    ]
  }
}''',
    errors=[['404','Cədvəl catalog-da yoxdur'],['500','Xarici DB-yə bağlantı uğursuz']],
    curl='''curl "http://localhost/datarover/backend.php?action=profile_table&table=fact_sales&layer=Mart" \\
  -H "Authorization: Bearer $TOKEN"'''
)

# ----- External import -----
endpoint_block(
    '2.4.13',
    'POST',
    'backend.php?action=external_sources_import',
    'Xarici mənbədən seçilmiş cədvəlləri və sütunları catalog-a idxal edir (upsert məntiqı).',
    'Session (catalog.import)',
    '''JSON body:
{
  "source_id": 3,
  "tables": [
    {
      "schema_name": "public",
      "table_name": "orders",
      "object_type": "TABLE",
      "row_count": 150000,
      "comment": "İmport edildi",
      "columns": [
        {"column_name":"id","data_type":"INT","full_type":"INT","is_primary_key":true,"comment":"PK"},
        {"column_name":"amount","data_type":"DECIMAL","full_type":"DECIMAL(10,2)","is_primary_key":false,"comment":null}
      ]
    }
  ]
}''',
    [
        ('Addım 1 — Cədvəl mövcudluğunu yoxla',
'''SELECT id FROM catalog_tables WHERE name = ?;
-- param: "$schema.$table" (fullTableName)'''),
        ('Addım 2A — Varsa row_count yenilə',
'''UPDATE catalog_tables SET row_count = ? WHERE id = ?;
-- params: $row_count, $existingId'''),
        ('Addım 2B — Yoxdursa: əvvəlcə layer tap',
'''SELECT id FROM catalog_layers WHERE name = ?;
-- param: $schemaName'''),
        ('Addım 2C — Layer yoxdursa yarat',
'''INSERT INTO catalog_layers (name, description)
VALUES (?, ?);
-- params: $schemaName, 'Imported from external source' '''),
        ('Addım 2D — Yeni cədvəl yarat',
'''INSERT INTO catalog_tables
  (layer_id, name, description, row_count, source_id, object_type, created_at)
VALUES (?, ?, ?, ?, ?, ?, NOW());'''),
        ('Addım 3 — Hər sütun üçün: mövcudluğu yoxla',
'''SELECT id FROM catalog_columns
WHERE table_id = ? AND name = ?;'''),
        ('Addım 3A — Varsa UPDATE (comment qorunur)',
'''UPDATE catalog_columns SET
  data_type = ?,
  is_pk = ?,
  description = COALESCE(?, description)
WHERE id = ?;'''),
        ('Addım 3B — Yoxdursa INSERT',
'''INSERT INTO catalog_columns
  (table_id, name, data_type, is_pk, description)
VALUES (?, ?, ?, ?, ?);'''),
    ],
    params_py='''foreach ($tables as $table) {
  $fullTableName = $table['schema_name'] . '.' . $table['table_name'];
  // ... (yuxarıdakı addımlara uyğun)
  foreach ($table['columns'] as $column) {
      // upsert
  }
}''',
    errors=[['400','source_id və ya tables boş'],['500','Inconsistent state — transaction yox']]
)

doc.add_heading('2.5 Permission Matrix', level=2)
tbl(['Action','SUPER_ADMIN','ENGINEER','STEWARD','ANALYST','VIEWER'],[
    ['catalog.view','✓','✓','✓','✓','✓'],
    ['catalog.create','✓','✓','—','—','—'],
    ['catalog.edit','✓','✓','✓','—','—'],
    ['catalog.delete','✓','✓','—','—','—'],
    ['catalog.profile','✓','✓','—','—','—'],
    ['catalog.import','✓','✓','—','—','—'],
])

doc.add_heading('2.6 Fayl Strukturu', level=2)
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','3913–3970','catalog_layers CRUD'],
    ['backend.php','3976–4037','catalog_tables CRUD'],
    ['backend.php','4043–4084','catalog_columns CRUD'],
    ['backend.php','4090–4236','profile_table'],
    ['backend.php','5219–5321','external_sources_import'],
    ['backend.php','105–214','getRealColumnStatsFromExternal() helper'],
    ['js/app.js','~4000–4250','Catalog UI'],
])

doc.add_heading('2.7 İnteqrasiya', level=2)
tbl(['Modul','Əlaqə'],[
    ['External Sources','catalog_tables.source_id FK → external_sources.id'],
    ['Glossary','physical_attributes JSON → catalog_columns ref'],
    ['Mapping','column_mappings source/target → catalog refs'],
    ['Reports','report_data_sources.table_id FK → catalog_tables.id'],
    ['Data Quality','soda_checks.table_name + column_name → catalog refs'],
])

pb()

# ============================================================
# MODULE 3: REPORTING CATALOG
# ============================================================
doc.add_heading('3. Modul: Reporting Catalog', level=1)

doc.add_heading('3.1 Overview', level=2)
p('Reporting Catalog BI alətlərindəki (Power BI, Tableau, Looker) hesabatları və onların tam metadata-sını saxlayır. Hər hesabat üçün 4 əlaqəli resource mövcuddur: data sources (cədvəl-əsaslı mənbələr), KPIs (metrik tərifləri), users (giriş hüquqları), terms (glossary bağlantıları). Certification workflow: draft → under_review → certified → deprecated.')

doc.add_heading('3.2 Cədvəllər (Data Model)', level=2)

doc.add_heading('3.2.1 reports', level=3)
code('''CREATE TABLE reports (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id VARCHAR(50) UNIQUE COMMENT 'RPT-0001 formatında kod',
  report_name VARCHAR(255) NOT NULL,
  short_description VARCHAR(500),
  long_description TEXT,
  business_purpose TEXT,
  domain_id INT,
  sub_domain VARCHAR(255),
  report_type VARCHAR(100) DEFAULT 'analytical',
  business_owner VARCHAR(255),
  data_owner VARCHAR(255),
  technical_owner VARCHAR(255),
  data_steward VARCHAR(255),
  report_maintainer VARCHAR(255),
  update_frequency VARCHAR(100) DEFAULT 'daily',
  certification_status VARCHAR(100) DEFAULT 'draft',
  bi_tool VARCHAR(100),
  workspace_location VARCHAR(500),
  dashboard_url VARCHAR(500),
  report_version VARCHAR(20) DEFAULT '1.0',
  pages_count INT DEFAULT 1,
  visuals_count INT,
  target_audience VARCHAR(500),
  access_level VARCHAR(50) DEFAULT 'restricted',
  has_pii BOOLEAN DEFAULT FALSE,
  has_financial_data BOOLEAN DEFAULT FALSE,
  regulatory_flags JSON,
  retention_policy VARCHAR(255),
  overall_quality_score DECIMAL(5,2),
  known_issues TEXT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  FOREIGN KEY (domain_id) REFERENCES domains(id) ON DELETE SET NULL,
  INDEX idx_report_id (report_id),
  INDEX idx_status (certification_status),
  INDEX idx_bi_tool (bi_tool)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('3.2.2 report_data_sources', level=3)
code('''CREATE TABLE report_data_sources (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  source_type VARCHAR(50) DEFAULT 'table',
  layer_name VARCHAR(255),
  table_id INT NULL,
  table_name VARCHAR(255),
  schema_name VARCHAR(255),
  database_name VARCHAR(255),
  used_columns JSON,
  etl_process_name VARCHAR(255),
  transformation_logic TEXT,
  refresh_frequency VARCHAR(100),
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
  FOREIGN KEY (table_id) REFERENCES catalog_tables(id) ON DELETE SET NULL,
  INDEX idx_report (report_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('3.2.3 report_kpis', level=3)
code('''CREATE TABLE report_kpis (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  kpi_name VARCHAR(255),
  kpi_code VARCHAR(100),
  business_definition TEXT,
  technical_formula TEXT,
  aggregation_rule VARCHAR(255),
  unit_of_measure VARCHAR(100),
  target_value DECIMAL(15,2),
  threshold_warning DECIMAL(15,2),
  threshold_critical DECIMAL(15,2),
  current_value DECIMAL(15,2),
  trend VARCHAR(50) DEFAULT 'stable',
  filters_applied JSON,
  dependencies JSON,
  term_id INT NULL,
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE SET NULL
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('3.2.4 report_users', level=3)
code('''CREATE TABLE report_users (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  user_name VARCHAR(255),
  user_email VARCHAR(255),
  user_role VARCHAR(100) DEFAULT 'viewer',
  department VARCHAR(255),
  access_granted_by VARCHAR(255),
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('3.2.5 report_terms', level=3)
code('''CREATE TABLE report_terms (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  term_id INT NOT NULL,
  relationship_type VARCHAR(100) DEFAULT 'uses',
  notes TEXT,
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE CASCADE,
  UNIQUE KEY unique_report_term (report_id, term_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('3.3 ER Diagram', level=2)
code('''erDiagram
    domains ||--o{ reports : categorizes
    reports ||--o{ report_data_sources : uses
    reports ||--o{ report_kpis : contains
    reports ||--o{ report_users : shared_with
    reports ||--o{ report_terms : references
    catalog_tables ||--o{ report_data_sources : supplies
    glossary_terms ||--o{ report_terms : linked
    glossary_terms ||--o{ report_kpis : defines''')

doc.add_heading('3.4 API Endpoint-ləri', level=2)

# ----- reports list -----
endpoint_block(
    '3.4.1',
    'GET',
    'backend.php?action=reports',
    'Hesabat siyahısı — dinamik filter (search, domain, bi_tool, status).',
    'Session (reports.view)',
    '''Query parametrlər:
  search    — report_name / short_description / report_id axtarışı
  domain    — domain_id filteri
  bi_tool   — power_bi | tableau | looker | qlik | metabase
  status    — certification_status filteri''',
    [
        ('Əsas SELECT (domain JOIN)',
'''SELECT r.*, d.name AS domain_name, d.icon AS domain_icon
FROM reports r
LEFT JOIN domains d ON r.domain_id = d.id
WHERE 1=1'''),
        ('+ search',
'''AND (r.report_name LIKE ? OR r.short_description LIKE ? OR r.report_id LIKE ?)
-- params: "%$search%" × 3'''),
        ('+ domain',
'AND r.domain_id = ?'),
        ('+ bi_tool',
'AND r.bi_tool = ?'),
        ('+ status',
'AND r.certification_status = ?'),
        ('+ sıralama',
'ORDER BY r.report_name;'),
    ],
    params_py='''$sql = "SELECT r.*, d.name as domain_name, d.icon as domain_icon
        FROM reports r LEFT JOIN domains d ON r.domain_id = d.id WHERE 1=1";
$params = [];

if ($search) {
    $sql .= " AND (r.report_name LIKE ? OR r.short_description LIKE ? OR r.report_id LIKE ?)";
    $t = "%$search%";
    $params = array_merge($params, [$t, $t, $t]);
}
if ($domain)  { $sql .= " AND r.domain_id = ?";            $params[] = $domain; }
if ($biTool)  { $sql .= " AND r.bi_tool = ?";              $params[] = $biTool; }
if ($status)  { $sql .= " AND r.certification_status = ?"; $params[] = $status; }

$sql .= " ORDER BY r.report_name";
$stmt = db()->prepare($sql); $stmt->execute($params);''',
    response='''{
  "success": true,
  "data": [
    {
      "id": 1,
      "report_id": "RPT-0001",
      "report_name": "Sales Executive Dashboard",
      "bi_tool": "power_bi",
      "certification_status": "certified",
      "domain_name": "Sales",
      "overall_quality_score": 96.5
    }
  ]
}''',
    curl='''curl "http://localhost/datarover/backend.php?action=reports&bi_tool=power_bi&status=certified" \\
  -H "Authorization: Bearer $TOKEN"'''
)

# ----- reports create -----
endpoint_block(
    '3.4.2',
    'POST',
    'backend.php?action=reports',
    'Yeni hesabat qeydiyyatı. report_id verilməzsə avtomatik "RPT-NNNN" şəklində generasiya olunur.',
    'Session (reports.create)',
    '''JSON body (27 sahə):
{
  "report_id": "RPT-0042",
  "report_name": "Customer 360",
  "short_description": "...",
  "long_description": "...",
  "business_purpose": "...",
  "domain_id": 3,
  "sub_domain": "B2C",
  "report_type": "analytical",
  "business_owner": "cso@x.com",
  "data_owner": "...",
  "technical_owner": "...",
  "data_steward": "...",
  "report_maintainer": "...",
  "update_frequency": "daily",
  "certification_status": "draft",
  "bi_tool": "power_bi",
  "workspace_location": "...",
  "dashboard_url": "https://app.powerbi.com/...",
  "report_version": "1.0",
  "pages_count": 5,
  "visuals_count": 28,
  "target_audience": "...",
  "access_level": "restricted",
  "has_pii": false,
  "has_financial_data": true,
  "regulatory_flags": ["SOX"],
  "retention_policy": "7 years"
}''',
    [
        ('INSERT — 27 sahə',
'''INSERT INTO reports (
  report_id, report_name, short_description, long_description, business_purpose,
  domain_id, sub_domain, report_type, business_owner, data_owner, technical_owner,
  data_steward, report_maintainer, update_frequency, certification_status,
  bi_tool, workspace_location, dashboard_url, report_version, pages_count, visuals_count,
  target_audience, access_level, has_pii, has_financial_data, regulatory_flags, retention_policy
) VALUES (?, ?, ?, ?, ?,
          ?, ?, ?, ?, ?, ?,
          ?, ?, ?, ?,
          ?, ?, ?, ?, ?, ?,
          ?, ?, ?, ?, ?, ?);'''),
    ],
    params_py='''// report_id generasiyası
$reportId = $data['report_id'] ?? 'RPT-' . str_pad(rand(1, 9999), 4, '0', STR_PAD_LEFT);

$stmt->execute([
    $reportId,
    $data['report_name'],
    $data['short_description']   ?? '',
    $data['long_description']    ?? '',
    $data['business_purpose']    ?? '',
    $data['domain_id']           ?: null,
    $data['sub_domain']          ?? '',
    $data['report_type']         ?? 'analytical',
    $data['business_owner']      ?? '',
    $data['data_owner']          ?? '',
    $data['technical_owner']     ?? '',
    $data['data_steward']        ?? '',
    $data['report_maintainer']   ?? '',
    $data['update_frequency']    ?? 'daily',
    $data['certification_status']?? 'draft',
    $data['bi_tool']             ?? 'power_bi',
    $data['workspace_location']  ?? '',
    $data['dashboard_url']       ?? '',
    $data['report_version']      ?? '1.0',
    $data['pages_count']         ?? 1,
    $data['visuals_count']       ?? 0,
    $data['target_audience']     ?? '',
    $data['access_level']        ?? 'restricted',
    $data['has_pii']             ?? 0,
    $data['has_financial_data']  ?? 0,
    isset($data['regulatory_flags']) ? json_encode($data['regulatory_flags']) : null,
    $data['retention_policy']    ?? ''
]);''',
    response='{"success":true,"data":{"id":42,"report_id":"RPT-0042"}}',
    errors=[['409','report_id duplicate']]
)

# ----- reports update (dynamic) -----
endpoint_block(
    '3.4.3',
    'PUT',
    'backend.php?action=reports&id={id}',
    'Yalnız body-də verilən sahələri yeniləyir (partial update). UPDATE clause-u dinamik tikilir.',
    'Session (reports.edit)',
    '''JSON body (yalnız yeniləmək istədiyin sahələr):
{
  "certification_status": "certified",
  "overall_quality_score": 96.5,
  "known_issues": "..."
}''',
    [
        ('Dinamik UPDATE şablonu',
'''UPDATE reports SET <field1> = ?, <field2> = ?, ... WHERE id = ?;

-- Icazə verilən sahələr:
-- report_name, short_description, long_description, business_purpose, domain_id,
-- sub_domain, report_type, business_owner, data_owner, technical_owner,
-- data_steward, report_maintainer, update_frequency, certification_status,
-- bi_tool, workspace_location, dashboard_url, report_version, pages_count,
-- visuals_count, target_audience, access_level, has_pii, has_financial_data,
-- retention_policy, overall_quality_score, known_issues, regulatory_flags'''),
    ],
    params_py='''$fields = ['report_name','short_description','long_description','business_purpose',
    'domain_id','sub_domain','report_type','business_owner','data_owner',
    'technical_owner','data_steward','report_maintainer','update_frequency',
    'certification_status','bi_tool','workspace_location','dashboard_url',
    'report_version','pages_count','visuals_count','target_audience','access_level',
    'has_pii','has_financial_data','retention_policy','overall_quality_score','known_issues'];

$updates = []; $params = [];
foreach ($fields as $f) {
    if (isset($data[$f])) { $updates[] = "$f = ?"; $params[] = $data[$f]; }
}

// JSON sahə xüsusi
if (isset($data['regulatory_flags'])) {
    $updates[] = 'regulatory_flags = ?';
    $params[] = json_encode($data['regulatory_flags']);
}

if (empty($updates)) {
    respond(['message' => 'Nothing to update']);
} else {
    $params[] = $id;
    $stmt = db()->prepare("UPDATE reports SET " . implode(', ', $updates) . " WHERE id = ?");
    $stmt->execute($params);
}''',
    response='{"success":true,"data":{"message":"Report updated"}}'
)

# ----- reports delete -----
endpoint_block(
    '3.4.4',
    'DELETE',
    'backend.php?action=reports&id={id}',
    'Hesabatı silir. CASCADE FK sayəsində bütün data_sources, kpis, users, terms də silinir.',
    'Session (reports.delete)',
    'URL param: id',
    [
        ('DELETE',
'''DELETE FROM reports WHERE id = ?;
-- param: $id
-- ON DELETE CASCADE: report_data_sources, report_kpis, report_users, report_terms də silinir'''),
    ]
)

# ----- report_data_sources -----
doc.add_heading('3.4.5  report_data_sources (sub-resource)', level=3)

endpoint_block(
    '3.4.5.1',
    'GET',
    'backend.php?action=report_data_sources&report_id={id}',
    'Hesabatın data mənbələrini catalog məlumatları ilə birgə qaytarır.',
    'Session',
    'Query: report_id (int, məcburi)',
    [
        ('JOIN ilə catalog məlumatı',
'''SELECT rds.*,
       ct.name AS catalog_table_name,
       cl.name AS catalog_layer_name
FROM report_data_sources rds
LEFT JOIN catalog_tables ct ON rds.table_id = ct.id
LEFT JOIN catalog_layers cl ON ct.layer_id = cl.id
WHERE rds.report_id = ?;
-- param: $reportId'''),
    ],
    errors=[['400','report_id tələb olunur']]
)

endpoint_block(
    '3.4.5.2',
    'POST',
    'backend.php?action=report_data_sources',
    'Hesabata data mənbəyi əlavə edir.',
    'Session',
    '''JSON body:
{
  "report_id": 42,
  "source_type": "table",
  "layer_name": "Mart",
  "table_id": 145,
  "table_name": "fact_sales",
  "schema_name": "mart",
  "database_name": "dwh",
  "used_columns": ["customer_id","amount","order_date"],
  "etl_process_name": "etl_sales_daily",
  "transformation_logic": "SUM(amount) GROUP BY customer",
  "refresh_frequency": "daily"
}''',
    [
        ('INSERT (11 sahə)',
'''INSERT INTO report_data_sources (
  report_id, source_type, layer_name, table_id, table_name,
  schema_name, database_name, used_columns,
  etl_process_name, transformation_logic, refresh_frequency
) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);'''),
    ],
    params_py='''$stmt->execute([
    $data['report_id'],
    $data['source_type']    ?? 'table',
    $data['layer_name']     ?? '',
    $data['table_id']       ?: null,
    $data['table_name']     ?? '',
    $data['schema_name']    ?? '',
    $data['database_name']  ?? '',
    isset($data['used_columns']) ? json_encode($data['used_columns']) : null,
    $data['etl_process_name']    ?? '',
    $data['transformation_logic']?? '',
    $data['refresh_frequency']   ?? ''
]);'''
)

endpoint_block(
    '3.4.5.3',
    'DELETE',
    'backend.php?action=report_data_sources&id={id}',
    'Data mənbəyini silir.',
    'Session',
    'URL param: id',
    [
        ('DELETE',
'DELETE FROM report_data_sources WHERE id = ?;'),
    ]
)

# ----- report_kpis -----
doc.add_heading('3.4.6  report_kpis (sub-resource)', level=3)

endpoint_block(
    '3.4.6.1',
    'GET',
    'backend.php?action=report_kpis&report_id={id}',
    'Hesabatın KPI-larını glossary term adı ilə birgə qaytarır.',
    'Session',
    'Query: report_id (məcburi)',
    [
        ('JOIN glossary',
'''SELECT rk.*, gt.name AS term_name
FROM report_kpis rk
LEFT JOIN glossary_terms gt ON rk.term_id = gt.id
WHERE rk.report_id = ?;
-- param: $reportId'''),
    ]
)

endpoint_block(
    '3.4.6.2',
    'POST',
    'backend.php?action=report_kpis',
    'Hesabata yeni KPI əlavə edir.',
    'Session',
    '''JSON body:
{
  "report_id": 42,
  "kpi_name": "Monthly Active Users",
  "kpi_code": "MAU",
  "business_definition": "...",
  "technical_formula": "COUNT(DISTINCT user_id) ...",
  "aggregation_rule": "sum",
  "unit_of_measure": "users",
  "target_value": 50000,
  "threshold_warning": 45000,
  "threshold_critical": 40000,
  "current_value": 48250,
  "trend": "up",
  "filters_applied": {"date":"last_30_days"},
  "dependencies": ["user_events","sessions"],
  "term_id": 27
}''',
    [
        ('INSERT (15 sahə)',
'''INSERT INTO report_kpis (
  report_id, kpi_name, kpi_code, business_definition, technical_formula,
  aggregation_rule, unit_of_measure, target_value, threshold_warning, threshold_critical,
  current_value, trend, filters_applied, dependencies, term_id
) VALUES (?, ?, ?, ?, ?,
          ?, ?, ?, ?, ?,
          ?, ?, ?, ?, ?);'''),
    ],
    params_py='''$stmt->execute([
    $data['report_id'],
    $data['kpi_name'],
    $data['kpi_code']            ?? '',
    $data['business_definition'] ?? '',
    $data['technical_formula']   ?? '',
    $data['aggregation_rule']    ?? '',
    $data['unit_of_measure']     ?? '',
    $data['target_value']        ?: null,
    $data['threshold_warning']   ?: null,
    $data['threshold_critical']  ?: null,
    $data['current_value']       ?: null,
    $data['trend']               ?? 'stable',
    isset($data['filters_applied']) ? json_encode($data['filters_applied']) : null,
    isset($data['dependencies'])    ? json_encode($data['dependencies'])    : null,
    $data['term_id']             ?: null
]);'''
)

endpoint_block(
    '3.4.6.3',
    'PUT',
    'backend.php?action=report_kpis&id={id}',
    'KPI yeniləyir (trend, current_value monitoring üçün).',
    'Session',
    'POST ilə eyni body',
    [
        ('UPDATE (13 sahə — report_id, kpi_code yenilənmir)',
'''UPDATE report_kpis SET
  kpi_name = ?, kpi_code = ?, business_definition = ?, technical_formula = ?,
  aggregation_rule = ?, unit_of_measure = ?, target_value = ?,
  threshold_warning = ?, threshold_critical = ?, current_value = ?,
  trend = ?, term_id = ?
WHERE id = ?;'''),
    ],
    params_py='''$stmt->execute([
    $data['kpi_name'],
    $data['kpi_code']            ?? '',
    $data['business_definition'] ?? '',
    $data['technical_formula']   ?? '',
    $data['aggregation_rule']    ?? '',
    $data['unit_of_measure']     ?? '',
    $data['target_value']        ?: null,
    $data['threshold_warning']   ?: null,
    $data['threshold_critical']  ?: null,
    $data['current_value']       ?: null,
    $data['trend']               ?? 'stable',
    $data['term_id']             ?: null,
    $id
]);'''
)

endpoint_block(
    '3.4.6.4',
    'DELETE',
    'backend.php?action=report_kpis&id={id}',
    'KPI-nı silir.',
    'Session',
    'URL param: id',
    [('DELETE','DELETE FROM report_kpis WHERE id = ?;')]
)

# ----- report_users -----
doc.add_heading('3.4.7  report_users (sub-resource)', level=3)

endpoint_block(
    '3.4.7.1',
    'GET',
    'backend.php?action=report_users&report_id={id}',
    'Hesabata girişi olan istifadəçilər.',
    'Session',
    'Query: report_id',
    [
        ('SELECT',
'''SELECT * FROM report_users
WHERE report_id = ?
ORDER BY user_name;'''),
    ]
)

endpoint_block(
    '3.4.7.2',
    'POST',
    'backend.php?action=report_users',
    'Hesabata istifadəçi giriş hüququ verir.',
    'Session',
    '''JSON body:
{
  "report_id": 42,
  "user_name": "John Doe",
  "user_email": "john@corp.com",
  "user_role": "viewer",
  "department": "Finance",
  "access_granted_by": "admin@corp.com"
}''',
    [
        ('INSERT',
'''INSERT INTO report_users
  (report_id, user_name, user_email, user_role, department, access_granted_by)
VALUES (?, ?, ?, ?, ?, ?);'''),
    ],
    params_py='''$stmt->execute([
    $data['report_id'],
    $data['user_name'],
    $data['user_email']        ?? '',
    $data['user_role']         ?? 'viewer',
    $data['department']        ?? '',
    $data['access_granted_by'] ?? ''
]);'''
)
p('Qeyd: report_users üçün PUT endpoint-i yoxdur — yeniləmək üçün DELETE + POST edilməlidir.')

endpoint_block(
    '3.4.7.3',
    'DELETE',
    'backend.php?action=report_users&id={id}',
    'İstifadəçi girişini ləğv edir.',
    'Session',
    'URL param: id',
    [('DELETE','DELETE FROM report_users WHERE id = ?;')]
)

# ----- report_terms -----
doc.add_heading('3.4.8  report_terms (sub-resource)', level=3)

endpoint_block(
    '3.4.8.1',
    'GET',
    'backend.php?action=report_terms&report_id={id}',
    'Hesabata bağlı glossary termləri (ad, tərif, domain ilə).',
    'Session',
    'Query: report_id',
    [
        ('JOIN glossary',
'''SELECT rt.*,
       gt.name       AS term_name,
       gt.definition AS term_definition,
       gt.domain     AS term_domain
FROM report_terms rt
JOIN glossary_terms gt ON rt.term_id = gt.id
WHERE rt.report_id = ?;'''),
    ]
)

endpoint_block(
    '3.4.8.2',
    'POST',
    'backend.php?action=report_terms',
    'Hesabata term əlavə edir (upsert — (report_id, term_id) unikal key-dir).',
    'Session',
    '''JSON body:
{
  "report_id": 42,
  "term_id": 13,
  "relationship_type": "uses",
  "notes": "CLV metric primary source"
}''',
    [
        ('UPSERT (INSERT ... ON DUPLICATE KEY UPDATE)',
'''INSERT INTO report_terms
  (report_id, term_id, relationship_type, notes)
VALUES (?, ?, ?, ?)
ON DUPLICATE KEY UPDATE
  relationship_type = VALUES(relationship_type),
  notes             = VALUES(notes);'''),
    ],
    params_py='''$stmt->execute([
    $data['report_id'],
    $data['term_id'],
    $data['relationship_type'] ?? 'uses',
    $data['notes']             ?? ''
]);'''
)

endpoint_block(
    '3.4.8.3',
    'DELETE',
    'backend.php?action=report_terms&id={id}',
    'Term bağlantısını silir.',
    'Session',
    'URL param: id',
    [('DELETE','DELETE FROM report_terms WHERE id = ?;')]
)

# ----- search_reports -----
endpoint_block(
    '3.4.9',
    'GET',
    'backend.php?action=search_reports&q={query}',
    'Global hesabat axtarışı (case-insensitive, 4 sahədə).',
    'Session',
    'Query: q (search term, məcburi)',
    [
        ('Case-insensitive axtarış',
'''SELECT id,
       report_name       AS name,
       short_description AS description,
       report_type,
       created_at
FROM reports
WHERE LOWER(report_name)       LIKE ?
   OR LOWER(short_description) LIKE ?
   OR LOWER(long_description)  LIKE ?
   OR LOWER(report_type)       LIKE ?
LIMIT 20;
-- params: "%$q%" × 4'''),
    ],
    params_py='''$searchPattern = '%' . strtolower($query) . '%';
$stmt->execute([$searchPattern, $searchPattern, $searchPattern, $searchPattern]);''',
    response='''{
  "reports": [
    {"id":42,"name":"Customer 360","description":"...","report_type":"analytical"}
  ]
}'''
)

doc.add_heading('3.5 Business Logic', level=2)
tbl(['Qayda','Təsvir'],[
    ['report_id auto-generate','Verilməzsə: "RPT-" + rand(1,9999) 4-rəqəmli'],
    ['certification workflow','draft → under_review → certified → deprecated'],
    ['regulatory_flags','JSON array, məs: ["GDPR","SOX","HIPAA"]'],
    ['overall_quality_score','Data quality moduldan gündəlik refresh oluna bilər'],
    ['CASCADE delete','reports silinsə 4 sub-resource də silinir (FK CASCADE)'],
    ['report_terms unique','(report_id, term_id) unikal key — upsert ilə idempotent INSERT'],
    ['report_users','PUT yoxdur — DELETE + POST patterni'],
])

doc.add_heading('3.6 Permission Matrix', level=2)
tbl(['Action','SUPER_ADMIN','ANALYST','STEWARD','ENGINEER','VIEWER'],[
    ['reports.view','✓','✓','✓','✓','✓'],
    ['reports.create','✓','✓','—','—','—'],
    ['reports.edit','✓','✓','✓','—','—'],
    ['reports.certify','✓','—','✓','—','—'],
    ['reports.delete','✓','—','✓','—','—'],
    ['reports.grant_access','✓','✓','✓','—','—'],
])

doc.add_heading('3.7 Fayl Strukturu', level=2)
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','4369–4492','reports CRUD (dinamik UPDATE)'],
    ['backend.php','4494–4537','report_data_sources'],
    ['backend.php','4539–4614','report_kpis'],
    ['backend.php','4616–4649','report_users'],
    ['backend.php','4651–4688','report_terms'],
    ['backend.php','5642–5671','search_reports'],
    ['js/app.js','Reports tab','loadReports, saveReport, linkKpi'],
])

doc.add_heading('3.8 İnteqrasiya', level=2)
tbl(['Modul','Əlaqə'],[
    ['Domain','reports.domain_id FK → domains.id'],
    ['Catalog','report_data_sources.table_id FK → catalog_tables.id'],
    ['Glossary','report_terms junction + report_kpis.term_id FK'],
    ['Data Quality','overall_quality_score hesablamasında quality_rule_results istifadə'],
    ['Auth','report_users.user_email → users.email ilə cross-reference'],
])

doc.add_heading('3.9 Ümumi SELECT Nümunələri (əlavə — developer reference)', level=2)
p('Hesabat + bütün sub-resource-lər bir sorğuda:')
code('''SELECT r.*,
  d.name AS domain_name,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT(
      'id',id,'layer_name',layer_name,'table_name',table_name,'table_id',table_id,
      'used_columns',used_columns,'refresh_frequency',refresh_frequency))
   FROM report_data_sources WHERE report_id = r.id) AS data_sources,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT(
      'id',id,'kpi_name',kpi_name,'kpi_code',kpi_code,
      'target_value',target_value,'current_value',current_value,'trend',trend))
   FROM report_kpis WHERE report_id = r.id) AS kpis,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT(
      'id',id,'user_name',user_name,'user_email',user_email,'user_role',user_role))
   FROM report_users WHERE report_id = r.id) AS users,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT(
      'term_id',rt.term_id,'term_name',gt.name,'relationship_type',rt.relationship_type))
   FROM report_terms rt JOIN glossary_terms gt ON gt.id = rt.term_id
   WHERE rt.report_id = r.id) AS terms
FROM reports r
LEFT JOIN domains d ON d.id = r.domain_id
WHERE r.id = ?;''')

p('Domain-based report statistikası:')
code('''SELECT d.name AS domain,
       COUNT(r.id) AS total_reports,
       SUM(CASE WHEN r.certification_status='certified' THEN 1 ELSE 0 END) AS certified,
       AVG(r.overall_quality_score) AS avg_quality
FROM domains d
LEFT JOIN reports r ON r.domain_id = d.id
GROUP BY d.id, d.name
ORDER BY total_reports DESC;''')

p('Bir term neçə hesabatda istifadə olunur:')
code('''SELECT gt.id, gt.name,
       COUNT(DISTINCT rt.report_id) AS reports_using
FROM glossary_terms gt
LEFT JOIN report_terms rt ON rt.term_id = gt.id
LEFT JOIN report_kpis rk ON rk.term_id = gt.id
GROUP BY gt.id
ORDER BY reports_using DESC;''')

pb()

# ============================================================
# APPENDIX: Cross-cutting & Setup
# ============================================================
doc.add_heading('Əlavə: Modullararası Əlaqələr', level=1)

p('Üç modulun bir-biri ilə qarşılıqlı təsir diaqramı:')
code('''erDiagram
    catalog_layers ||--o{ catalog_tables : "contains"
    catalog_tables ||--o{ catalog_columns : "has"
    catalog_tables ||--o{ report_data_sources : "supplies"

    glossary_terms ||--o{ glossary_term_history : "changes"
    glossary_terms ||--o{ term_physical_attributes : "maps_to_catalog"
    glossary_terms ||--o{ report_terms : "referenced_by"
    glossary_terms ||--o{ report_kpis : "defines"

    reports ||--o{ report_data_sources : "uses"
    reports ||--o{ report_kpis : "measures"
    reports ||--o{ report_users : "shared_with"
    reports ||--o{ report_terms : "documents"

    domains ||--o{ reports : "categorizes"
    domains }o..o{ glossary_terms : "domain_field_match"''')

doc.add_heading('Setup — Tam Migration Ardıcıllığı', level=2)
code('''-- 1. Base cədvəllər (əvvəlcə)
CREATE TABLE domains ( ... );           -- (dependency)
CREATE TABLE users ( ... );             -- (dependency)
CREATE TABLE external_sources ( ... );  -- (dependency)

-- 2. Glossary (bir-birinə bağlı 3 cədvəl)
CREATE TABLE glossary_terms ( ... );
CREATE TABLE glossary_term_history ( ... );
CREATE TABLE term_physical_attributes ( ... );

-- 3. Catalog (3 layer)
CREATE TABLE catalog_layers ( ... );
CREATE TABLE catalog_tables ( ... );    -- FK → catalog_layers, external_sources
CREATE TABLE catalog_columns ( ... );   -- FK → catalog_tables

-- 4. Reports (1 main + 4 sub-resource)
CREATE TABLE reports ( ... );           -- FK → domains
CREATE TABLE report_data_sources ( ... ); -- FK → reports, catalog_tables
CREATE TABLE report_kpis ( ... );       -- FK → reports, glossary_terms
CREATE TABLE report_users ( ... );      -- FK → reports
CREATE TABLE report_terms ( ... );      -- FK → reports, glossary_terms

-- 5. Seed
INSERT INTO catalog_layers (name, icon, color, `order`) VALUES
  ('Raw',       '📥', '#ef4444', 0),
  ('Staging',   '🔄', '#f59e0b', 1),
  ('Warehouse', '🏛️', '#6366f1', 2),
  ('Mart',      '🎯', '#10b981', 3);

INSERT INTO domains (name, icon, color) VALUES
  ('Sales',     '💰', '#10b981'),
  ('Finance',   '💵', '#3b82f6'),
  ('Operations','⚙️', '#f59e0b');''')

doc.add_heading('Verification Sorğuları', level=2)
code('''-- Cədvəllər yaradılıbmı?
SHOW TABLES LIKE 'glossary_%';
SHOW TABLES LIKE 'catalog_%';
SHOW TABLES LIKE 'report_%';

-- Seed datası
SELECT COUNT(*) FROM catalog_layers;       -- ≥ 4
SELECT COUNT(*) FROM domains;               -- ≥ 3

-- FK konstraintlər aktivdirmi?
SELECT table_name, constraint_name, referenced_table_name
FROM information_schema.referential_constraints
WHERE constraint_schema = 'datarover';

-- Smoke test (həqiqi row-larla)
SELECT
  (SELECT COUNT(*) FROM glossary_terms) AS terms,
  (SELECT COUNT(*) FROM catalog_tables) AS tables,
  (SELECT COUNT(*) FROM reports)        AS reports;''')

doc.add_heading('Təhlükəsizlik Qeydləri', level=2)
tbl(['Risk','Tövsiyə'],[
    ['Transaction yoxdur','Multi-statement operasiyalar (hard delete, external_sources_import) üçün START TRANSACTION əlavə et'],
    ['Dinamik identifier quoting','profile_table-da $queryTableName istifadəçi input-dan gəlirsə backtick escape məcburi'],
    ['Password plain-text','external_sources.password — AES-256-CBC encryption tətbiq et (APP_ENCRYPTION_KEY env)'],
    ['JSON field validation','Frontend-dən gələn stewards/qualityRules/etc. json_encode edilməzdən əvvəl array olduğundan əmin ol'],
    ['SQL injection (LIKE)','search termində % və _ wildcard escape olunmalıdır: addcslashes($q, "%_")'],
    ['CASCADE vs audit','CASCADE DELETE audit_log qeydini əlində saxlamır — silmə əvvəli manual INSERT audit_log'],
])

doc.add_heading('Dəyişiklik Nümunələri (Developer üçün)', level=2)
p('A) Yeni sahə əlavə etmək (məs: reports.favorite_count):')
code('''-- 1. Migration
ALTER TABLE reports ADD COLUMN favorite_count INT DEFAULT 0;

-- 2. backend.php 3.4.2 POST-a parametr əlavə:
// $fields array-ına 'favorite_count' əlavə et (dinamik UPDATE üçün)

-- 3. app.js-də UI binding
// Report form-una yeni input əlavə et''')

p('B) Yeni endpoint əlavə etmək (məs: report_tags):')
code('''-- 1. CREATE TABLE report_tags (
--      id INT PK, report_id INT FK, tag VARCHAR(100), ...)
-- 2. backend.php sonuna əlavə:
if ($action === 'report_tags') {
  if ($method === 'GET')    { /* SELECT */ }
  if ($method === 'POST')   { /* INSERT */ }
  if ($method === 'DELETE') { /* DELETE */ }
}
-- 3. app.js-də fetch(API + '?action=report_tags')''')

# ============================================================
# SAVE
# ============================================================
out = r'c:\xampp\htdocs\datarover\Technical_Spec_Detailed.docx'
doc.save(out)
print(f'Saved: {out}')
