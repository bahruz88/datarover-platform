#!/usr/bin/env python3
"""Generate API Documentation DOCX file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(18)
style_h1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
style_h1.font.bold = True

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(14)
style_h2.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(12)
style_h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('API Documentation')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Data Governance Platform — REST API Reference')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Version 1.0  |  {datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Authentication & Authorization',
    '3. General Information',
    '4. Authentication Endpoints',
    '5. LDAP Integration Endpoints',
    '6. User Management Endpoints',
    '7. Role & Permission Endpoints',
    '8. Session Management Endpoints',
    '9. Password Policy Endpoints',
    '10. Business Glossary Endpoints',
    '11. Data Catalog Endpoints',
    '12. Data Mapping Endpoints',
    '13. Reporting Endpoints',
    '14. Domain Management Endpoints',
    '15. Governance Endpoints',
    '16. External Data Sources Endpoints',
    '17. Platform Statistics',
    '18. Export',
    '19. Error Handling',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This document provides a comprehensive reference for the REST API of the '
    'Data Governance Platform. The platform offers a centralized solution for '
    'managing business glossaries, data catalogs, data quality rules, report '
    'cataloging, governance workflows, and user administration.'
)
doc.add_paragraph(
    'The API follows a query-parameter based routing pattern where endpoints '
    'are accessed through a single entry point with an "action" parameter '
    'specifying the desired operation.'
)

# ============================================================
# 2. AUTHENTICATION & AUTHORIZATION
# ============================================================
doc.add_heading('2. Authentication & Authorization', level=1)

doc.add_heading('2.1 Authentication Methods', level=2)
doc.add_paragraph(
    'The platform supports the following authentication mechanisms:'
)
add_table(doc, ['Method', 'Description', 'Usage'],
[
    ['Session Token', 'Traditional username/password login returns a session token', 'Authorization: Bearer <token>'],
    ['API Key', 'For external system integrations (quality tools, ETL, etc.)', 'X-Api-Key: <key> or Authorization: Bearer <key>'],
    ['SSO / Keycloak', 'Single Sign-On via external identity provider', 'Redirected through SSO flow'],
    ['LDAP', 'Active Directory / LDAP authentication', 'Configured via admin settings'],
])

doc.add_heading('2.2 Role-Based Access Control (RBAC)', level=2)
doc.add_paragraph(
    'The platform implements a fine-grained RBAC system. Users are assigned roles, '
    'and each role contains a set of permissions in the format module.action '
    '(e.g., "catalog.view", "glossary.create"). Super administrators bypass all permission checks.'
)
doc.add_paragraph('Key authorization functions:')
add_table(doc, ['Function', 'Description'],
[
    ['requireAuth()', 'Ensures user is logged in; returns 401 if not'],
    ['requirePermission(module, action)', 'Checks specific permission; returns 403 if denied'],
    ['hasPermission(module, action)', 'Boolean check for permission'],
    ['hasRole(roleCode)', 'Checks if user has a specific role'],
    ['hasAnyPermission(permissions[])', 'True if user has at least one of the listed permissions'],
    ['hasAllPermissions(permissions[])', 'True if user has all listed permissions'],
])

# ============================================================
# 3. GENERAL INFORMATION
# ============================================================
doc.add_heading('3. General Information', level=1)

doc.add_heading('3.1 Base URL', level=2)
code_block(doc, 'https://<your-domain>/backend.php?action=<endpoint>')

doc.add_heading('3.2 Request / Response Format', level=2)
doc.add_paragraph('All requests and responses use JSON format.')
doc.add_paragraph('Success response:')
code_block(doc, '{\n  "success": true,\n  "data": { ... }\n}')
doc.add_paragraph('Error response:')
code_block(doc, '{\n  "success": false,\n  "error": "Error message description"\n}')

doc.add_heading('3.3 HTTP Methods', level=2)
add_table(doc, ['Method', 'Purpose'],
[
    ['GET', 'Retrieve resources'],
    ['POST', 'Create new resources'],
    ['PUT', 'Update existing resources'],
    ['DELETE', 'Remove resources'],
])

doc.add_heading('3.4 Common HTTP Status Codes', level=2)
add_table(doc, ['Code', 'Meaning'],
[
    ['200', 'Success'],
    ['400', 'Bad Request — missing or invalid parameters'],
    ['401', 'Unauthorized — authentication required'],
    ['403', 'Forbidden — insufficient permissions'],
    ['404', 'Not Found'],
    ['405', 'Method Not Allowed'],
    ['500', 'Internal Server Error'],
])

doc.add_heading('3.5 Timezone', level=2)
doc.add_paragraph('All timestamps are in the Asia/Baku timezone (UTC+4). Datetime fields use ISO 8601 format.')

doc.add_page_break()

# ============================================================
# 4. AUTHENTICATION ENDPOINTS
# ============================================================
doc.add_heading('4. Authentication Endpoints', level=1)

doc.add_heading('4.1 Login', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'auth/login'],
    ['Method', 'POST'],
    ['Auth Required', 'No'],
    ['Description', 'Authenticate user with username and password'],
])
doc.add_paragraph('Request Body:')
code_block(doc, '{\n  "username": "string (required)",\n  "password": "string (required)"\n}')
doc.add_paragraph('Response: Token, user information, roles, and permissions.')

doc.add_heading('4.2 Logout', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'auth/logout'],
    ['Method', 'GET / POST'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Logout user and invalidate session'],
])

doc.add_heading('4.3 SSO Login', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'auth/sso_login'],
    ['Method', 'POST'],
    ['Auth Required', 'No'],
    ['Description', 'Login via SSO provider (Keycloak, etc.)'],
])
doc.add_paragraph('Request Body:')
code_block(doc, '{\n  "username": "string",\n  "email": "string",\n  "first_name": "string",\n  "last_name": "string",\n  "sso_provider": "string",\n  "sso_id": "string",\n  "default_role": "string"\n}')

doc.add_heading('4.4 Get Current User', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'auth/me'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Returns current authenticated user info'],
])

doc.add_heading('4.5 Change Password', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'auth/change-password'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Change current user password'],
])

doc.add_heading('4.6 Initial Setup', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'auth/setup'],
    ['Method', 'POST'],
    ['Auth Required', 'No (first-run only)'],
    ['Description', 'Initial system setup and admin account creation'],
])

doc.add_page_break()

# ============================================================
# 5. LDAP
# ============================================================
doc.add_heading('5. LDAP Integration Endpoints', level=1)

doc.add_heading('5.1 Test LDAP Connection', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'ldap/test'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (Admin)'],
    ['Description', 'Test connectivity to LDAP server'],
])
doc.add_paragraph('Request Body:')
code_block(doc, '{\n  "url": "string",\n  "base_dn": "string",\n  "bind_dn": "string",\n  "bind_password": "string"\n}')

doc.add_heading('5.2 List LDAP Users', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'ldap/users'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (Admin)'],
    ['Description', 'Retrieve users from LDAP directory'],
])

doc.add_heading('5.3 Verify LDAP Credentials', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'ldap/verify'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes'],
    ['Description', 'Verify user credentials against LDAP'],
])

doc.add_heading('5.4 Sync LDAP Users', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'ldap/sync'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (Admin)'],
    ['Description', 'Synchronize users from LDAP into the platform'],
])

doc.add_page_break()

# ============================================================
# 6. USER MANAGEMENT
# ============================================================
doc.add_heading('6. User Management Endpoints', level=1)

doc.add_heading('6.1 List Users', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'users'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (users.view)'],
    ['Description', 'List all users with optional filtering'],
])
doc.add_paragraph('Query Parameters:')
add_table(doc, ['Parameter', 'Type', 'Description'],
[
    ['id', 'integer', 'Get specific user by ID'],
    ['q', 'string', 'Search by username, email, or name'],
    ['role', 'string', 'Filter by role code'],
    ['status', 'string', 'Filter: active | inactive | locked'],
])

doc.add_heading('6.2 Create User', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'users'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (users.create)'],
    ['Description', 'Create a new user account'],
])
doc.add_paragraph('Request Body:')
code_block(doc, '{\n  "username": "string (required)",\n  "email": "string (required)",\n  "password": "string (required)",\n  "first_name": "string",\n  "last_name": "string",\n  "roles": [1, 2]\n}')

doc.add_heading('6.3 Update User', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'users&id={id}'],
    ['Method', 'PUT'],
    ['Auth Required', 'Yes (users.edit)'],
    ['Description', 'Update existing user details'],
])
doc.add_paragraph('Request Body: email, first_name, last_name, phone, is_active, is_locked, roles')

doc.add_heading('6.4 Delete User', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'users&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes (users.delete)'],
    ['Description', 'Delete a user account'],
])

doc.add_page_break()

# ============================================================
# 7. ROLES & PERMISSIONS
# ============================================================
doc.add_heading('7. Role & Permission Endpoints', level=1)

doc.add_heading('7.1 List Roles', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'roles'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (roles.view)'],
    ['Description', 'List all roles; pass id for details with permissions'],
])

doc.add_heading('7.2 Create Role', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'roles'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (roles.create)'],
    ['Description', 'Create a new role with permissions'],
])
doc.add_paragraph('Request Body:')
code_block(doc, '{\n  "name": "string (required)",\n  "code": "string (required)",\n  "description": "string",\n  "permissions": [1, 2, 3]\n}')

doc.add_heading('7.3 Update Role', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'roles&id={id}'],
    ['Method', 'PUT'],
    ['Auth Required', 'Yes (roles.edit)'],
    ['Description', 'Update role name, description, and permissions'],
])

doc.add_heading('7.4 Delete Role', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'roles&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes (roles.delete)'],
    ['Description', 'Delete a role'],
])

doc.add_heading('7.5 List Permissions', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'permissions'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (roles.view)'],
    ['Description', 'List all available permissions grouped by module'],
])

doc.add_page_break()

# ============================================================
# 8. SESSION MANAGEMENT
# ============================================================
doc.add_heading('8. Session Management Endpoints', level=1)

doc.add_heading('8.1 List Sessions', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'sessions'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Get all active sessions for current user'],
])

doc.add_heading('8.2 Terminate Session', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'sessions&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Terminate a specific session. Use id=all to terminate all other sessions.'],
])

# ============================================================
# 9. PASSWORD POLICY
# ============================================================
doc.add_heading('9. Password Policy Endpoints', level=1)

doc.add_heading('9.1 Get Password Policy', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'password-policy'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Retrieve current password policy settings'],
])

doc.add_heading('9.2 Update Password Policy', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'password-policy'],
    ['Method', 'PUT'],
    ['Auth Required', 'Yes (settings.edit)'],
    ['Description', 'Update password policy settings'],
])

doc.add_page_break()

# ============================================================
# 10. BUSINESS GLOSSARY
# ============================================================
doc.add_heading('10. Business Glossary Endpoints', level=1)

doc.add_heading('10.1 List Glossary Terms', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'glossary_terms'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'List glossary terms with optional filtering'],
])
doc.add_paragraph('Query Parameters:')
add_table(doc, ['Parameter', 'Type', 'Description'],
[
    ['id', 'integer', 'Get specific term (includes history)'],
    ['q', 'string', 'Search by name or definition'],
    ['status', 'string', 'Filter: draft | approved | under_review | deleted'],
    ['domain', 'string', 'Filter by domain'],
])

doc.add_heading('10.2 Create Glossary Term', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'glossary_terms'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Create a new glossary term'],
])
doc.add_paragraph('Request Body:')
code_block(doc, '{\n  "name": "string (required)",\n  "definition": "string (required)",\n  "domain": "string",\n  "owner": "string",\n  "stewards": ["string"],\n  "status": "draft | approved | under_review",\n  "physicalAttributes": { ... },\n  "qualityRules": [ ... ],\n  "synonyms": ["string"],\n  "relatedTerms": ["string"],\n  "securityClassification": "string",\n  "dataType": "string",\n  "example": "string",\n  "formula": "string",\n  "businessLogic": "string",\n  "technicalDescription": "string",\n  "sourceSystem": "string",\n  "notes": "string"\n}')

doc.add_heading('10.3 Update Glossary Term', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'glossary_terms&id={id}'],
    ['Method', 'PUT'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Update an existing glossary term'],
])

doc.add_heading('10.4 Delete Glossary Term', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'glossary_terms&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Delete a glossary term. Use soft=true for soft delete.'],
])

doc.add_heading('10.5 Get Term History', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'glossary_history&id={id}'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Get change history for a specific term'],
])

doc.add_heading('10.6 Restore Deleted Term', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'restore_term&id={id}'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes (Session)'],
    ['Description', 'Restore a soft-deleted term'],
])

doc.add_page_break()

# ============================================================
# 11. DATA CATALOG
# ============================================================
doc.add_heading('11. Data Catalog Endpoints', level=1)

doc.add_heading('11.1 Layers', level=2)
doc.add_paragraph('Layers represent logical groupings of data tables (e.g., Raw, Staging, Curated).')

doc.add_heading('11.1.1 List Layers', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_layers'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'List all data layers. Pass id for specific layer with tables.'],
])

doc.add_heading('11.1.2 Create Layer', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_layers'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes'],
    ['Description', 'Create a new data layer'],
])
code_block(doc, '{\n  "name": "string (required)",\n  "description": "string",\n  "order": "integer"\n}')

doc.add_heading('11.1.3 Update Layer', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_layers&id={id}'],
    ['Method', 'PUT'],
    ['Auth Required', 'Yes'],
    ['Description', 'Update layer details'],
])

doc.add_heading('11.1.4 Delete Layer', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_layers&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes'],
    ['Description', 'Delete a layer'],
])

doc.add_heading('11.2 Tables', level=2)

doc.add_heading('11.2.1 List Tables', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_tables'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'List tables with columns. Filter by layer name.'],
])

doc.add_heading('11.2.2 Create Table', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_tables'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes'],
    ['Description', 'Register a new table in the catalog'],
])
code_block(doc, '{\n  "name": "string (required)",\n  "layer_id": "integer (required)",\n  "description": "string",\n  "row_count": "integer",\n  "owner": "string"\n}')

doc.add_heading('11.2.3 Update Table', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_tables&id={id}'],
    ['Method', 'PUT'],
    ['Auth Required', 'Yes'],
    ['Description', 'Update table metadata'],
])

doc.add_heading('11.2.4 Delete Table', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_tables&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes'],
    ['Description', 'Delete table and all its columns'],
])

doc.add_heading('11.3 Columns', level=2)

doc.add_heading('11.3.1 List Columns', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_columns'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'List columns. Filter by table_id.'],
])

doc.add_heading('11.3.2 Create Column', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_columns'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes'],
    ['Description', 'Add a column to a table'],
])
code_block(doc, '{\n  "name": "string (required)",\n  "table_id": "integer (required)",\n  "data_type": "string",\n  "description": "string",\n  "is_pk": "boolean",\n  "is_fk": "boolean",\n  "is_nullable": "boolean",\n  "order": "integer"\n}')

doc.add_heading('11.3.3 Delete Column', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'catalog_columns&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes'],
    ['Description', 'Delete a column'],
])

doc.add_heading('11.4 Table Profiling', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'profile_table'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'Get profiling data for a table (statistics, row counts, data samples)'],
])
doc.add_paragraph('Query Parameters: table (name), layer (name)')

doc.add_page_break()

# ============================================================
# 12. DATA MAPPING
# ============================================================
doc.add_heading('12. Data Mapping Endpoints', level=1)

doc.add_heading('12.1 Mappings', level=2)
add_table(doc, ['Endpoint', 'Method', 'Description'],
[
    ['mappings', 'GET', 'List all column mappings (pass id for specific)'],
    ['mappings', 'POST', 'Create a new mapping'],
    ['mappings&id={id}', 'PUT', 'Update mapping'],
    ['mappings&id={id}', 'DELETE', 'Delete mapping'],
])

doc.add_heading('12.2 Join Conditions', level=2)
add_table(doc, ['Endpoint', 'Method', 'Description'],
[
    ['join_conditions', 'GET', 'List join conditions'],
    ['join_conditions', 'POST', 'Create join condition'],
    ['join_conditions&id={id}', 'PUT', 'Update join condition'],
    ['join_conditions&id={id}', 'DELETE', 'Delete join condition'],
])

doc.add_page_break()

# ============================================================
# 13. REPORTING
# ============================================================
doc.add_heading('13. Reporting Endpoints', level=1)

doc.add_heading('13.1 Reports', level=2)

doc.add_heading('13.1.1 List Reports', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'reports'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'List reports with optional filtering'],
])
doc.add_paragraph('Query Parameters:')
add_table(doc, ['Parameter', 'Type', 'Description'],
[
    ['search', 'string', 'Search report name/description'],
    ['domain', 'integer', 'Filter by domain ID'],
    ['bi_tool', 'string', 'Filter: power_bi | tableau | looker | etc.'],
    ['status', 'string', 'Filter by certification status'],
])

doc.add_heading('13.1.2 Create Report', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'reports'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes'],
    ['Description', 'Register a new report'],
])
code_block(doc, '{\n  "report_name": "string (required)",\n  "domain_id": "integer",\n  "business_owner": "string",\n  "data_owner": "string",\n  "bi_tool": "string",\n  "report_type": "string",\n  "certification_status": "string"\n}')

doc.add_heading('13.1.3 Update Report', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'reports&id={id}'],
    ['Method', 'PUT'],
    ['Auth Required', 'Yes'],
    ['Description', 'Update report metadata'],
])

doc.add_heading('13.1.4 Delete Report', level=3)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'reports&id={id}'],
    ['Method', 'DELETE'],
    ['Auth Required', 'Yes'],
    ['Description', 'Delete a report'],
])

doc.add_heading('13.2 Report Sub-resources', level=2)
add_table(doc, ['Endpoint', 'Method', 'Description'],
[
    ['report_data_sources?report_id={id}', 'GET', 'Get report data sources'],
    ['report_data_sources', 'POST', 'Add data source to report'],
    ['report_kpis?report_id={id}', 'GET', 'Get report KPIs'],
    ['report_kpis', 'POST', 'Add KPI to report'],
    ['report_users?report_id={id}', 'GET', 'Get report users/viewers'],
    ['report_users', 'POST', 'Add user to report'],
    ['report_terms?report_id={id}', 'GET', 'Get glossary terms used in report'],
    ['search_reports?q={query}', 'GET', 'Search reports by name/description'],
])

doc.add_page_break()

# ============================================================
# 14. DOMAIN MANAGEMENT
# ============================================================
doc.add_heading('14. Domain Management Endpoints', level=1)

doc.add_heading('14.1 Domains', level=2)
add_table(doc, ['Endpoint', 'Method', 'Description'],
[
    ['domains', 'GET', 'List all domains (pass id for details with stakeholders)'],
    ['domains', 'POST', 'Create domain (body: name, icon, color, description)'],
    ['domains&id={id}', 'PUT', 'Update domain'],
    ['domains&id={id}', 'DELETE', 'Delete domain'],
])

doc.add_heading('14.2 Domain Stakeholders', level=2)
add_table(doc, ['Endpoint', 'Method', 'Description'],
[
    ['domain_stakeholders', 'GET', 'Get stakeholders (params: domain_id, role_id)'],
    ['domain_stakeholders', 'POST', 'Add stakeholder to domain'],
    ['domain_stakeholders', 'DELETE', 'Remove stakeholder from domain'],
])

doc.add_page_break()

# ============================================================
# 15. GOVERNANCE
# ============================================================
doc.add_heading('15. Governance Endpoints', level=1)

add_table(doc, ['Endpoint', 'Method', 'Description'],
[
    ['governance_roles', 'GET', 'List governance roles'],
    ['governance_roles', 'POST', 'Create governance role'],
    ['governance_steps', 'GET', 'Get workflow steps'],
    ['governance_steps', 'POST', 'Create workflow step'],
    ['governance_stakeholders', 'GET', 'Get governance stakeholders'],
    ['governance_stakeholders', 'POST', 'Add governance stakeholder'],
])

doc.add_page_break()

# ============================================================
# 16. EXTERNAL DATA SOURCES
# ============================================================
doc.add_heading('16. External Data Sources Endpoints', level=1)

doc.add_heading('16.1 Manage Sources', level=2)
add_table(doc, ['Endpoint', 'Method', 'Description'],
[
    ['external_sources', 'GET', 'List external database sources (pass id for details)'],
    ['external_sources', 'POST', 'Create external source connection'],
    ['external_sources&id={id}', 'PUT', 'Update source connection'],
    ['external_sources&id={id}', 'DELETE', 'Delete source'],
])
doc.add_paragraph('Create/Update Request Body:')
code_block(doc, '{\n  "name": "string (required)",\n  "db_type": "mysql | postgresql | oracle | mssql",\n  "host": "string",\n  "port": "integer",\n  "username": "string",\n  "password": "string",\n  "database_name": "string",\n  "sid": "string (Oracle only)"\n}')

doc.add_heading('16.2 Import Tables', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'external_sources_import'],
    ['Method', 'POST'],
    ['Auth Required', 'Yes'],
    ['Description', 'Import tables from external source into the data catalog'],
])
code_block(doc, '{\n  "source_id": "integer (required)",\n  "tables": [\n    {\n      "schema_name": "string",\n      "table_name": "string",\n      "columns": [ ... ],\n      "row_count": "integer"\n    }\n  ]\n}')

doc.add_heading('16.3 Get Connections', level=2)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'get_connections'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'Get available database connections (for schedule creation)'],
])

doc.add_page_break()

# ============================================================
# 17. PLATFORM STATISTICS
# ============================================================
doc.add_heading('17. Platform Statistics', level=1)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'stats'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'Get comprehensive platform statistics'],
])
doc.add_paragraph('Response includes counts and metrics for:')
items = [
    'Glossary: total, approved, draft, under_review',
    'Catalog: tables, columns, active, documented',
    'Mappings: total, validated, pending',
    'Quality: score, total checks, passed, failed',
    'Reports: total, certified, KPIs, quality score',
    'Users, Roles, External Sources, Domains',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 18. EXPORT
# ============================================================
doc.add_heading('18. Export', level=1)
add_table(doc, ['Property', 'Value'],
[
    ['Endpoint', 'export_all'],
    ['Method', 'GET'],
    ['Auth Required', 'Yes'],
    ['Description', 'Export all platform data as JSON (glossary, catalog, governance, mappings)'],
])

# ============================================================
# 19. ERROR HANDLING
# ============================================================
doc.add_heading('19. Error Handling', level=1)
doc.add_paragraph(
    'All error responses follow a consistent JSON format with appropriate HTTP status codes.'
)
doc.add_paragraph('Example error response:')
code_block(doc, '{\n  "success": false,\n  "error": "Detailed error message"\n}')

doc.add_paragraph('Common error scenarios:')
add_table(doc, ['Scenario', 'Status Code', 'Example Error'],
[
    ['Missing required field', '400', 'Field "name" is required'],
    ['Invalid session', '401', 'Authentication required'],
    ['Insufficient permissions', '403', 'Permission denied: catalog.create required'],
    ['Resource not found', '404', 'Term not found'],
    ['Wrong HTTP method', '405', 'Method not allowed'],
    ['Server error', '500', 'Internal server error'],
])

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\xampp\htdocs\datarover\API_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
