#!/usr/bin/env python3
"""Generate comprehensive Technical Specification DOCX for all 9 modules (AZ)."""

import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ---- Styles ----
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)
h1 = doc.styles['Heading 1']; h1.font.size = Pt(20); h1.font.color.rgb = RGBColor(0x1a,0x56,0xdb); h1.font.bold = True
h2 = doc.styles['Heading 2']; h2.font.size = Pt(15); h2.font.color.rgb = RGBColor(0x1a,0x56,0xdb)
h3 = doc.styles['Heading 3']; h3.font.size = Pt(12); h3.font.color.rgb = RGBColor(0x33,0x33,0x33)
h4 = doc.styles['Heading 4']; h4.font.size = Pt(11); h4.font.color.rgb = RGBColor(0x55,0x55,0x55)

def para(text, bold=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    return p

def code(text, lang=''):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5); pf.space_before = Pt(3); pf.space_after = Pt(6)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(8.5)
    doc.add_paragraph()

def section_break():
    doc.add_page_break()

# ============================================================================
# COVER
# ============================================================================
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Texniki Spesifikasiya'); r.font.size = Pt(34); r.bold=True; r.font.color.rgb=RGBColor(0x1a,0x56,0xdb)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Data Governance & Quality Platform — v56'); r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x55,0x55,0x55)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Backend Developer üçün Tam Modul Sənədləşməsi'); r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x77,0x77,0x77)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Versiya 1.0  |  {datetime.date.today().strftime("%B %Y")}'); r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
section_break()

# ============================================================================
# GLOBAL INTRO
# ============================================================================
doc.add_heading('Giriş', level=1)
para('Bu sənəd platformaya yeni qoşulan backend developerin modulları başa düşməsi və kod yazmağa başlaması üçün tam istinad materialıdır. Sənəd 9 əsas modulu əhatə edir və hər modul üçün ayrıca bölmə mövcuddur.')

doc.add_heading('Texnoloji Stek', level=2)
tbl(['Komponent','Texnologiya','Versiya / Qeyd'],[
    ['Backend','PHP','8.x / XAMPP runtime, backend.php tək giriş nöqtəsi'],
    ['Verilənlər bazası','MySQL','8.x, InnoDB, utf8mb4_unicode_ci'],
    ['Frontend','Vanilla JavaScript','app.js (~22,500 sətir), SPA'],
    ['Scheduler','Python','dq_schedule_simple.py (cron-style polling)'],
    ['Auth','Session + API Key + SSO (Keycloak) + LDAP',''],
    ['Zona','Asia/Baku (UTC+4)',''],
])

doc.add_heading('Routing Konvensiyası', level=2)
para('Bütün API çağırışları tək giriş nöqtəsindən keçir:')
code('backend.php?action=<endpoint_name>&id=<optional>&method=GET|POST|PUT|DELETE')
para('Cavab formatı həmişə JSON:')
code('{\n  "success": true,\n  "data": { ... }\n}\n\n// xəta:\n{\n  "success": false,\n  "error": "Error message"\n}')

doc.add_heading('İstifadəçi Rolları (RBAC)', level=2)
tbl(['Rol Kodu','Ad','Təyinat'],[
    ['SUPER_ADMIN','Super Administrator','Bütün icazələr, sistem tənzimləmələri'],
    ['GOV_MANAGER','Governance Manager','Glossary, domain, governance idarəetməsi'],
    ['QUALITY_MANAGER','Quality Manager','Data quality qaydaları və icraları'],
    ['DATA_STEWARD','Data Steward','Term və catalog elementlərinin təsdiqi'],
    ['DATA_ENGINEER','Data Engineer','Catalog, mapping, external source işləri'],
    ['ANALYST','Business Analyst','Report yaratma və baxış'],
    ['COMPLIANCE','Compliance Officer','Audit, governance nəzarəti'],
    ['VIEWER','Report Viewer','Yalnız oxuma icazəsi'],
])
para('İcazələr module.action formatında saxlanılır (məs: glossary.create, catalog.edit, quality.run). SUPER_ADMIN bütün yoxlamaları bypass edir.')

doc.add_heading('Modulların Siyahısı', level=2)
tbl(['№','Modul','Əsas cədvəllər'],[
    ['1','Authentication & RBAC','users, user_roles, permissions, user_sessions, api_keys'],
    ['2','Business Glossary','glossary_terms, glossary_term_history, term_physical_attributes'],
    ['3','Data Catalog','catalog_layers, catalog_tables, catalog_columns'],
    ['4','Data Quality','data_quality_rules, soda_checks, soda_scan_results, quality_rule_results'],
    ['5','Column Mapping','column_mappings, table_join_conditions'],
    ['6','Reporting','reports, report_data_sources, report_kpis, report_users, report_terms'],
    ['7','Domain Management','domains, domain_stakeholders'],
    ['8','Governance Workflow','governance_roles, governance_workflow_steps, governance_stakeholders'],
    ['9','External Data Sources','external_sources, external_source_tables'],
])
section_break()

# ============================================================================
# HELPERS FOR EACH MODULE
# ============================================================================
def module_header(num, name_az, name_en):
    doc.add_heading(f'{num}. Modul: {name_az} ({name_en})', level=1)

def sec(title):
    doc.add_heading(title, level=2)

def sub(title):
    doc.add_heading(title, level=3)

# ============================================================================
# MODULE 1: AUTHENTICATION & RBAC
# ============================================================================
module_header(1, 'Autentifikasiya və Rol-Əsaslı İcazələr', 'Authentication & RBAC')

sec('1.1 Module Overview')
para('Bu modul bütün platformanın təhlükəsizlik təməlidir: istifadəçi qeydiyyatı, giriş (local/SSO/LDAP), session idarəetməsi, parol siyasəti, rol-icazə matrisi və audit jurnalı. Əsas rollar: Super Admin (sistem), Data Steward (təsdiq), Data Engineer (texniki), Analyst (oxuma), Compliance (audit). Bütün digər modullar auth_middleware.php vasitəsilə bu moduldan keçir.')

sec('1.2 Data Model — Cədvəllər')

sub('1.2.1 users (MODULE-OWNED)')
para('İstifadəçi hesablarını saxlayır. auth_type sahəsi local/sso/ldap dəyərlərindən birini alır.')
code('''CREATE TABLE users (
  id INT PRIMARY KEY AUTO_INCREMENT,
  tenant_id INT,
  username VARCHAR(255) UNIQUE NOT NULL,
  email VARCHAR(255) UNIQUE NOT NULL,
  password_hash VARCHAR(255),
  first_name VARCHAR(255),
  last_name VARCHAR(255),
  phone VARCHAR(20),
  avatar_url VARCHAR(500),
  auth_type VARCHAR(50) NOT NULL DEFAULT 'local',
  is_active BOOLEAN DEFAULT TRUE,
  is_locked BOOLEAN DEFAULT FALSE,
  locked_until DATETIME NULL,
  must_change_password BOOLEAN DEFAULT FALSE,
  last_login DATETIME NULL,
  last_password_change DATETIME NULL,
  sso_provider VARCHAR(50) NULL,
  sso_id VARCHAR(255) NULL,
  created_by INT NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  INDEX idx_username (username),
  INDEX idx_email (email),
  INDEX idx_tenant_id (tenant_id),
  INDEX idx_auth_type (auth_type),
  CONSTRAINT fk_users_created_by FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE SET NULL
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('1.2.2 user_roles (MODULE-OWNED)')
code('''CREATE TABLE user_roles (
  id INT PRIMARY KEY AUTO_INCREMENT,
  tenant_id INT,
  name VARCHAR(255) NOT NULL,
  code VARCHAR(50) UNIQUE NOT NULL,
  description TEXT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_code (code),
  INDEX idx_tenant_id (tenant_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('1.2.3 permissions (MODULE-OWNED)')
code('''CREATE TABLE permissions (
  id INT PRIMARY KEY AUTO_INCREMENT,
  module VARCHAR(100) NOT NULL,
  action VARCHAR(100) NOT NULL,
  description VARCHAR(255),
  UNIQUE KEY unique_module_action (module, action),
  INDEX idx_module (module)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('1.2.4 role_permissions')
code('''CREATE TABLE role_permissions (
  id INT PRIMARY KEY AUTO_INCREMENT,
  role_id INT NOT NULL,
  permission_id INT NOT NULL,
  UNIQUE KEY unique_role_perm (role_id, permission_id),
  FOREIGN KEY (role_id) REFERENCES user_roles(id) ON DELETE CASCADE,
  FOREIGN KEY (permission_id) REFERENCES permissions(id) ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('1.2.5 user_role_assignments')
code('''CREATE TABLE user_role_assignments (
  id INT PRIMARY KEY AUTO_INCREMENT,
  user_id INT NOT NULL,
  role_id INT NOT NULL,
  assigned_by INT NULL,
  assigned_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  UNIQUE KEY unique_user_role (user_id, role_id),
  FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
  FOREIGN KEY (role_id) REFERENCES user_roles(id) ON DELETE CASCADE,
  FOREIGN KEY (assigned_by) REFERENCES users(id) ON DELETE SET NULL
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('1.2.6 user_sessions')
code('''CREATE TABLE user_sessions (
  id INT PRIMARY KEY AUTO_INCREMENT,
  user_id INT NOT NULL,
  session_token VARCHAR(255) UNIQUE NOT NULL,
  ip_address VARCHAR(45),
  user_agent VARCHAR(500),
  expires_at DATETIME NOT NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
  INDEX idx_token (session_token),
  INDEX idx_expires (expires_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('1.2.7 password_policies, password_history')
code('''CREATE TABLE password_policies (
  id INT PRIMARY KEY AUTO_INCREMENT,
  tenant_id INT,
  min_length INT DEFAULT 8,
  require_uppercase BOOLEAN DEFAULT TRUE,
  require_lowercase BOOLEAN DEFAULT TRUE,
  require_number BOOLEAN DEFAULT TRUE,
  require_special BOOLEAN DEFAULT TRUE,
  max_age_days INT DEFAULT 90,
  prevent_reuse_count INT DEFAULT 5,
  lockout_attempts INT DEFAULT 5,
  lockout_duration_minutes INT DEFAULT 30,
  session_timeout_minutes INT DEFAULT 60,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE password_history (
  id INT PRIMARY KEY AUTO_INCREMENT,
  user_id INT NOT NULL,
  password_hash VARCHAR(255),
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
  INDEX idx_user_id (user_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('1.2.8 api_keys, audit_log (SHARED)')
code('''CREATE TABLE api_keys (
  id INT PRIMARY KEY AUTO_INCREMENT,
  api_key VARCHAR(255) UNIQUE NOT NULL,
  user_id INT NULL,
  description VARCHAR(500),
  is_active BOOLEAN DEFAULT TRUE,
  expires_at DATETIME NULL,
  last_used DATETIME NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE SET NULL,
  INDEX idx_api_key (api_key)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE audit_log (
  id INT PRIMARY KEY AUTO_INCREMENT,
  tenant_id INT,
  user_id INT,
  username VARCHAR(255),
  action VARCHAR(255),
  module VARCHAR(100),
  entity_type VARCHAR(100),
  entity_id INT,
  old_values JSON,
  new_values JSON,
  ip_address VARCHAR(45),
  user_agent VARCHAR(500),
  status VARCHAR(50),
  error_message TEXT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_user_id (user_id),
  INDEX idx_created_at (created_at),
  INDEX idx_entity (entity_type, entity_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('1.3 ER Diagram')
code('''erDiagram
    users ||--o{ user_role_assignments : has
    user_roles ||--o{ user_role_assignments : assigned_to
    user_roles ||--o{ role_permissions : has
    permissions ||--o{ role_permissions : granted_via
    users ||--o{ user_sessions : owns
    users ||--o{ password_history : records
    users ||--o{ api_keys : generates
    users ||--o{ audit_log : triggers''')

sec('1.4 Əsas SQL Sorğuları')

sub('1.4.1 Login — password verification')
code('''SELECT u.*, GROUP_CONCAT(ur.code) as roles
FROM users u
LEFT JOIN user_role_assignments ura ON u.id = ura.user_id
LEFT JOIN user_roles ur ON ura.role_id = ur.id
WHERE u.username = :username
  AND u.is_active = 1
  AND (u.is_locked = 0 OR u.locked_until < NOW())
GROUP BY u.id;
-- sonra PHP tərəfində password_verify($password, $row['password_hash'])''')

sub('1.4.2 Session yaratma')
code('''INSERT INTO user_sessions (user_id, session_token, ip_address, user_agent, expires_at)
VALUES (:user_id, :token, :ip, :agent, DATE_ADD(NOW(), INTERVAL :mins MINUTE));''')

sub('1.4.3 Session validation (hər request)')
code('''SELECT s.user_id, s.expires_at, u.username, u.is_active
FROM user_sessions s
JOIN users u ON u.id = s.user_id
WHERE s.session_token = :token
  AND s.expires_at > NOW()
  AND u.is_active = 1
LIMIT 1;''')

sub('1.4.4 İcazə yoxlaması')
code('''SELECT COUNT(*) AS allowed
FROM user_role_assignments ura
JOIN role_permissions rp ON rp.role_id = ura.role_id
JOIN permissions p ON p.id = rp.permission_id
WHERE ura.user_id = :uid
  AND p.module = :module
  AND p.action = :action;''')

sub('1.4.5 Failed login counter + lockout')
code('''-- pseudo-transaction
START TRANSACTION;
  UPDATE users SET
    failed_attempts = failed_attempts + 1,
    is_locked = IF(failed_attempts + 1 >= :max_attempts, 1, 0),
    locked_until = IF(failed_attempts + 1 >= :max_attempts,
                      DATE_ADD(NOW(), INTERVAL :lockout_mins MINUTE), locked_until)
  WHERE username = :username;
COMMIT;''')

sub('1.4.6 Password history — re-use yoxlaması')
code('''SELECT password_hash FROM password_history
WHERE user_id = :uid
ORDER BY created_at DESC
LIMIT :reuse_count;
-- PHP: loop edib password_verify($new, $hash) hər biri ilə''')

sec('1.5 Business Logic Qaydaları')
tbl(['Qayda','Təsvir'],[
    ['Session TTL','password_policies.session_timeout_minutes (default 60 dəq)'],
    ['Lockout','lockout_attempts uğursuz girişdən sonra locked_until = NOW() + lockout_duration_minutes'],
    ['Password reuse','Son N parolun hash-i ilə müqayisə (prevent_reuse_count)'],
    ['Password age','last_password_change + max_age_days < NOW() → must_change_password = TRUE'],
    ['SSO auto-provision','İstifadəçi mövcud deyilsə, default_role ilə yaradılır (auth_type=sso)'],
    ['LDAP sync','ldap/sync yalnız yeni istifadəçilər yaradır, mövcudları update edir'],
])

para('Parol kompleksliyi (regex):')
code('''// PHP validation
function validatePassword($pwd, $policy) {
  if (strlen($pwd) < $policy['min_length']) return false;
  if ($policy['require_uppercase'] && !preg_match('/[A-Z]/', $pwd)) return false;
  if ($policy['require_lowercase'] && !preg_match('/[a-z]/', $pwd)) return false;
  if ($policy['require_number']   && !preg_match('/[0-9]/', $pwd)) return false;
  if ($policy['require_special']  && !preg_match('/[^A-Za-z0-9]/', $pwd)) return false;
  return true;
}''')

sec('1.6 REST API Endpoint-ləri')

sub('1.6.1 POST auth/login')
tbl(['Sahə','Dəyər'],[
    ['Yol','backend.php?action=auth/login'],
    ['Auth','Tələb olunmur'],
    ['İstifadə edilən SQL','1.4.1, 1.4.2, 1.4.5'],
])
para('Request:')
code('{\n  "username": "steward01",\n  "password": "Demo@123"\n}')
para('Response (200):')
code('''{
  "success": true,
  "data": {
    "token": "sess_a4f...",
    "user": { "id": 5, "username": "steward01", "email": "s@x.com" },
    "roles": ["DATA_STEWARD"],
    "permissions": ["glossary.view", "glossary.edit", ...]
  }
}''')
para('cURL:')
code('''curl -X POST "http://localhost/datarover/backend.php?action=auth/login" \\
  -H "Content-Type: application/json" \\
  -d '{"username":"steward01","password":"Demo@123"}' ''')
para('Frontend binding: app.js → login(username, password) (sətir 159–191)')

sub('1.6.2 POST auth/logout')
para('Body tələb olunmur. Authorization: Bearer <token> header-i vacibdir.')
code('curl -X POST "http://localhost/datarover/backend.php?action=auth/logout" -H "Authorization: Bearer $TOKEN"')
para('Frontend: app.js → logout() (193–215)')

sub('1.6.3 POST auth/sso_login')
para('Keycloak callback üçün. İstifadəçi yoxdursa avtomatik yaradılır.')
code('''{
  "username": "john.doe",
  "email": "john@corp.com",
  "first_name": "John",
  "last_name": "Doe",
  "sso_provider": "keycloak",
  "sso_id": "uuid-from-keycloak",
  "default_role": "ANALYST"
}''')

sub('1.6.4 POST auth/me')
para('Cari istifadəçi məlumatı + icazələr. UI permission yoxlamaları üçün əsas mənbə.')

sub('1.6.5 POST auth/change-password')
code('''{
  "current_password": "...",
  "new_password": "..."
}''')
para('Xətalar: 400 (zəif parol), 403 (cari parol səhv), 409 (parol yenidən istifadə).')

sub('1.6.6 CRUD users')
tbl(['Method','Yol','İcazə','Təsvir'],[
    ['GET','users','users.view','Siyahı / axtarış (q, role, status filter)'],
    ['GET','users&id={id}','users.view','Bir istifadəçi + rolları'],
    ['POST','users','users.create','Yeni istifadəçi'],
    ['PUT','users&id={id}','users.edit','Yeniləmə (rolları da yeniləyir)'],
    ['DELETE','users&id={id}','users.delete','Silmə (SET NULL audit_log-da)'],
])

sub('1.6.7 CRUD roles & GET permissions')
tbl(['Method','Yol','İcazə'],[
    ['GET','roles','roles.view'],
    ['POST','roles','roles.create'],
    ['PUT','roles&id={id}','roles.edit'],
    ['DELETE','roles&id={id}','roles.delete'],
    ['GET','permissions','roles.view'],
])

sub('1.6.8 Sessions & Password Policy & Audit')
tbl(['Method','Yol','Təyinat'],[
    ['GET','sessions','Cari user-in aktiv sessiyaları'],
    ['DELETE','sessions&id={id}','Sessiya terminasiyası (id=all → hamısı)'],
    ['GET','password-policy','Parol siyasəti'],
    ['PUT','password-policy','settings.edit tələb olunur'],
    ['GET','audit','audit.view — page, limit, user_id, module, date_from/to'],
])

sub('1.6.9 LDAP endpoint-ləri')
tbl(['Method','Yol','Təsvir'],[
    ['POST','ldap/test','LDAP bağlantı testi'],
    ['POST','ldap/users','Base DN-dən istifadəçi siyahısı'],
    ['POST','ldap/verify','Credential yoxlama'],
    ['POST','ldap/sync','Toplu sync (mövcud olanı yeniləyir, yenisini yaradır)'],
])

sec('1.7 Python Scheduler Contract')
para('Bu modul scheduler istifadə etmir. Session expiration və lockout_until lazy-evaluation ilə hər request-də yoxlanılır. Köhnə session-ların fiziki silinməsi üçün ayrı bir cron task planlaşdırılmalıdır:')
code('DELETE FROM user_sessions WHERE expires_at < DATE_SUB(NOW(), INTERVAL 7 DAY);')

sec('1.8 Permission Matrix')
tbl(['Action','SUPER_ADMIN','GOV_MGR','Q_MGR','STEWARD','ENGINEER','ANALYST','COMPLIANCE','VIEWER'],[
    ['users.view','✓','✓','—','—','—','—','✓','—'],
    ['users.create','✓','—','—','—','—','—','—','—'],
    ['users.edit','✓','—','—','—','—','—','—','—'],
    ['users.delete','✓','—','—','—','—','—','—','—'],
    ['roles.edit','✓','—','—','—','—','—','—','—'],
    ['audit.view','✓','✓','—','—','—','—','✓','—'],
    ['settings.edit','✓','—','—','—','—','—','—','—'],
])

sec('1.9 Fayl Strukturu')
tbl(['Fayl','Sətir aralığı','Məzmun'],[
    ['backend.php','1263–1979','Bütün auth/* handler-ləri'],
    ['backend.php','1985–2521','users, roles, permissions, sessions, audit, password-policy'],
    ['auth_middleware.php','tam','requireAuth, requirePermission, hasPermission, hasRole'],
    ['js/app.js','112–334','Frontend auth flow'],
    ['js/permissions.js','tam','UI-da permission yoxlaması'],
])

sec('1.10 İnteqrasiya Nöqtələri')
para('Bütün digər modullar auth_middleware::requireAuth() və requirePermission() ilə qorunur. audit_log cədvəli bütün modullardan INSERT qəbul edir (SHARED). api_keys Data Quality xarici API-si (Great Expectations, dbt, Soda) tərəfindən istifadə olunur.')

sec('1.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','Migration SQL-lərini icra et (users → permissions → role_permissions → user_role_assignments ardıcıllığı ilə)'],
    ['2','auth/setup endpoint-i ilə ilk SUPER_ADMIN yarat'],
    ['3','Demo user seed (14 nəfər, Demo@123) — RBAC_README.md'],
    ['4','Verification: SELECT COUNT(*) FROM users; SELECT COUNT(*) FROM permissions;'],
    ['5','Login testi: curl ilə auth/login, response-də token olmalıdır'],
])

section_break()

# ============================================================================
# MODULE 2: BUSINESS GLOSSARY
# ============================================================================
module_header(2, 'Biznes Lüğəti', 'Business Glossary')

sec('2.1 Module Overview')
para('Business Glossary biznes terminlərini (KPI, metrik, entity) onların tərifləri, sahibləri, fiziki atributları və keyfiyyət qaydaları ilə birlikdə saxlayır. Workflow: draft → under_review → approved (və ya rejected). Hər dəyişiklik glossary_term_history-də saxlanılır. Əsas rollar: Data Steward (yaradır, redaktə edir), Governance Manager (təsdiq edir), Viewer (oxuyur).')

sec('2.2 Data Model')

sub('2.2.1 glossary_terms (MODULE-OWNED)')
code('''CREATE TABLE glossary_terms (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255) NOT NULL UNIQUE,
  abbreviation VARCHAR(100),
  definition TEXT,
  domain VARCHAR(255),
  data_type VARCHAR(100),
  example TEXT,
  formula TEXT,
  business_logic TEXT,
  technical_description TEXT,
  owner VARCHAR(255),
  stewards JSON,
  security_classification VARCHAR(100),
  physical_attributes JSON,
  quality_rules JSON,
  synonyms JSON,
  related_terms JSON,
  source_system VARCHAR(255),
  notes TEXT,
  history JSON,
  status VARCHAR(50) DEFAULT 'draft',
  reject_reason TEXT,
  process_comments TEXT,
  deleted_at DATETIME NULL,
  deleted_by VARCHAR(255) NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  INDEX idx_name (name),
  INDEX idx_status (status),
  INDEX idx_domain (domain),
  INDEX idx_deleted (deleted_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')
para('Qeyd: status mümkün dəyərlər — draft, under_review, approved, rejected, deleted. Soft delete üçün deleted_at istifadə edilir.')

sub('2.2.2 glossary_term_history')
code('''CREATE TABLE glossary_term_history (
  id INT PRIMARY KEY AUTO_INCREMENT,
  term_id INT NOT NULL,
  action VARCHAR(100),
  comment TEXT,
  user VARCHAR(255),
  from_status VARCHAR(50),
  to_status VARCHAR(50),
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE CASCADE,
  INDEX idx_term_id (term_id),
  INDEX idx_created_at (created_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('2.2.3 term_physical_attributes')
code('''CREATE TABLE term_physical_attributes (
  id INT PRIMARY KEY AUTO_INCREMENT,
  term_id INT NOT NULL,
  layer VARCHAR(255),
  table_name VARCHAR(255),
  column_name VARCHAR(255),
  attribute_value TEXT,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE CASCADE,
  INDEX idx_term (term_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('2.3 ER Diagram')
code('''erDiagram
    glossary_terms ||--o{ glossary_term_history : changes
    glossary_terms ||--o{ term_physical_attributes : maps
    glossary_terms }o--|| domains : belongs_to
    glossary_terms ||--o{ report_terms : used_in
    catalog_columns }o--o{ glossary_terms : semantic_link''')

sec('2.4 Əsas SQL Sorğuları')

sub('2.4.1 Axtarış + filter')
code('''SELECT id, name, definition, domain, status, owner, updated_at
FROM glossary_terms
WHERE deleted_at IS NULL
  AND (:q IS NULL OR name LIKE CONCAT('%', :q, '%') OR definition LIKE CONCAT('%', :q, '%'))
  AND (:status IS NULL OR status = :status)
  AND (:domain IS NULL OR domain = :domain)
ORDER BY updated_at DESC
LIMIT 100;''')

sub('2.4.2 Term detal + tarixçə')
code('''SELECT t.*,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT(
      'action', action, 'comment', comment, 'user', user,
      'from_status', from_status, 'to_status', to_status, 'at', created_at))
   FROM glossary_term_history WHERE term_id = t.id) AS history
FROM glossary_terms t
WHERE t.id = :id;''')

sub('2.4.3 Status transition (transaction)')
code('''START TRANSACTION;
  UPDATE glossary_terms
  SET status = :new_status, updated_at = NOW()
  WHERE id = :id AND status = :expected_status;

  INSERT INTO glossary_term_history
    (term_id, action, comment, user, from_status, to_status)
  VALUES (:id, 'status_change', :comment, :user, :expected_status, :new_status);
COMMIT;''')

sub('2.4.4 Soft delete & restore')
code('''-- soft delete
UPDATE glossary_terms
SET status = 'deleted', deleted_at = NOW(), deleted_by = :user
WHERE id = :id AND deleted_at IS NULL;

-- restore
UPDATE glossary_terms
SET status = 'draft', deleted_at = NULL, deleted_by = NULL
WHERE id = :id AND deleted_at IS NOT NULL;''')

sec('2.5 Business Logic Qaydaları')
para('Status lifecycle:')
tbl(['From','To','Kim','Şərt'],[
    ['draft','under_review','Steward','name + definition boş olmamalı'],
    ['under_review','approved','Gov Manager','Minimum 1 stakeholder approval'],
    ['under_review','rejected','Gov Manager','reject_reason məcburidir'],
    ['rejected','draft','Steward','Resubmit zamanı'],
    ['approved','under_review','Steward','Dəyişiklik təklifi (re-review)'],
    ['*','deleted','Admin/Gov','Soft delete'],
])

para('Validation qaydaları:')
tbl(['Sahə','Qayda'],[
    ['name','Unikal, 3–255 simvol, boş ola bilməz'],
    ['definition','Minimum 10 simvol'],
    ['domain','domains cədvəlində mövcud olmalı (və ya NULL)'],
    ['status','Enum: draft, under_review, approved, rejected, deleted'],
    ['physical_attributes','JSON array — hər element {layer, table, column}'],
    ['quality_rules','JSON array — rule_id refs (data_quality_rules-də mövcud olmalı)'],
])

sec('2.6 REST API Endpoint-ləri')

sub('2.6.1 GET glossary_terms')
tbl(['Parametr','Tip','Təsvir'],[
    ['q','string','name/definition axtarışı'],
    ['status','string','draft|under_review|approved|rejected|deleted'],
    ['domain','string','Domain filteri'],
    ['id','int','Tək term (tarixçə ilə)'],
])

sub('2.6.2 POST glossary_terms')
code('''{
  "name": "Customer Lifetime Value",
  "abbreviation": "CLV",
  "definition": "Müştərinin platforma üzrə...",
  "domain": "Sales",
  "owner": "ceo@corp.com",
  "stewards": ["steward1@corp.com"],
  "physical_attributes": [
    {"layer":"dwh","table":"customer","column":"ltv_score"}
  ],
  "quality_rules": ["RULE-001","RULE-042"],
  "synonyms": ["LTV"],
  "related_terms": [15, 27],
  "status": "draft"
}''')

sub('2.6.3 PUT glossary_terms&id={id}')
para('Eyni body. Status dəyişsə history-yə INSERT edir (2.4.3).')

sub('2.6.4 DELETE glossary_terms&id={id}')
para('Default soft delete. ?soft=false → hard delete (yalnız Super Admin).')

sub('2.6.5 GET glossary_history&id={id}')
para('Term üçün bütün dəyişiklik tarixçəsi.')

sub('2.6.6 GET restore_term&id={id}')
para('Silinmiş termi bərpa edir.')

para('cURL nümunəsi:')
code('''curl "http://localhost/datarover/backend.php?action=glossary_terms&status=approved" \\
  -H "Authorization: Bearer $TOKEN"''')

sec('2.7 Python Scheduler Contract')
para('Glossary modulu scheduler istifadə etmir. Bütün status keçidləri manual (UI) və ya API ilə idarə olunur. Gələcəkdə stale term xəbərdarlığı üçün cron əlavə oluna bilər:')
code('''SELECT id, name, owner FROM glossary_terms
WHERE status = 'under_review'
  AND updated_at < DATE_SUB(NOW(), INTERVAL 30 DAY);''')

sec('2.8 Permission Matrix')
tbl(['Action','SUPER_ADMIN','GOV_MGR','STEWARD','ENGINEER','ANALYST','VIEWER'],[
    ['glossary.view','✓','✓','✓','✓','✓','✓'],
    ['glossary.create','✓','✓','✓','—','—','—'],
    ['glossary.edit','✓','✓','✓','—','—','—'],
    ['glossary.approve','✓','✓','—','—','—','—'],
    ['glossary.delete','✓','✓','—','—','—','—'],
    ['glossary.restore','✓','✓','—','—','—','—'],
])

sec('2.9 Fayl Strukturu')
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','2522–2706','CRUD handler'],
    ['backend.php','3864–3907','restore_term, glossary_history'],
    ['js/app.js','3567–3705','loadGlossary()'],
    ['js/app.js','4206–4310','saveTermToDB, updateTermInDB'],
    ['js/app.js','4903–5168','Status workflow (advance, reject, resubmit)'],
])

sec('2.10 İnteqrasiya Nöqtələri')
tbl(['Modul','Əlaqə'],[
    ['Data Catalog','physical_attributes JSON → catalog_columns-a link'],
    ['Data Quality','quality_rules JSON → data_quality_rules.rule_id refs'],
    ['Reports','report_terms cədvəli vasitəsilə link'],
    ['Domain Mgmt','domain sahəsi domains.name ilə match edir'],
    ['Governance','Status workflow governance_workflow_steps ilə idarə olunur'],
])

sec('2.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE glossary_terms, glossary_term_history, term_physical_attributes'],
    ['2','Seed: ilk 8–10 bazaya məxsus termin'],
    ['3','GOV_MGR, STEWARD rollarına glossary.* icazələri ver'],
    ['4','Verification: SELECT status, COUNT(*) FROM glossary_terms GROUP BY status;'],
    ['5','UI test: Business Glossary tab → Add Term → status workflow'],
])

section_break()

# ============================================================================
# MODULE 3: DATA CATALOG
# ============================================================================
module_header(3, 'Data Kataloqu', 'Data Catalog')

sec('3.1 Module Overview')
para('Data Catalog platformanın metadata repozitoriyasıdır. Hər data varlığı 3 səviyyəli iyerarxiyada saxlanılır: Layer (məs: Raw, Staging, Mart) → Table → Column. External mənbələrdən cədvəllər idxal olunur və live profilləşdirilir (row count, null%, min/max, distinct values). Əsas rollar: Data Engineer (idxal və profilləşdirmə), Data Steward (metadata zənginləşdirməsi), Analyst (oxuma).')

sec('3.2 Data Model')

sub('3.2.1 catalog_layers')
code('''CREATE TABLE catalog_layers (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255) UNIQUE NOT NULL,
  icon VARCHAR(50),
  color VARCHAR(10),
  description TEXT,
  `order` INT DEFAULT 0,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_order (`order`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('3.2.2 catalog_tables')
code('''CREATE TABLE catalog_tables (
  id INT PRIMARY KEY AUTO_INCREMENT,
  layer_id INT NOT NULL,
  name VARCHAR(500) NOT NULL,
  description TEXT,
  row_count BIGINT DEFAULT 0,
  owner VARCHAR(255),
  source_id INT NULL,
  schema_name VARCHAR(255),
  object_type VARCHAR(50) DEFAULT 'TABLE',
  is_active BOOLEAN DEFAULT TRUE,
  last_profiled_at DATETIME NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  UNIQUE KEY unique_layer_table (layer_id, name),
  FOREIGN KEY (layer_id) REFERENCES catalog_layers(id) ON DELETE CASCADE,
  FOREIGN KEY (source_id) REFERENCES external_sources(id) ON DELETE SET NULL,
  INDEX idx_name (name),
  INDEX idx_source (source_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('3.2.3 catalog_columns')
code('''CREATE TABLE catalog_columns (
  id INT PRIMARY KEY AUTO_INCREMENT,
  table_id INT NOT NULL,
  name VARCHAR(255) NOT NULL,
  data_type VARCHAR(100),
  description TEXT,
  is_pk BOOLEAN DEFAULT FALSE,
  is_fk BOOLEAN DEFAULT FALSE,
  is_nullable BOOLEAN DEFAULT TRUE,
  default_value VARCHAR(255),
  icon VARCHAR(50),
  `order` INT DEFAULT 0,
  profile_stats JSON,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  UNIQUE KEY unique_table_column (table_id, name),
  FOREIGN KEY (table_id) REFERENCES catalog_tables(id) ON DELETE CASCADE,
  INDEX idx_order (`order`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')
para('profile_stats JSON sxemi:')
code('''{
  "null_count": 142,
  "null_percent": 1.4,
  "distinct_count": 9855,
  "min": "2020-01-01",
  "max": "2024-12-31",
  "sample_values": ["A","B","C"],
  "profiled_at": "2026-04-19 10:00:00"
}''')

sec('3.3 ER Diagram')
code('''erDiagram
    catalog_layers ||--o{ catalog_tables : contains
    catalog_tables ||--o{ catalog_columns : has
    external_sources ||--o{ catalog_tables : sourced_from
    catalog_tables ||--o{ column_mappings : source_or_target
    catalog_columns }o--o{ glossary_terms : semantic_ref''')

sec('3.4 Əsas SQL Sorğuları')

sub('3.4.1 Layer tree (nested)')
code('''SELECT l.id, l.name, l.icon, l.color,
       (SELECT JSON_ARRAYAGG(JSON_OBJECT(
           'id', t.id, 'name', t.name, 'row_count', t.row_count))
        FROM catalog_tables t WHERE t.layer_id = l.id) AS tables
FROM catalog_layers l
ORDER BY l.`order`;''')

sub('3.4.2 Table + columns join')
code('''SELECT t.*, l.name AS layer_name,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT(
      'id',c.id,'name',c.name,'data_type',c.data_type,
      'is_pk',c.is_pk,'is_nullable',c.is_nullable))
   FROM catalog_columns c WHERE c.table_id = t.id ORDER BY c.`order`) AS columns
FROM catalog_tables t
JOIN catalog_layers l ON l.id = t.layer_id
WHERE t.id = :id;''')

sub('3.4.3 Live profiling (information_schema)')
code('''-- External source (MySQL) üçün column metadata
SELECT column_name, data_type, is_nullable, column_default, column_key
FROM information_schema.columns
WHERE table_schema = :schema AND table_name = :table
ORDER BY ordinal_position;

-- Row count
SELECT COUNT(*) AS cnt FROM `:schema`.`:table`;

-- Null stats üçün dinamik SQL
SELECT
  SUM(CASE WHEN `:col` IS NULL THEN 1 ELSE 0 END) AS null_cnt,
  COUNT(DISTINCT `:col`) AS distinct_cnt,
  MIN(`:col`) AS min_v,
  MAX(`:col`) AS max_v
FROM `:schema`.`:table`;''')

sub('3.4.4 External import (batch)')
code('''INSERT INTO catalog_tables (layer_id, name, schema_name, source_id, row_count)
VALUES (:layer, :name, :schema, :source, :rows)
ON DUPLICATE KEY UPDATE
  schema_name = VALUES(schema_name),
  source_id   = VALUES(source_id),
  row_count   = VALUES(row_count),
  updated_at  = NOW();

-- sonra columns üçün loop
INSERT INTO catalog_columns (table_id, name, data_type, is_pk, is_nullable, `order`)
VALUES (:tid, :cname, :dtype, :pk, :nullable, :idx)
ON DUPLICATE KEY UPDATE data_type = VALUES(data_type);''')

sec('3.5 Business Logic Qaydaları')
tbl(['Qayda','Təsvir'],[
    ['Layer order','UI-da layer-lər `order` ASC ilə göstərilir'],
    ['Unique constraint','Eyni layer daxilində table adı unikaldır'],
    ['Profile staleness','last_profiled_at > 7 gün əvvəldirsə UI-da "stale" göstərilir'],
    ['Hard delete cascade','Layer silinsə → bütün table + column da silinir'],
    ['External link','source_id NULL olmayan table-ın adı dəyişsə sync zamanı yenidən yaradılır'],
])

sec('3.6 REST API Endpoint-ləri')

sub('3.6.1 Layers')
tbl(['Method','Path','İcazə'],[
    ['GET','catalog_layers','catalog.view'],
    ['GET','catalog_layers&id={id}','catalog.view (table list ilə)'],
    ['POST','catalog_layers','catalog.create'],
    ['PUT','catalog_layers&id={id}','catalog.edit'],
    ['DELETE','catalog_layers&id={id}','catalog.delete'],
])

sub('3.6.2 Tables')
code('''POST body:
{
  "name": "fact_sales",
  "layer_id": 2,
  "description": "Daily sales facts",
  "row_count": 1500000,
  "owner": "dwh_team"
}''')

sub('3.6.3 Columns')
code('''POST body:
{
  "table_id": 45,
  "name": "customer_id",
  "data_type": "INT",
  "is_pk": false,
  "is_fk": true,
  "is_nullable": false,
  "description": "FK → dim_customer.id",
  "order": 3
}''')

sub('3.6.4 GET profile_table')
para('Query: table (name), layer (name). External source-a bağlanır və live profil qaytarır. Cache-lənmir.')
code('''curl "http://localhost/datarover/backend.php?action=profile_table&table=fact_sales&layer=Mart" \\
  -H "Authorization: Bearer $TOKEN"''')

para('Frontend binding: app.js → loadCatalog() təxminən 4000-ci sətirlər.')

sec('3.7 Python Scheduler Contract')
para('Catalog profiling bu anda sinxron endpoint (profile_table) üzərindən işləyir. Böyük table-lar üçün async profiling dq_schedule_simple.py-a əlavə oluna bilər:')
code('''-- cədvəl
CREATE TABLE catalog_profile_jobs (
  id INT PRIMARY KEY AUTO_INCREMENT,
  table_id INT NOT NULL,
  status ENUM('pending','running','done','error') DEFAULT 'pending',
  requested_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  started_at DATETIME NULL,
  finished_at DATETIME NULL,
  error TEXT,
  FOREIGN KEY (table_id) REFERENCES catalog_tables(id) ON DELETE CASCADE,
  INDEX idx_status (status, requested_at)
);

-- scheduler pickup
SELECT * FROM catalog_profile_jobs
WHERE status = 'pending'
ORDER BY requested_at
LIMIT 5 FOR UPDATE SKIP LOCKED;''')

sec('3.8 Permission Matrix')
tbl(['Action','SUPER_ADMIN','ENGINEER','STEWARD','ANALYST','VIEWER'],[
    ['catalog.view','✓','✓','✓','✓','✓'],
    ['catalog.create','✓','✓','—','—','—'],
    ['catalog.edit','✓','✓','✓','—','—'],
    ['catalog.delete','✓','✓','—','—','—'],
    ['catalog.profile','✓','✓','—','—','—'],
    ['catalog.import','✓','✓','—','—','—'],
])

sec('3.9 Fayl Strukturu')
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','3913–4237','Layers, tables, columns, profile_table'],
    ['backend.php','5219–5328','external_sources_import (catalog-a idxal)'],
    ['js/app.js','təxminən 4000–4250','Catalog UI'],
])

sec('3.10 İnteqrasiya Nöqtələri')
tbl(['Modul','Əlaqə'],[
    ['External Sources','source_id FK → external_sources.id'],
    ['Glossary','physical_attributes JSON → catalog_columns ref'],
    ['Mapping','column_mappings source/target table+column ref'],
    ['Quality','soda_checks.table_name + column_name → catalog-a ref'],
    ['Reports','report_data_sources.table_id FK → catalog_tables.id'],
])

sec('3.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE catalog_layers, catalog_tables, catalog_columns'],
    ['2','Seed 4–5 layer: Raw, Staging, Warehouse, Mart, Sandbox'],
    ['3','External source əlavə et və external_sources_import testlə'],
    ['4','profile_table endpoint testi'],
    ['5','SELECT l.name, COUNT(t.id) FROM catalog_layers l LEFT JOIN catalog_tables t ON t.layer_id=l.id GROUP BY l.id;'],
])

section_break()

# ============================================================================
# MODULE 4: DATA QUALITY
# ============================================================================
module_header(4, 'Data Keyfiyyəti', 'Data Quality')

sec('4.1 Module Overview')
para('Data Quality modulu catalog-dakı table/column-lara qarşı icra olunan keyfiyyət qaydalarını saxlayır və onların nəticələrini toplayır. İki paralel sistem mövcuddur: (1) native data_quality_rules (xarici alətlərdən — Great Expectations, dbt — nəticə qəbul edir), (2) Soda-integrated soda_checks (daxili icra, custom SQL). Icra dq_schedule_simple.py vasitəsilə sinxron (endpoint) və ya planlı (cron-style) ola bilər. Rollar: Quality Manager (qayda yaradır), Steward (nəzarət), Engineer (icra).')

sec('4.2 Data Model')

sub('4.2.1 data_quality_rules')
code('''CREATE TABLE data_quality_rules (
  id INT PRIMARY KEY AUTO_INCREMENT,
  rule_id VARCHAR(255) UNIQUE NOT NULL,
  name VARCHAR(255) NOT NULL,
  type VARCHAR(100) COMMENT 'completeness, uniqueness, validity, accuracy, timeliness, consistency',
  severity VARCHAR(50) DEFAULT 'medium' COMMENT 'low, medium, high, critical',
  description TEXT,
  expression TEXT COMMENT 'Custom SQL or rule definition',
  table_name VARCHAR(255),
  column_name VARCHAR(255),
  linked_columns JSON,
  status VARCHAR(50) DEFAULT 'active',
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  INDEX idx_rule_id (rule_id),
  INDEX idx_type (type),
  INDEX idx_status (status)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('4.2.2 quality_rule_results')
code('''CREATE TABLE quality_rule_results (
  id INT PRIMARY KEY AUTO_INCREMENT,
  rule_id VARCHAR(255) NOT NULL,
  pass_rate DECIMAL(5,2),
  total_records INT,
  passed_records INT,
  failed_records INT,
  run_date DATETIME,
  source_system VARCHAR(255),
  table_name VARCHAR(255),
  column_name VARCHAR(255),
  metadata JSON,
  api_key VARCHAR(255),
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (rule_id) REFERENCES data_quality_rules(rule_id) ON DELETE CASCADE ON UPDATE CASCADE,
  INDEX idx_rule_id (rule_id),
  INDEX idx_run_date (run_date)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('4.2.3 soda_checks (daxili icra)')
code('''CREATE TABLE soda_checks (
  id INT PRIMARY KEY AUTO_INCREMENT,
  check_name VARCHAR(255) NOT NULL,
  check_type VARCHAR(100) COMMENT 'row_count, null, unique, custom_sql',
  layer_name VARCHAR(255),
  table_name VARCHAR(255),
  column_name VARCHAR(255),
  kpi_name VARCHAR(255),
  custom_sql TEXT,
  target_value DECIMAL(15,4),
  threshold_value DECIMAL(15,4),
  status VARCHAR(50) DEFAULT 'active',
  schedule_cron VARCHAR(100) NULL,
  next_run_at DATETIME NULL,
  last_run_at DATETIME NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  INDEX idx_status (status),
  INDEX idx_next_run (next_run_at),
  INDEX idx_table (layer_name, table_name)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('4.2.4 soda_scan_results')
code('''CREATE TABLE soda_scan_results (
  id INT PRIMARY KEY AUTO_INCREMENT,
  scan_id VARCHAR(255),
  check_id INT,
  check_name VARCHAR(255),
  table_name VARCHAR(255),
  column_name VARCHAR(255),
  result VARCHAR(50) COMMENT 'pass, fail, error',
  outcome VARCHAR(100),
  metric_value DECIMAL(15,4),
  threshold DECIMAL(15,4),
  row_count INT,
  failed_row_count INT,
  metrics JSON,
  log TEXT,
  failed_rows JSON,
  duration_ms INT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (check_id) REFERENCES soda_checks(id) ON DELETE CASCADE,
  INDEX idx_scan_id (scan_id),
  INDEX idx_check_id (check_id),
  INDEX idx_created_at (created_at),
  INDEX idx_result (result)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('4.2.5 soda_check_results (aggregated summary)')
code('''CREATE TABLE soda_check_results (
  id INT PRIMARY KEY AUTO_INCREMENT,
  check_key VARCHAR(255),
  check_name VARCHAR(255),
  result VARCHAR(50),
  evaluation DECIMAL(5,2),
  metric_value DECIMAL(15,4),
  threshold DECIMAL(15,4),
  row_count INT,
  failed_row_count INT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_result (result),
  INDEX idx_check_key (check_key)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('4.2.6 dq_schedule_jobs (scheduler üçün — YENİ)')
code('''CREATE TABLE dq_schedule_jobs (
  id INT PRIMARY KEY AUTO_INCREMENT,
  check_id INT NOT NULL,
  status ENUM('pending','running','passed','failed','error') DEFAULT 'pending',
  scheduled_at DATETIME NOT NULL,
  started_at DATETIME NULL,
  finished_at DATETIME NULL,
  retries INT DEFAULT 0,
  error_message TEXT,
  scan_id VARCHAR(255) NULL,
  FOREIGN KEY (check_id) REFERENCES soda_checks(id) ON DELETE CASCADE,
  INDEX idx_status_scheduled (status, scheduled_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('4.3 ER Diagram')
code('''erDiagram
    data_quality_rules ||--o{ quality_rule_results : produces
    soda_checks ||--o{ soda_scan_results : executes
    soda_checks ||--o{ dq_schedule_jobs : scheduled_as
    soda_checks }o--|| catalog_tables : targets
    soda_checks }o--o| catalog_columns : targets
    api_keys }o--o{ quality_rule_results : submits_via''')

sec('4.4 Əsas SQL Sorğuları')

sub('4.4.1 Rule icrası (completeness nümunəsi)')
code('''-- soda_check-də saxlanan custom_sql
SELECT COUNT(*) AS total,
       SUM(CASE WHEN `:col` IS NOT NULL THEN 1 ELSE 0 END) AS passed,
       SUM(CASE WHEN `:col` IS NULL THEN 1 ELSE 0 END) AS failed
FROM `:schema`.`:table`;

-- pass_rate hesablanması
pass_rate = (passed / total) * 100''')

sub('4.4.2 Uniqueness')
code('''SELECT COUNT(*) AS total,
       COUNT(DISTINCT `:col`) AS distinct_cnt
FROM `:schema`.`:table`;
-- fail_rate = 1 - distinct_cnt/total''')

sub('4.4.3 Validity (regex / enum)')
code('''SELECT COUNT(*) AS total,
       SUM(CASE WHEN `:col` REGEXP '^[A-Z]{2}[0-9]{6}$' THEN 1 ELSE 0 END) AS passed
FROM `:schema`.`:table`;''')

sub('4.4.4 Latest result — hər rule üçün')
code('''SELECT r.rule_id, r.name, r.severity,
       res.pass_rate, res.run_date, res.failed_records
FROM data_quality_rules r
LEFT JOIN (
  SELECT rule_id, pass_rate, run_date, failed_records,
         ROW_NUMBER() OVER (PARTITION BY rule_id ORDER BY run_date DESC) rn
  FROM quality_rule_results
) res ON res.rule_id = r.rule_id AND res.rn = 1
WHERE r.status = 'active';''')

sub('4.4.5 Domain-səviyyə rollup')
code('''SELECT d.name AS domain,
       AVG(res.pass_rate) AS avg_pass_rate,
       COUNT(DISTINCT r.rule_id) AS rule_count
FROM domains d
JOIN catalog_tables t ON t.owner = d.name  -- domain-to-table map
JOIN data_quality_rules r ON r.table_name = t.name
JOIN quality_rule_results res ON res.rule_id = r.rule_id
WHERE res.run_date >= DATE_SUB(NOW(), INTERVAL 7 DAY)
GROUP BY d.name;''')

sub('4.4.6 API key authentication (external API)')
code('''SELECT user_id, is_active, expires_at
FROM api_keys
WHERE api_key = :key
  AND is_active = 1
  AND (expires_at IS NULL OR expires_at > NOW())
LIMIT 1;

-- istifadə vaxtını qeyd et
UPDATE api_keys SET last_used = NOW() WHERE api_key = :key;''')

sec('4.5 Business Logic Qaydaları')

para('Evaluation lifecycle:')
tbl(['Status','Təyinat'],[
    ['pending','Scheduler-in götürməyini gözləyir'],
    ['running','Scheduler icradadır'],
    ['passed','pass_rate >= threshold'],
    ['failed','pass_rate < threshold'],
    ['error','SQL icrası və ya connection xətası'],
])

para('Scoring formulları:')
code('''# Column quality %
column_score = AVG(pass_rate) for all rules on that column

# Table quality %
table_score = AVG(column_score) for all columns of that table

# Domain quality %
domain_score = AVG(table_score) for tables in domain

# Ağırlıqlı variant
weighted = SUM(pass_rate * severity_weight) / SUM(severity_weight)
# severity_weight: critical=4, high=3, medium=2, low=1''')

para('Nümunə hesablama:')
code('''Table: fact_sales
  customer_id: not-null 99.8% + unique 100% → 99.9%
  amount:      not-null 100%  + positive 97.5% → 98.75%
  order_date:  not-null 100%  + valid-range 99.2% → 99.6%

Table score = (99.9 + 98.75 + 99.6) / 3 = 99.42%''')

para('Edge case-lər:')
tbl(['Hal','Davranış'],[
    ['NULL column','completeness check 0% qaytarır, status=failed'],
    ['Connection fail','status=error, error_message saxlanılır, retry sayı artır'],
    ['Timeout (30s)','status=error, next_run-da retry ediləcək'],
    ['Partial run','scan_id fərqli olacaq, yeni scan kimi qeyd olunur'],
    ['0 row table','total=0 → pass_rate=100% (vacuous truth)'],
])

sec('4.6 REST API Endpoint-ləri')

sub('4.6.1 External API — POST api/quality-results')
para('Authentication: X-Api-Key və ya Authorization: Bearer <key>')
code('''curl -X POST "http://localhost/datarover/backend.php?action=api/quality-results" \\
  -H "X-Api-Key: $API_KEY" \\
  -H "Content-Type: application/json" \\
  -d '{
    "rule_id": "RULE-001",
    "pass_rate": 98.5,
    "total_records": 10000,
    "passed_records": 9850,
    "failed_records": 150,
    "run_date": "2026-04-19T10:30:00Z",
    "source_system": "great_expectations",
    "table_name": "fact_sales",
    "column_name": "customer_id"
  }' ''')

sub('4.6.2 CRUD quality_rules')
tbl(['Method','Path','İcazə'],[
    ['GET','quality_rules','quality.view'],
    ['POST','quality_rules','quality.create'],
    ['PUT','quality_rules&id={id}','quality.edit'],
    ['DELETE','quality_rules&id={id}','quality.delete'],
])

sub('4.6.3 soda_checks CRUD + icra')
tbl(['Method','Path','İcazə','İşlək'],[
    ['GET','soda_checks','quality.view','Check siyahısı'],
    ['POST','soda_checks','quality.create','Yeni check'],
    ['POST','run_check','quality.run','body: {check_id}'],
    ['POST','run_all_checks','quality.run','Bütün aktiv check-lər'],
    ['POST','run_selected_checks','quality.run','body: {check_ids:[...]}'],
    ['GET','soda_scans','quality.view','Scan history (limit)'],
    ['GET','soda_summary','quality.view','Aggregate metric-lər'],
    ['GET','download_failed_rows','quality.view','CSV/Excel export'],
])

sub('4.6.4 Schedule endpoint-ləri')
tbl(['Method','Path','Təyinat'],[
    ['GET','quality_schedules','Cari schedule-lar'],
    ['POST','quality_schedules','Yeni schedule (cron + check_id)'],
    ['PUT','quality_schedules&id={id}','Schedule yeniləmə'],
    ['DELETE','quality_schedules&id={id}','Ləğv'],
    ['POST','run_schedule','body: {schedule_id} — dərhal icra'],
    ['GET','quality_executions','Icra tarixçəsi'],
])

para('Frontend binding: app.js → runCheck(checkId) (3253), runAllChecks() (3396), saveCheck() (3057), fetchSodaScans() (1295).')

sec('4.7 Python Scheduler Contract (dq_schedule_simple.py)')

sub('4.7.1 Oxunan/yazılan cədvəllər')
tbl(['Cədvəl','Əməliyyat'],[
    ['dq_schedule_jobs','SELECT pending, UPDATE running/passed/failed/error'],
    ['soda_checks','SELECT check definition (custom_sql, table, column)'],
    ['soda_scan_results','INSERT nəticə'],
    ['soda_check_results','INSERT aggregated summary'],
    ['external_sources','SELECT bağlantı credentials'],
])

sub('4.7.2 Job pickup query')
code('''SELECT j.id, j.check_id, c.custom_sql, c.layer_name, c.table_name,
       c.column_name, c.threshold_value, c.target_value, c.check_type
FROM dq_schedule_jobs j
JOIN soda_checks c ON c.id = j.check_id
WHERE j.status = 'pending'
  AND j.scheduled_at <= NOW()
ORDER BY j.scheduled_at ASC
LIMIT 10
FOR UPDATE SKIP LOCKED;

-- status = 'running' et
UPDATE dq_schedule_jobs
SET status='running', started_at=NOW()
WHERE id IN (:ids);''')

sub('4.7.3 Status transition-ları')
tbl(['From','To','Şərt'],[
    ['pending','running','Scheduler tərəfindən pick-up'],
    ['running','passed','pass_rate >= threshold və SQL uğurla icra'],
    ['running','failed','pass_rate < threshold'],
    ['running','error','Exception (SQL xətası, timeout, conn fail)'],
    ['error','pending','Retry (retries < 3, next_run = NOW() + 5min)'],
])

sub('4.7.4 Nəticə qeyd etmə')
code('''# Python pseudo-code (dq_schedule_simple.py)
import mysql.connector, time, uuid

POLL_INTERVAL = 10  # saniyə
MAX_RETRIES = 3
TIMEOUT = 30

def poll():
    while True:
        jobs = pickup_jobs()
        for j in jobs:
            try:
                scan_id = str(uuid.uuid4())
                result = execute_check(j, timeout=TIMEOUT)
                insert_scan_result(scan_id, j, result)
                update_job(j['id'], 'passed' if result['pass'] else 'failed', scan_id)
            except Exception as e:
                if j['retries'] < MAX_RETRIES:
                    requeue(j['id'], delay_minutes=5)
                else:
                    update_job(j['id'], 'error', error=str(e))
        time.sleep(POLL_INTERVAL)

def execute_check(job, timeout):
    src = get_source_for_table(job['table_name'])
    conn = connect(src, timeout=timeout)
    cur = conn.cursor()
    cur.execute(job['custom_sql'])
    row = cur.fetchone()
    return evaluate(row, job)''')

sub('4.7.5 Failure & Retry policy')
tbl(['Halı','Reaksiya'],[
    ['SQL syntax error','status=error, retry etməz (sabit xəta)'],
    ['Connection refused','retry 3 dəfə, arada 5 dəq gözləyir'],
    ['Timeout','status=error, retries++, yenidən sıraya qoyulur'],
    ['Deadlock','10s sonra retry (MySQL spesifik)'],
])

sec('4.8 Permission Matrix')
tbl(['Action','SUPER_ADMIN','Q_MGR','ENGINEER','STEWARD','ANALYST','VIEWER'],[
    ['quality.view','✓','✓','✓','✓','✓','✓'],
    ['quality.create','✓','✓','✓','—','—','—'],
    ['quality.edit','✓','✓','✓','—','—','—'],
    ['quality.delete','✓','✓','—','—','—','—'],
    ['quality.run','✓','✓','✓','—','—','—'],
    ['quality.schedule','✓','✓','—','—','—','—'],
])

sec('4.9 Fayl Strukturu')
tbl(['Fayl','Sətir/Qeyd','Məzmun'],[
    ['backend.php','1118–1262','External api/quality-results'],
    ['backend.php','2742–3863','CRUD, run_check, download_failed_rows'],
    ['dq_schedule_simple.py','tam','Scheduler daemon (pip: mysql-connector, pyyaml)'],
    ['js/app.js','1295–3541','UI: scans, checks, runner'],
    ['docs/QUALITY_API.md','tam','Xarici API istifadə təlimatı'],
])

sec('4.10 İnteqrasiya Nöqtələri')
tbl(['Modul','Əlaqə'],[
    ['Catalog','soda_checks.table_name + column_name → catalog_columns'],
    ['External Sources','soda_check icrası external_sources.credentials-dən istifadə'],
    ['Glossary','data_quality_rules.rule_id ↔ glossary_terms.quality_rules JSON'],
    ['Reports','overall_quality_score hesablamasında iştirak'],
    ['Auth (API Key)','External tool-lar api_keys ilə autentifikasiya'],
])

sec('4.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE 6 cədvəl (data_quality_rules, quality_rule_results, soda_checks, soda_scan_results, soda_check_results, dq_schedule_jobs)'],
    ['2','pip install mysql-connector-python pyyaml'],
    ['3','dq_schedule_simple.py-ı Windows Task Scheduler və ya cron ilə arxada işlət'],
    ['4','İlk api_key yarat və xarici sistemə ver'],
    ['5','İlk soda_check yarat, POST run_check ilə test et'],
    ['6','Verification: SELECT check_id, COUNT(*) FROM soda_scan_results GROUP BY check_id;'],
])

section_break()

# ============================================================================
# MODULE 5: COLUMN MAPPING
# ============================================================================
module_header(5, 'Sütun Xəritələməsi', 'Column Mapping')

sec('5.1 Module Overview')
para('Mapping modulu ETL transformasiyalarını sənədləşdirir: source layer/table/column → target layer/table/column, transformation tipi və join şərtləri. Əsasən lineage (kimin haradan gəldiyi) və impact analysis üçün istifadə olunur. Rollar: Data Engineer (yaradır), Steward (təsdiq), Analyst (oxuyur).')

sec('5.2 Data Model')

sub('5.2.1 column_mappings')
code('''CREATE TABLE column_mappings (
  id INT PRIMARY KEY AUTO_INCREMENT,
  source_layer VARCHAR(255),
  source_table VARCHAR(255),
  source_column VARCHAR(255),
  target_layer VARCHAR(255),
  target_table VARCHAR(255),
  target_column VARCHAR(255),
  transformation VARCHAR(255) DEFAULT 'Direct Copy' COMMENT 'Direct Copy, Cast, Concat, Aggregate, Custom',
  custom_expression TEXT,
  description TEXT,
  is_validated BOOLEAN DEFAULT FALSE,
  validated_by VARCHAR(255),
  validated_at DATETIME NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  INDEX idx_source (source_layer, source_table, source_column),
  INDEX idx_target (target_layer, target_table, target_column),
  INDEX idx_validated (is_validated)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('5.2.2 table_join_conditions')
code('''CREATE TABLE table_join_conditions (
  id INT PRIMARY KEY AUTO_INCREMENT,
  source_layer VARCHAR(255),
  source_table VARCHAR(255),
  source_column VARCHAR(255),
  target_layer VARCHAR(255),
  target_table VARCHAR(255),
  target_column VARCHAR(255),
  join_type VARCHAR(20) DEFAULT 'INNER' COMMENT 'INNER, LEFT, RIGHT, FULL',
  additional_condition TEXT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_source (source_layer, source_table),
  INDEX idx_target (target_layer, target_table)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('5.3 ER Diagram')
code('''erDiagram
    catalog_tables ||--o{ column_mappings : source
    catalog_tables ||--o{ column_mappings : target
    catalog_tables ||--o{ table_join_conditions : joins''')

sec('5.4 Əsas SQL Sorğuları')

sub('5.4.1 Upstream lineage (source tapılması)')
code('''SELECT source_layer, source_table, source_column, transformation
FROM column_mappings
WHERE target_layer = :layer
  AND target_table = :table
  AND target_column = :column;''')

sub('5.4.2 Downstream impact (column dəyişsə kimə təsir edir)')
code('''SELECT target_layer, target_table, target_column, transformation
FROM column_mappings
WHERE source_layer = :layer
  AND source_table = :table
  AND source_column = :column;''')

sub('5.4.3 Recursive lineage chain (N hop)')
code('''WITH RECURSIVE lineage AS (
  SELECT 0 AS depth, source_layer, source_table, source_column,
         target_layer, target_table, target_column
  FROM column_mappings
  WHERE source_layer = :layer AND source_table = :table AND source_column = :col

  UNION ALL

  SELECT l.depth + 1, m.source_layer, m.source_table, m.source_column,
         m.target_layer, m.target_table, m.target_column
  FROM column_mappings m
  JOIN lineage l
    ON m.source_layer = l.target_layer
   AND m.source_table = l.target_table
   AND m.source_column = l.target_column
  WHERE l.depth < 10
)
SELECT * FROM lineage ORDER BY depth;''')

sub('5.4.4 Unvalidated mappings')
code('''SELECT COUNT(*) FROM column_mappings WHERE is_validated = 0;''')

sec('5.5 Business Logic Qaydaları')
tbl(['Qayda','Təsvir'],[
    ['Transformation enum','Direct Copy, Cast, Concat, Aggregate, Custom (SQL)'],
    ['Custom expression','Yalnız transformation=Custom olanda məcburidir'],
    ['Validation','Steward is_validated=1 qoyur, validated_by + validated_at doldurur'],
    ['Uniqueness','Eyni source+target column pair üçün 1 mapping (biznes qaydası, DB-də unikal deyil)'],
    ['Circular check','UI-da target→source dövrəsi aşkar edilsə xəbərdarlıq'],
])

sec('5.6 REST API')

sub('5.6.1 CRUD mappings')
tbl(['Method','Path','İcazə'],[
    ['GET','mappings','mapping.view — optional filter: source_table, target_table, is_validated'],
    ['GET','mappings&id={id}','mapping.view'],
    ['POST','mappings','mapping.create'],
    ['PUT','mappings&id={id}','mapping.edit'],
    ['DELETE','mappings&id={id}','mapping.delete'],
])
code('''POST body:
{
  "source_layer": "Staging",
  "source_table": "stg_orders",
  "source_column": "order_amt",
  "target_layer": "Mart",
  "target_table": "fact_sales",
  "target_column": "amount",
  "transformation": "Cast",
  "custom_expression": "CAST(order_amt AS DECIMAL(10,2))"
}''')

sub('5.6.2 CRUD join_conditions')
code('''POST body:
{
  "source_layer": "Mart",
  "source_table": "fact_sales",
  "source_column": "customer_id",
  "target_layer": "Mart",
  "target_table": "dim_customer",
  "target_column": "id",
  "join_type": "LEFT"
}''')

para('cURL:')
code('curl "http://localhost/datarover/backend.php?action=mappings" -H "Authorization: Bearer $TOKEN"')
para('Frontend: app.js mapping UI-dakı saveMapping() və loadMappings() funksiyaları.')

sec('5.7 Python Scheduler Contract')
para('Mapping modulu scheduler istifadə etmir — tam manual CRUD. Auto-discovery üçün gələcəkdə SQL parser əlavə oluna bilər (ETL CREATE TABLE AS SELECT skriptlərini parse edib mapping-ə INSERT).')

sec('5.8 Permission Matrix')
tbl(['Action','ADMIN','ENGINEER','STEWARD','ANALYST','VIEWER'],[
    ['mapping.view','✓','✓','✓','✓','✓'],
    ['mapping.create','✓','✓','—','—','—'],
    ['mapping.edit','✓','✓','—','—','—'],
    ['mapping.validate','✓','—','✓','—','—'],
    ['mapping.delete','✓','✓','—','—','—'],
])

sec('5.9 Fayl Strukturu')
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','4243–4363','mappings + join_conditions CRUD'],
    ['js/app.js','Mapping tab','saveMapping, loadMappings, validateMapping'],
])

sec('5.10 İnteqrasiya Nöqtələri')
para('Catalog (source/target table refs), Lineage view (UI-da D3.js / vis-network ilə graph). Reports modulu data source təhlili üçün bu cədvəlləri oxuya bilər.')

sec('5.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE column_mappings, table_join_conditions'],
    ['2','Seed: 10–15 örnək mapping'],
    ['3','Lineage query test: 5.4.3 recursive CTE'],
    ['4','Verification: SELECT COUNT(*) FROM column_mappings WHERE is_validated=1;'],
])

section_break()

# ============================================================================
# MODULE 6: REPORTING
# ============================================================================
module_header(6, 'Hesabatlar', 'Reporting')

sec('6.1 Module Overview')
para('Reports modulu BI alətlərindəki (Power BI, Tableau, Looker) dashboard-ları və onların metadata-sını (sahibi, data mənbələri, KPI-lər, istifadəçilər, glossary term-ləri) saxlayır. Certification workflow: draft → under_review → certified → deprecated. overall_quality_score data quality modulundan hesablanır. Rollar: Analyst (yaradır), Data Owner (sertifikatlaşdırır), Viewer (istifadə edir).')

sec('6.2 Data Model')

sub('6.2.1 reports')
code('''CREATE TABLE reports (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id VARCHAR(50) UNIQUE,
  report_name VARCHAR(255) NOT NULL,
  short_description VARCHAR(500),
  long_description TEXT,
  business_purpose TEXT,
  domain_id INT,
  sub_domain VARCHAR(255),
  report_type VARCHAR(100) DEFAULT 'analytical' COMMENT 'analytical, operational, regulatory',
  business_owner VARCHAR(255),
  data_owner VARCHAR(255),
  technical_owner VARCHAR(255),
  data_steward VARCHAR(255),
  report_maintainer VARCHAR(255),
  update_frequency VARCHAR(100) DEFAULT 'daily',
  certification_status VARCHAR(100) DEFAULT 'draft' COMMENT 'draft, under_review, certified, deprecated',
  bi_tool VARCHAR(100) COMMENT 'power_bi, tableau, looker, qlik, metabase',
  workspace_location VARCHAR(500),
  dashboard_url VARCHAR(500),
  report_version VARCHAR(20) DEFAULT '1.0',
  pages_count INT DEFAULT 1,
  visuals_count INT,
  target_audience VARCHAR(500),
  access_level VARCHAR(50) DEFAULT 'restricted' COMMENT 'public, restricted, confidential',
  has_pii BOOLEAN DEFAULT FALSE,
  has_financial_data BOOLEAN DEFAULT FALSE,
  regulatory_flags JSON,
  retention_policy VARCHAR(255),
  overall_quality_score DECIMAL(5,2),
  last_refreshed_at DATETIME NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  FOREIGN KEY (domain_id) REFERENCES domains(id) ON DELETE SET NULL,
  INDEX idx_report_id (report_id),
  INDEX idx_status (certification_status),
  INDEX idx_domain (domain_id),
  INDEX idx_bi_tool (bi_tool)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('6.2.2 report_data_sources')
code('''CREATE TABLE report_data_sources (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  source_type VARCHAR(50) DEFAULT 'table',
  layer_name VARCHAR(255),
  table_id INT NULL,
  table_name VARCHAR(255),
  schema_name VARCHAR(255),
  database_name VARCHAR(255),
  used_columns JSON,
  etl_process_name VARCHAR(255),
  transformation_logic TEXT,
  refresh_frequency VARCHAR(100),
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
  FOREIGN KEY (table_id) REFERENCES catalog_tables(id) ON DELETE SET NULL,
  INDEX idx_report (report_id),
  INDEX idx_table (table_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('6.2.3 report_kpis')
code('''CREATE TABLE report_kpis (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  kpi_name VARCHAR(255),
  kpi_code VARCHAR(100),
  business_definition TEXT,
  technical_formula TEXT,
  aggregation_rule VARCHAR(255),
  unit_of_measure VARCHAR(100),
  target_value DECIMAL(15,2),
  threshold_warning DECIMAL(15,2),
  threshold_critical DECIMAL(15,2),
  current_value DECIMAL(15,2),
  trend VARCHAR(50) DEFAULT 'stable' COMMENT 'up, down, stable',
  filters_applied JSON,
  dependencies JSON,
  term_id INT NULL,
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE SET NULL,
  INDEX idx_report (report_id),
  INDEX idx_term (term_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('6.2.4 report_users, report_terms')
code('''CREATE TABLE report_users (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  user_name VARCHAR(255),
  user_email VARCHAR(255),
  user_role VARCHAR(100) DEFAULT 'viewer' COMMENT 'viewer, editor, owner',
  department VARCHAR(255),
  access_granted_by VARCHAR(255),
  access_granted_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
  INDEX idx_report (report_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE report_terms (
  id INT PRIMARY KEY AUTO_INCREMENT,
  report_id INT NOT NULL,
  term_id INT NOT NULL,
  relationship_type VARCHAR(100) DEFAULT 'references',
  notes TEXT,
  FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
  FOREIGN KEY (term_id) REFERENCES glossary_terms(id) ON DELETE CASCADE,
  UNIQUE KEY unique_report_term (report_id, term_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('6.3 ER Diagram')
code('''erDiagram
    domains ||--o{ reports : categorizes
    reports ||--o{ report_data_sources : uses
    reports ||--o{ report_kpis : contains
    reports ||--o{ report_users : shared_with
    reports ||--o{ report_terms : references
    catalog_tables ||--o{ report_data_sources : supplies
    glossary_terms ||--o{ report_terms : linked
    glossary_terms ||--o{ report_kpis : defined_as''')

sec('6.4 Əsas SQL Sorğuları')

sub('6.4.1 Filter və search')
code('''SELECT r.*, d.name AS domain_name
FROM reports r
LEFT JOIN domains d ON d.id = r.domain_id
WHERE (:q IS NULL OR r.report_name LIKE CONCAT('%',:q,'%') OR r.short_description LIKE CONCAT('%',:q,'%'))
  AND (:domain IS NULL OR r.domain_id = :domain)
  AND (:bi_tool IS NULL OR r.bi_tool = :bi_tool)
  AND (:status IS NULL OR r.certification_status = :status)
ORDER BY r.updated_at DESC;''')

sub('6.4.2 Report detal + sub-resource')
code('''SELECT r.*, d.name AS domain_name,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT('id',id,'table_name',table_name,'used_columns',used_columns))
   FROM report_data_sources WHERE report_id = r.id) AS data_sources,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT('id',id,'name',kpi_name,'formula',technical_formula))
   FROM report_kpis WHERE report_id = r.id) AS kpis,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT('email',user_email,'role',user_role))
   FROM report_users WHERE report_id = r.id) AS users,
  (SELECT JSON_ARRAYAGG(t.name)
   FROM report_terms rt JOIN glossary_terms t ON t.id = rt.term_id
   WHERE rt.report_id = r.id) AS terms
FROM reports r
LEFT JOIN domains d ON d.id = r.domain_id
WHERE r.id = :id;''')

sub('6.4.3 Overall quality score hesablanması')
code('''-- Bu report-un data source table-larının keyfiyyət skoru
SELECT r.id, r.report_name,
       AVG(res.pass_rate) AS overall_quality_score
FROM reports r
JOIN report_data_sources rds ON rds.report_id = r.id
JOIN data_quality_rules dqr ON dqr.table_name = rds.table_name
JOIN quality_rule_results res ON res.rule_id = dqr.rule_id
WHERE res.run_date >= DATE_SUB(NOW(), INTERVAL 7 DAY)
GROUP BY r.id;

-- Update
UPDATE reports SET overall_quality_score = :score WHERE id = :id;''')

sec('6.5 Business Logic Qaydaları')
tbl(['Qayda','Təsvir'],[
    ['certification_status lifecycle','draft → under_review → certified → deprecated'],
    ['has_pii','TRUE olanda access_level avtomatik "restricted" qoyulur'],
    ['regulatory_flags JSON','["GDPR","SOX","HIPAA"] — audit report-ları üçün'],
    ['overall_quality_score','Gündəlik cron ilə 6.4.3 SQL-dən refresh olunur'],
    ['update_frequency','daily, hourly, weekly, monthly, realtime, manual'],
])

sec('6.6 REST API')

sub('6.6.1 CRUD reports')
tbl(['Method','Path','İcazə'],[
    ['GET','reports','reports.view (filters: search, domain, bi_tool, status)'],
    ['POST','reports','reports.create'],
    ['PUT','reports&id={id}','reports.edit'],
    ['DELETE','reports&id={id}','reports.delete'],
])

sub('6.6.2 Sub-resource endpoint-ləri')
tbl(['Method','Path','Təyinat'],[
    ['GET/POST','report_data_sources?report_id={id}','Data source əlavə etmək'],
    ['GET/POST/PUT','report_kpis?report_id={id}','KPI idarəsi'],
    ['GET/POST','report_users?report_id={id}','Access list'],
    ['GET/POST','report_terms?report_id={id}','Glossary term linki'],
    ['GET','search_reports?q={term}','Qısa axtarış'],
])

code('''POST report nümunəsi:
{
  "report_name": "Sales Executive Dashboard",
  "domain_id": 3,
  "business_owner": "cso@corp.com",
  "bi_tool": "power_bi",
  "report_type": "analytical",
  "update_frequency": "daily",
  "access_level": "restricted",
  "has_pii": false,
  "dashboard_url": "https://app.powerbi.com/..."
}''')

sec('6.7 Python Scheduler Contract')
para('overall_quality_score gündəlik yenilənməsi üçün scheduler tapşırığı (dq_schedule_simple.py-da və ya ayrı script):')
code('''-- Gündəlik cron task-ı (01:00 UTC+4)
UPDATE reports r
JOIN (
  SELECT rds.report_id, AVG(res.pass_rate) AS avg_score
  FROM report_data_sources rds
  JOIN data_quality_rules dqr ON dqr.table_name = rds.table_name
  JOIN quality_rule_results res ON res.rule_id = dqr.rule_id
  WHERE res.run_date >= DATE_SUB(NOW(), INTERVAL 7 DAY)
  GROUP BY rds.report_id
) agg ON agg.report_id = r.id
SET r.overall_quality_score = agg.avg_score;''')

sec('6.8 Permission Matrix')
tbl(['Action','ADMIN','ANALYST','STEWARD','ENGINEER','VIEWER'],[
    ['reports.view','✓','✓','✓','✓','✓'],
    ['reports.create','✓','✓','—','—','—'],
    ['reports.edit','✓','✓','✓','—','—'],
    ['reports.certify','✓','—','✓','—','—'],
    ['reports.delete','✓','—','✓','—','—'],
])

sec('6.9 Fayl Strukturu')
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','4369–4688','reports + sub-resource handler'],
    ['js/app.js','Reports tab','loadReports, saveReport, linkKpi'],
])

sec('6.10 İnteqrasiya Nöqtələri')
para('Domain (domain_id FK), Catalog (report_data_sources.table_id FK), Glossary (report_terms junction, report_kpis.term_id), Quality (overall_quality_score hesablama).')

sec('6.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE 5 cədvəl (reports + 4 sub-resource)'],
    ['2','Domain-lər mövcud olduğundan əmin ol'],
    ['3','Seed 3–5 nümunə report'],
    ['4','Cron quality score update-i qur'],
    ['5','SELECT certification_status, COUNT(*) FROM reports GROUP BY 1;'],
])

section_break()

# ============================================================================
# MODULE 7: DOMAIN MANAGEMENT
# ============================================================================
module_header(7, 'Domen İdarəetməsi', 'Domain Management')

sec('7.1 Module Overview')
para('Domain modulu biznes sahələrini (Sales, Finance, HR, Operations, Marketing) saxlayır və hər domain-ə məsul şəxsləri (stakeholder) təyin edir. Glossary term-ləri və Reports domen-ə bağlanır, belə ki data governance məsuliyyəti bölüşdürülə bilər. Rollar: Gov Manager (domain yaradır), Steward (stakeholder kimi təyin olunur).')

sec('7.2 Data Model')

sub('7.2.1 domains')
code('''CREATE TABLE domains (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255) UNIQUE NOT NULL,
  icon VARCHAR(50),
  color VARCHAR(10) COMMENT 'HEX rəng, məs: #1a56db',
  description TEXT,
  parent_domain_id INT NULL COMMENT 'Nested domain üçün',
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  FOREIGN KEY (parent_domain_id) REFERENCES domains(id) ON DELETE SET NULL,
  INDEX idx_name (name),
  INDEX idx_parent (parent_domain_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('7.2.2 domain_stakeholders')
code('''CREATE TABLE domain_stakeholders (
  id INT PRIMARY KEY AUTO_INCREMENT,
  domain_id INT NOT NULL,
  role_id VARCHAR(50) COMMENT 'OWNER, STEWARD, APPROVER, VIEWER',
  stakeholder_name VARCHAR(255),
  stakeholder_email VARCHAR(255),
  user_id INT NULL,
  assigned_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (domain_id) REFERENCES domains(id) ON DELETE CASCADE,
  FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE SET NULL,
  UNIQUE KEY unique_domain_email_role (domain_id, stakeholder_email, role_id),
  INDEX idx_domain (domain_id),
  INDEX idx_user (user_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('7.3 ER Diagram')
code('''erDiagram
    domains ||--o{ domain_stakeholders : has
    domains ||--o{ domains : parent_child
    domains ||--o{ reports : categorizes
    users ||--o{ domain_stakeholders : assigned
    glossary_terms }o--|| domains : belongs_to''')

sec('7.4 Əsas SQL')

sub('7.4.1 Domain + stakeholder siyahısı')
code('''SELECT d.*,
  (SELECT JSON_ARRAYAGG(JSON_OBJECT(
      'id',ds.id,'role',ds.role_id,'name',ds.stakeholder_name,'email',ds.stakeholder_email))
   FROM domain_stakeholders ds WHERE ds.domain_id = d.id) AS stakeholders
FROM domains d
ORDER BY d.name;''')

sub('7.4.2 Domain-based statistics')
code('''SELECT d.name AS domain,
       (SELECT COUNT(*) FROM glossary_terms WHERE domain = d.name) AS term_count,
       (SELECT COUNT(*) FROM reports WHERE domain_id = d.id) AS report_count
FROM domains d;''')

sec('7.5 Business Logic')
tbl(['Qayda','Təsvir'],[
    ['Unique name','name UNIQUE constraint'],
    ['Hierarchy','parent_domain_id self-ref, max depth 3 UI-da'],
    ['Role-ID enum','OWNER, STEWARD, APPROVER, VIEWER (hər domain-də 1 OWNER olmalıdır)'],
    ['Delete protection','Əgər domain-ə bağlı reports varsa silmək qadağan (ON DELETE SET NULL ilə orphan olur)'],
])

sec('7.6 REST API')
tbl(['Method','Path','Təyinat'],[
    ['GET','domains','Bütün domains (id varsa stakeholders ilə)'],
    ['POST','domains','body: name, icon, color, description'],
    ['PUT','domains&id={id}','Yeniləmə'],
    ['DELETE','domains&id={id}','Silmə (reports orphan olur)'],
    ['GET','domain_stakeholders?domain_id={id}','Siyahı'],
    ['POST','domain_stakeholders','body: domain_id, role_id, name, email'],
    ['DELETE','domain_stakeholders&id={id}','Silmə'],
])

sec('7.7 Python Scheduler')
para('İstifadə etmir.')

sec('7.8 Permission Matrix')
tbl(['Action','ADMIN','GOV_MGR','STEWARD','ANALYST','VIEWER'],[
    ['domains.view','✓','✓','✓','✓','✓'],
    ['domains.create','✓','✓','—','—','—'],
    ['domains.edit','✓','✓','—','—','—'],
    ['domains.delete','✓','✓','—','—','—'],
    ['domains.assign_stakeholder','✓','✓','—','—','—'],
])

sec('7.9 Fayl Strukturu')
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','5000–5140','domains + domain_stakeholders'],
    ['js/app.js','Domains tab','loadDomains, saveDomain'],
])

sec('7.10 İnteqrasiya')
para('Glossary (domain VARCHAR sahə domains.name-ə match), Reports (domain_id FK), RBAC (user_id ilə domain_stakeholders.user_id link).')

sec('7.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE domains, domain_stakeholders'],
    ['2','Seed: Sales, Finance, HR, Operations, Marketing'],
    ['3','Hər domain-ə OWNER təyin et'],
    ['4','Verification: SELECT name, (SELECT COUNT(*) FROM domain_stakeholders WHERE domain_id=d.id) FROM domains d;'],
])

section_break()

# ============================================================================
# MODULE 8: GOVERNANCE WORKFLOW
# ============================================================================
module_header(8, 'Governance İş Axını', 'Governance Workflow')

sec('8.1 Module Overview')
para('Governance modulu yeni glossary term, report, catalog entry kimi artefaktların təsdiq iş axınını idarə edir. Workflow step-ləri (draft→review→approve→publish) və hər step üçün təsdiqləyiciləri (governance_stakeholders) saxlayır. Bu modul Business Glossary status keçidinin arxasındadır. Rollar: Gov Manager (workflow dizayn), Stakeholder (təsdiq verir).')

sec('8.2 Data Model')

sub('8.2.1 governance_roles')
code('''CREATE TABLE governance_roles (
  id INT PRIMARY KEY AUTO_INCREMENT,
  role_id INT NULL COMMENT 'user_roles cədvəlinə opsional ref',
  name VARCHAR(255) UNIQUE NOT NULL,
  icon VARCHAR(50),
  color VARCHAR(10),
  description TEXT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (role_id) REFERENCES user_roles(id) ON DELETE SET NULL
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('8.2.2 governance_workflow_steps')
code('''CREATE TABLE governance_workflow_steps (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255),
  step_order INT,
  description TEXT,
  from_status VARCHAR(100),
  to_status VARCHAR(100),
  required_approvers INT DEFAULT 1,
  approver_role_id INT NULL,
  timeout_days INT DEFAULT 7,
  auto_escalate BOOLEAN DEFAULT FALSE,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (approver_role_id) REFERENCES governance_roles(id) ON DELETE SET NULL,
  INDEX idx_step_order (step_order),
  INDEX idx_transition (from_status, to_status)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('8.2.3 governance_stakeholders')
code('''CREATE TABLE governance_stakeholders (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255),
  role_id INT,
  email VARCHAR(255),
  department VARCHAR(255),
  user_id INT NULL,
  is_active BOOLEAN DEFAULT TRUE,
  FOREIGN KEY (role_id) REFERENCES governance_roles(id) ON DELETE CASCADE,
  FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE SET NULL,
  INDEX idx_role (role_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('8.3 ER Diagram')
code('''erDiagram
    governance_roles ||--o{ governance_workflow_steps : approves
    governance_roles ||--o{ governance_stakeholders : assigned
    user_roles ||--o| governance_roles : maps_to
    users ||--o{ governance_stakeholders : linked''')

sec('8.4 Əsas SQL')

sub('8.4.1 Term üçün növbəti step')
code('''SELECT id, name, to_status, required_approvers, timeout_days
FROM governance_workflow_steps
WHERE from_status = :current_status
ORDER BY step_order ASC
LIMIT 1;''')

sub('8.4.2 Step üçün təsdiqləyicilər')
code('''SELECT s.name, s.email
FROM governance_stakeholders s
JOIN governance_workflow_steps w ON w.approver_role_id = s.role_id
WHERE w.id = :step_id AND s.is_active = 1;''')

sub('8.4.3 Overdue term-lər')
code('''SELECT t.id, t.name, t.status, t.updated_at, w.timeout_days
FROM glossary_terms t
JOIN governance_workflow_steps w ON w.from_status = t.status
WHERE t.deleted_at IS NULL
  AND DATEDIFF(NOW(), t.updated_at) > w.timeout_days;''')

sec('8.5 Business Logic')
tbl(['Qayda','Təsvir'],[
    ['Step sequence','step_order ASC sırası ilə'],
    ['required_approvers','Neçə stakeholder təsdiq etməlidir — 1, 2 və ya 3'],
    ['timeout_days','Step bu günlərdə başa çatmasa auto_escalate'],
    ['Escalation','auto_escalate=TRUE → parent governance_role-a bildirir'],
    ['user_id link','Auth moduldakı users ilə sync'],
])

sec('8.6 REST API')
tbl(['Method','Path','Təyinat'],[
    ['GET','governance_roles','Rol siyahısı'],
    ['POST','governance_roles','Yeni rol'],
    ['GET','governance_steps','Workflow step-ləri'],
    ['POST','governance_steps','Step əlavə et'],
    ['PUT','governance_steps&id={id}','Yeniləmə'],
    ['GET','governance_stakeholders','Bütün stakeholders'],
    ['POST','governance_stakeholders','Əlavə et'],
])

code('''POST governance_steps nümunəsi:
{
  "name": "Data Owner Approval",
  "step_order": 2,
  "from_status": "under_review",
  "to_status": "approved",
  "required_approvers": 1,
  "approver_role_id": 3,
  "timeout_days": 5
}''')

sec('8.7 Python Scheduler')
para('Overdue/escalation üçün scheduler task (gündəlik 09:00):')
code('''-- Timeout-a çatmış term-lər
SELECT t.id, t.name, w.to_status, GROUP_CONCAT(s.email) AS escalate_to
FROM glossary_terms t
JOIN governance_workflow_steps w ON w.from_status = t.status
JOIN governance_stakeholders s ON s.role_id = w.approver_role_id
WHERE t.deleted_at IS NULL
  AND DATEDIFF(NOW(), t.updated_at) > w.timeout_days
  AND w.auto_escalate = 1
GROUP BY t.id, w.id;
-- Hər sətir üçün email notification göndər''')

sec('8.8 Permission Matrix')
tbl(['Action','ADMIN','GOV_MGR','STEWARD','ANALYST','VIEWER'],[
    ['governance.view','✓','✓','✓','✓','✓'],
    ['governance.create_role','✓','✓','—','—','—'],
    ['governance.edit_workflow','✓','✓','—','—','—'],
    ['governance.approve','✓','✓','✓','—','—'],
])

sec('8.9 Fayl Strukturu')
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','4690–4846','3 sub-handler'],
    ['js/app.js','Governance tab','workflow UI'],
])

sec('8.10 İnteqrasiya')
para('Glossary status transitions bu moduldan işarələri oxuyur (from_status/to_status). Auth modulu ilə user_id link. Email notification üçün SMTP config tələb olunur (env: SMTP_HOST, SMTP_PORT, SMTP_USER).')

sec('8.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE governance_roles, governance_workflow_steps, governance_stakeholders'],
    ['2','Seed roles: Data Owner, Data Steward, Compliance'],
    ['3','Seed steps: draft→under_review (steward), under_review→approved (owner)'],
    ['4','Stakeholder-ləri user-ə bağla'],
    ['5','Email SMTP konfiqurasiyası'],
])

section_break()

# ============================================================================
# MODULE 9: EXTERNAL DATA SOURCES
# ============================================================================
module_header(9, 'Xarici Data Mənbələri', 'External Data Sources')

sec('9.1 Module Overview')
para('External Sources modulu 4 növ data bazasına (MySQL, PostgreSQL, Oracle, MSSQL) credential-larla birlikdə bağlantı saxlayır. Bu bağlantılar Catalog idxalı, live profilləşdirmə və Soda check icrası üçün istifadə edilir. Credentials DB-də saxlanılır (gələcəkdə encryption tələb olunur). Rollar: Data Engineer (yaradır), Admin (silər).')

sec('9.2 Data Model')

sub('9.2.1 external_sources')
code('''CREATE TABLE external_sources (
  id INT PRIMARY KEY AUTO_INCREMENT,
  name VARCHAR(255) UNIQUE NOT NULL,
  db_type VARCHAR(50) NOT NULL COMMENT 'mysql, postgresql, oracle, mssql',
  host VARCHAR(255) NOT NULL,
  port INT NOT NULL DEFAULT 3306,
  username VARCHAR(255),
  password VARCHAR(500) COMMENT 'Encrypted (AES-256-CBC recommended)',
  database_name VARCHAR(255),
  sid VARCHAR(255) COMMENT 'Oracle SID',
  service_name VARCHAR(255) COMMENT 'Oracle service alternative',
  ssl_mode VARCHAR(50) DEFAULT 'prefer',
  status VARCHAR(50) DEFAULT 'active' COMMENT 'active, inactive, error',
  connection_string TEXT,
  table_count INT DEFAULT 0,
  last_scan DATETIME NULL,
  last_error TEXT,
  created_by INT NULL,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at DATETIME NULL ON UPDATE CURRENT_TIMESTAMP,
  FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE SET NULL,
  INDEX idx_db_type (db_type),
  INDEX idx_status (status)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sub('9.2.2 external_source_tables')
code('''CREATE TABLE external_source_tables (
  id INT PRIMARY KEY AUTO_INCREMENT,
  source_id INT NOT NULL,
  schema_name VARCHAR(255),
  table_name VARCHAR(255) NOT NULL,
  object_type VARCHAR(50) DEFAULT 'TABLE' COMMENT 'TABLE, VIEW, MAT_VIEW',
  row_count BIGINT,
  column_count INT,
  imported_to_catalog BOOLEAN DEFAULT FALSE,
  catalog_table_id INT NULL,
  last_scanned_at DATETIME NULL,
  FOREIGN KEY (source_id) REFERENCES external_sources(id) ON DELETE CASCADE,
  FOREIGN KEY (catalog_table_id) REFERENCES catalog_tables(id) ON DELETE SET NULL,
  UNIQUE KEY unique_source_schema_table (source_id, schema_name, table_name),
  INDEX idx_source (source_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

sec('9.3 ER Diagram')
code('''erDiagram
    users ||--o{ external_sources : created_by
    external_sources ||--o{ external_source_tables : contains
    external_sources ||--o{ catalog_tables : sources
    external_source_tables }o--o| catalog_tables : imported_as''')

sec('9.4 Əsas SQL')

sub('9.4.1 Source siyahısı')
code('''SELECT id, name, db_type, host, status, table_count, last_scan
FROM external_sources
WHERE status != 'inactive'
ORDER BY name;''')

sub('9.4.2 Source-un table-ları')
code('''SELECT schema_name, table_name, row_count, imported_to_catalog
FROM external_source_tables
WHERE source_id = :id
ORDER BY schema_name, table_name;''')

sub('9.4.3 Table scan (MySQL)')
code('''-- information_schema üzərindən
SELECT table_schema, table_name, table_rows, table_type
FROM information_schema.tables
WHERE table_schema NOT IN ('information_schema','mysql','performance_schema','sys')
ORDER BY table_schema, table_name;''')

para('PostgreSQL üçün:')
code('''SELECT table_schema, table_name,
       (SELECT reltuples::BIGINT FROM pg_class WHERE relname = t.table_name) AS row_count
FROM information_schema.tables t
WHERE table_schema NOT IN ('pg_catalog','information_schema');''')

para('Oracle üçün:')
code('''SELECT owner AS schema_name, table_name, num_rows
FROM all_tables
WHERE owner NOT IN ('SYS','SYSTEM','ORDSYS','MDSYS','CTXSYS');''')

sub('9.4.4 Column scan (MySQL)')
code('''SELECT column_name, data_type, is_nullable, column_default,
       column_key, extra, ordinal_position
FROM information_schema.columns
WHERE table_schema = :schema AND table_name = :table
ORDER BY ordinal_position;''')

sec('9.5 Business Logic')
tbl(['Qayda','Təsvir'],[
    ['Password encryption','password sahəsi AES-256-CBC ilə şifrələnməli (key env-dən)'],
    ['Connection timeout','10 saniyə test üçün, 30 saniyə scan üçün'],
    ['Port defaults','MySQL 3306, PostgreSQL 5432, Oracle 1521, MSSQL 1433'],
    ['System schema filter','information_schema, pg_catalog, sys, mysql filter edilir'],
    ['Import duplicate','Eyni layer+table adı varsa UPDATE edir (ON DUPLICATE KEY)'],
    ['last_scan update','Hər successful scan-dan sonra timestamp yenilənir'],
])

para('PHP-də connection kod nümunəsi:')
code('''function connectExternal($source) {
  $pass = decrypt($source['password']);
  switch ($source['db_type']) {
    case 'mysql':
      return new PDO("mysql:host={$source['host']};port={$source['port']};dbname={$source['database_name']}",
                     $source['username'], $pass, [PDO::ATTR_TIMEOUT => 30]);
    case 'postgresql':
      return new PDO("pgsql:host={$source['host']};port={$source['port']};dbname={$source['database_name']}",
                     $source['username'], $pass);
    case 'oracle':
      $dsn = "oci:dbname=//{$source['host']}:{$source['port']}/{$source['sid']}";
      return new PDO($dsn, $source['username'], $pass);
    case 'mssql':
      return new PDO("sqlsrv:Server={$source['host']},{$source['port']};Database={$source['database_name']}",
                     $source['username'], $pass);
  }
}''')

sec('9.6 REST API')
tbl(['Method','Path','Təyinat'],[
    ['GET','external_sources','Bütün source-lar'],
    ['GET','external_sources&id={id}','Source detal + table list'],
    ['POST','external_sources','Yeni source (+ connection test)'],
    ['PUT','external_sources&id={id}','Yeniləmə'],
    ['DELETE','external_sources&id={id}','Silmə (cascade tables)'],
    ['POST','external_sources_import','Seçilmiş table-ları catalog-a idxal et'],
    ['GET','get_connections','Schedule yaratmaq üçün connection siyahısı'],
])

code('''POST external_sources nümunəsi:
{
  "name": "prod-dwh",
  "db_type": "postgresql",
  "host": "db.prod.internal",
  "port": 5432,
  "username": "readonly_user",
  "password": "secret",
  "database_name": "warehouse",
  "ssl_mode": "require"
}''')

code('''POST external_sources_import:
{
  "source_id": 3,
  "tables": [
    {
      "schema_name": "public",
      "table_name": "orders",
      "columns": [
        {"name":"id","data_type":"INT","is_pk":true},
        {"name":"customer_id","data_type":"INT","is_fk":true},
        {"name":"amount","data_type":"DECIMAL(10,2)"}
      ],
      "row_count": 150000
    }
  ],
  "layer_id": 1
}''')

sec('9.7 Python Scheduler')
para('Periodik scan üçün (həftəlik / gündəlik):')
code('''-- source-ları scan ardıcıllığı ilə sırala
SELECT id, name, db_type, host, port, username, password, database_name, sid
FROM external_sources
WHERE status = 'active'
  AND (last_scan IS NULL OR last_scan < DATE_SUB(NOW(), INTERVAL 1 DAY))
ORDER BY last_scan ASC NULLS FIRST
LIMIT 5;

-- scan-dan sonra
UPDATE external_sources
SET last_scan = NOW(), table_count = :cnt, status = 'active', last_error = NULL
WHERE id = :id;

-- error halında
UPDATE external_sources
SET status = 'error', last_error = :err
WHERE id = :id;''')

sec('9.8 Permission Matrix')
tbl(['Action','ADMIN','ENGINEER','STEWARD','ANALYST','VIEWER'],[
    ['sources.view','✓','✓','✓','—','—'],
    ['sources.create','✓','✓','—','—','—'],
    ['sources.edit','✓','✓','—','—','—'],
    ['sources.delete','✓','—','—','—','—'],
    ['sources.import','✓','✓','—','—','—'],
    ['sources.scan','✓','✓','—','—','—'],
])

sec('9.9 Fayl Strukturu')
tbl(['Fayl','Sətir','Məzmun'],[
    ['backend.php','5146–5328','external_sources CRUD + import'],
    ['backend.php','2713–2726','get_connections'],
    ['js/app.js','External tab','loadSources, testConnection, importTables'],
])

sec('9.10 İnteqrasiya')
para('Catalog (catalog_tables.source_id FK), Data Quality (soda_check icrası bu connection-ları istifadə edir), Auth (created_by user ref). Credentials saxlanması üçün XAMPP-da APP_ENCRYPTION_KEY environment variable tələb olunur.')

sec('9.11 Setup Checklist')
tbl(['Addım','Komanda'],[
    ['1','CREATE TABLE external_sources, external_source_tables'],
    ['2','APP_ENCRYPTION_KEY env variable qoy'],
    ['3','PDO driverları yoxla: mysql, pgsql, oci, sqlsrv'],
    ['4','İlk source əlavə et və ldap/test-ə bənzər test endpoint-i ilə yoxla'],
    ['5','external_sources_import ilə 1 table idxal et'],
    ['6','SELECT * FROM catalog_tables WHERE source_id IS NOT NULL;'],
])

section_break()

# ============================================================================
# APPENDIX: GLOBAL SETUP CHECKLIST
# ============================================================================
doc.add_heading('Əlavə: Ümumi Setup Checklist', level=1)

para('Yeni mühit üçün ardıcıllıq:')
tbl(['#','Addım','Detal'],[
    ['1','XAMPP quraşdır','PHP 8.x, MySQL 8.x, Apache'],
    ['2','DB yarat','CREATE DATABASE datarover DEFAULT CHARSET utf8mb4 COLLATE utf8mb4_unicode_ci;'],
    ['3','Timezone tənzimlə','SET GLOBAL time_zone = "+04:00";'],
    ['4','Schema migrasiyaları','1-ci növbədə auth (users, roles, permissions)'],
    ['5','Sonra domain + governance','Referens cədvəlləri'],
    ['6','Sonra catalog + external_sources','Data metadata'],
    ['7','Sonra glossary + mapping','Biznes layer'],
    ['8','Sonra quality','Quality rules + Soda'],
    ['9','Sonra reports','Reporting'],
    ['10','Seed dataları','RBAC_README.md-dən 14 demo user'],
    ['11','Python mühiti','pip install mysql-connector-python pyyaml'],
    ['12','dq_schedule_simple.py başlat','Background daemon kimi'],
    ['13','Frontend test','index.html-i brauzerdə aç, login ol'],
    ['14','Smoke test','Hər modul üçün 1 CRUD əməliyyatı'],
])

para('Ümumi verification sorğuları:')
code('''-- Bütün cədvəllərdə satır sayı
SELECT table_name, table_rows
FROM information_schema.tables
WHERE table_schema = 'datarover'
ORDER BY table_name;

-- Permission sayı (ən azı 40 olmalı)
SELECT COUNT(*) FROM permissions;

-- Aktiv user sayı
SELECT COUNT(*) FROM users WHERE is_active = 1;

-- Scheduler health
SELECT status, COUNT(*) FROM dq_schedule_jobs GROUP BY status;''')

# ============================================================================
# SAVE
# ============================================================================
out = r'c:\xampp\htdocs\datarover\Technical_Specification.docx'
doc.save(out)
print(f'Saved: {out}')
